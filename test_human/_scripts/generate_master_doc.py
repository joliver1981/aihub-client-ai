"""Generate MASTER_TEST_PLAN.docx — the human-tester's single starting document."""
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "MASTER_TEST_PLAN.docx"


def add_heading_with_color(doc, text, level=1, color=(31, 78, 120)):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = RGBColor(*color)
    return h


def add_kv_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Light Grid Accent 1"
    for i, (k, v) in enumerate(rows):
        cells = table.rows[i].cells
        cells[0].text = k
        cells[1].text = v
        for run in cells[0].paragraphs[0].runs:
            run.bold = True
    return table


def main():
    doc = Document()

    # Title page
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AI Hub — Human Test Plan")
    r.font.size = Pt(28)
    r.bold = True
    r.font.color.rgb = RGBColor(31, 78, 120)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Northwind Outdoor Co. — Retail / Wholesale / Ecommerce scenarios")
    r.font.size = Pt(14)
    r.italic = True
    r.font.color.rgb = RGBColor(102, 102, 102)

    doc.add_paragraph("")
    doc.add_paragraph("")
    add_kv_table(doc, [
        ("Audience",       "Human QA tester (single tester or small QA team)"),
        ("Estimated time", "Approximately one full working day end-to-end"),
        ("Difficulty",     "Mid-level — tester should be familiar with the AI Hub UI"),
        ("Owner",          "AI Hub Quality Team"),
        ("Last updated",   "2026-05-25"),
        ("Fictional co.",  "Northwind Outdoor Co. (omni-channel outdoor gear)"),
    ])
    doc.add_page_break()

    # Table of contents
    add_heading_with_color(doc, "Contents", level=1)
    for item in [
        "1. Purpose & scope",
        "2. Prerequisites (environment, credentials, services)",
        "3. The test suite at a glance",
        "4. How to score — PASS / PARTIAL / FAIL",
        "5. Step-by-step run order",
        "6. Section: Finance",
        "7. Section: Operations",
        "8. Section: IT",
        "9. Section: Planning",
        "10. Section: Data Assistant (NL→SQL)",
        "11. Section: Workflow Builder",
        "12. Section: Compliance / Integrations / MCP smoke",
        "13. Consolidated scoring sheet",
        "14. What to do when something fails",
        "15. Glossary & cleanup checklist",
    ]:
        doc.add_paragraph(item, style="List Number")
    doc.add_page_break()

    # 1. Purpose
    add_heading_with_color(doc, "1. Purpose & scope", level=1)
    doc.add_paragraph(
        "This document is the master human-runnable test plan for the AI "
        "Hub product. Unlike the automated suite in tests_v2/, this plan "
        "is meant to be executed by a human tester who walks through the "
        "UI, uploads documents, asks the agent questions, and judges the "
        "quality of the response against pinned 'fingerprint' answers.")
    doc.add_paragraph(
        "All test material is framed around a fictional omni-channel "
        "retailer/wholesaler/ecommerce business — Northwind Outdoor Co. "
        "— so the questions resemble work that real Finance, Operations, "
        "IT, and Planning teams might do day-to-day. The fixtures contain "
        "specific facts (named SKUs, named vendors, specific $ amounts, "
        "specific dates) so that pass / fail is unambiguous.")

    # 2. Prerequisites
    add_heading_with_color(doc, "2. Prerequisites", level=1)
    doc.add_paragraph("Before starting, confirm the following:")
    for item in [
        "A working AI Hub installation reachable in your browser (typically http://localhost:5001).",
        "An admin or developer login. Default local credentials: admin / admin.",
        "These services are running: main app (5001), document API (5011), vector API (5031), knowledge service (5051), executor (5061).",
        "The Command Center (5091) is required only for browser-driven journey tests; not needed for this plan.",
        "At least 500 MB of disk free for indexed document chunks.",
        "Your browser has DevTools available (Chrome or Edge recommended).",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph(
        "Tip: the easiest way to confirm the stack is healthy is to "
        "navigate to http://localhost:5001 and watch the page load with "
        "no red console errors.")

    # 3. Suite at a glance
    add_heading_with_color(doc, "3. The test suite at a glance", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Section", "What it tests", "Fixtures used", "Est. time"]):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in [
        ("6. Finance",              "Agent Knowledge QA on retail/wholesale finance docs",  "F1 xlsx, F2 pdf, F3 docx", "45 min"),
        ("7. Operations",           "Agent Knowledge QA on inventory + shipping + SOP",      "O1 xlsx (3 tabs), O2 pdf, O3 docx", "45 min"),
        ("8. IT",                   "Agent Knowledge QA on assets + audit + runbook",        "I1 xlsx, I2 pdf, I3 docx", "45 min"),
        ("9. Planning",             "Agent Knowledge QA on forecast + S&OP + policy",         "P1 xlsx, P2 pdf, P3 docx", "45 min"),
        ("10. Data Assistant",      "NL→SQL accuracy + adversarial safety",                  "(none — uses live DB)",     "45 min"),
        ("11. Workflow Builder",    "End-to-end workflow construction and run",              "(none)",                     "60 min"),
        ("12. Smoke tests",         "Compliance, Integrations, MCP basic CRUD",              "(none)",                     "30 min"),
    ]:
        r = table.add_row().cells
        for i, v in enumerate(row):
            r[i].text = v
    doc.add_paragraph("")
    doc.add_paragraph(
        "Total: approximately 5–6 hours of focused testing. Plan for a full "
        "day with breaks and write-up time.")

    # 4. Scoring
    add_heading_with_color(doc, "4. How to score — PASS / PARTIAL / FAIL", level=1)
    doc.add_paragraph("Use the following rubric for every question:")
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Verdict"
    hdr[1].text = "Definition"
    for p in hdr[0].paragraphs + hdr[1].paragraphs:
        for r in p.runs:
            r.bold = True
    for row in [
        ("PASS",    "The answer contains every key fact listed in 'Expected output'. Wording can vary."),
        ("PARTIAL", "Some but not all key facts are present. Or the answer is correct but the agent had to be prompted twice to get it."),
        ("FAIL",    "The answer is wrong, fabricated, or refused without good reason."),
    ]:
        r = table.add_row().cells
        for i, v in enumerate(row):
            r[i].text = v
    doc.add_paragraph(
        "If a question requires a numeric answer, accept any answer within "
        "±1% of the expected value as a PASS — exact match is not required "
        "as long as the rounding is sensible.")
    doc.add_paragraph(
        "For Data Assistant questions, score the generated SQL separately "
        "from the natural-language answer. The SQL should be plausible "
        "(correct shape, correct joins, correct filters) even if the "
        "result row count differs from your expectation due to DB state.")

    # 5. Run order
    add_heading_with_color(doc, "5. Step-by-step run order", level=1)
    doc.add_paragraph("Run sections in this order — the agent-knowledge sections (6–9) can be done in any order, but sections 10–12 should come AFTER, since they don't depend on uploaded fixtures.")
    for step in [
        "Open the AI Hub UI and confirm you can log in.",
        "(Optional) From the Agent Knowledge page, sort agents by name and delete any agents whose names start with 'HUMAN-TEST-' from previous runs.",
        "Section 6 (Finance): create the agent, upload the 3 Finance fixtures, wait for indexing, run the F1/F2/F3 question batches.",
        "Section 7 (Operations): same pattern.",
        "Section 8 (IT): same pattern.",
        "Section 9 (Planning): same pattern.",
        "Section 10 (Data Assistant): no upload required — open the Data Assistant and walk the questions.",
        "Section 11 (Workflow Builder): build the four workflow scenarios and run them.",
        "Section 12 (Smoke tests): CRUD checks on Compliance, Integrations, MCP.",
        "Fill in the consolidated scoring sheet at the end of this document.",
        "Cleanup: delete the four HUMAN-TEST-* agents and the four HUMAN-TEST-W* workflows.",
    ]:
        doc.add_paragraph(step, style="List Number")
    doc.add_paragraph(
        "Tip: indexing is asynchronous and can take 2–4 minutes per "
        "document set. While Finance is indexing, you can prep the next "
        "section in another browser tab.")

    # Sections 6-9: Department test plans
    department_sections = [
        ("6. Section: Finance",     "01_Finance",     "Finance",   "F", "fingerprint facts in this section come from the actual xlsx/pdf/docx files in this folder."),
        ("7. Section: Operations",  "02_Operations",  "Operations","O", "fingerprint facts in this section come from the actual xlsx/pdf/docx files in this folder."),
        ("8. Section: IT",          "03_IT",          "IT",        "I", "fingerprint facts in this section come from the actual xlsx/pdf/docx files in this folder."),
        ("9. Section: Planning",    "04_Planning",    "Planning",  "P", "fingerprint facts in this section come from the actual xlsx/pdf/docx files in this folder."),
    ]
    dept_details = {
        "Finance": {
            "intro": "Finance fixtures cover three classic Finance artifacts: a monthly sales spreadsheet by region/channel/family, a multi-page Q3 P&L statement (PDF with non-repeating headers on continuation pages — a known stress-test for OCR/table extraction), and a vendor payment-terms reference document.",
            "fixtures": [
                ("F1_monthly_sales_by_region.xlsx", "Sales spreadsheet, 60 rows. Tests grand-total reading, region/channel/family aggregation."),
                ("F2_Q3_PnL_statement.pdf",         "P&L statement, multi-page table with header on page 1 ONLY. Stress-tests context retention across pages."),
                ("F3_vendor_payment_terms.docx",    "10-vendor reference doc. Tests narrative + table extraction."),
            ],
            "qcount": 23,
        },
        "Operations": {
            "intro": "Operations fixtures cover three day-to-day Ops artifacts: a multi-tab inventory turnover spreadsheet (three quarters × three warehouses), a multi-page shipping carrier manifest with 90 rows, and the Returns SOP.",
            "fixtures": [
                ("O1_inventory_turnover.xlsx", "Three sheets (Q1, Q2, Q3 FY25), 36 SKUs per sheet. Tests multi-tab indexing and cross-tab time-series reasoning."),
                ("O2_carrier_manifest.pdf",    "90 shipments across 5 pages. Header on page 1 only. Tests row lookup on later pages."),
                ("O3_returns_SOP.docx",        "Returns SOP. Tests channel-specific policy extraction."),
            ],
            "qcount": 23,
        },
        "IT": {
            "intro": "IT fixtures cover three IT artifacts: asset inventory, an annual security audit report with a multi-page findings table, and a runbook for the ERP-to-Shopify integration.",
            "fixtures": [
                ("I1_asset_inventory.xlsx", "70 assets, 5 sites. Tests aggregation + min/max."),
                ("I2_security_audit.pdf",   "24 findings split across 3 pages, no repeating header. Tests cross-page table reading."),
                ("I3_integration_runbook.docx", "Integration runbook. Tests procedural and contact-info extraction."),
            ],
            "qcount": 24,
        },
        "Planning": {
            "intro": "Planning fixtures cover three Planning artifacts: a 12-month demand forecast by family × region, the annual Sales & Operations Plan (S&OP) with a multi-page capacity table, and the capacity & replenishment policy.",
            "fixtures": [
                ("P1_demand_forecast.xlsx", "20 rows (5 families × 4 regions) × 12 months. Tests seasonality detection and aggregation."),
                ("P2_annual_SOP.pdf",       "Multi-page plan with capacity-allocation table spanning two pages, no repeating header."),
                ("P3_capacity_policy.docx", "Capacity policy. Tests formula extraction and policy-rule retrieval."),
            ],
            "qcount": 28,
        },
    }

    for sec_title, folder, dept, prefix, _ in department_sections:
        add_heading_with_color(doc, sec_title, level=1)
        details = dept_details[dept]
        doc.add_paragraph(details["intro"])
        add_heading_with_color(doc, "Fixtures", level=2)
        for fname, desc in details["fixtures"]:
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(fname)
            r.bold = True
            p.add_run(" — " + desc)

        add_heading_with_color(doc, "How to run", level=2)
        for step in [
            f"In the AI Hub UI, create a new agent named 'HUMAN-TEST-{dept}'.",
            f"Upload all three fixtures from test_human/{folder}/fixtures/.",
            "Wait for indexing to complete (2–4 minutes typically).",
            f"Open the per-fixture test plan at test_human/{folder}/test_plan.md and ask each question in order.",
            "Score PASS / PARTIAL / FAIL for each question and record in the consolidated scoring sheet at the end of this doc.",
        ]:
            doc.add_paragraph(step, style="List Number")
        doc.add_paragraph(
            f"The full question list ({details['qcount']} questions) is in "
            f"test_human/{folder}/test_plan.md. Each question is paired "
            f"with the exact expected answer keyed off facts pinned in the "
            f"fixture itself.")
        doc.add_paragraph("")
        p = doc.add_paragraph()
        r = p.add_run("Pass criteria: ")
        r.bold = True
        p.add_run(f"≥ 85% pass across all {details['qcount']} questions.")
        doc.add_page_break()

    # 10. Data Assistant
    add_heading_with_color(doc, "10. Section: Data Assistant (NL→SQL)", level=1)
    doc.add_paragraph(
        "This section tests the AI Hub's natural-language to SQL features "
        "(Data Assistant and Data Explorer v2). No fixture files are "
        "required — the agent queries whatever live database is connected "
        "in your install.")
    add_heading_with_color(doc, "How to run", level=2)
    for step in [
        "Confirm a Data Assistant agent is connected to the sample retail DB.",
        "Open test_human/05_Data_Assistant/data_assistant_questions.md.",
        "Walk through five categories of questions: basic aggregation, ranking, multi-dim grouping, joins, and adversarial safety.",
        "For each question, score both the generated SQL and the natural-language answer.",
        "Pay extra attention to category E (adversarial): refusal behavior on PII and DDL questions is non-negotiable.",
    ]:
        doc.add_paragraph(step, style="List Number")
    p = doc.add_paragraph()
    r = p.add_run("Pass criteria: ")
    r.bold = True
    p.add_run("≥ 80% on categories A–D; 100% on safety questions E1 and E4.")
    doc.add_page_break()

    # 11. Workflow
    add_heading_with_color(doc, "11. Section: Workflow Builder", level=1)
    doc.add_paragraph(
        "This section tests the Workflow Builder end-to-end through four "
        "real-world retail/wholesale/ecommerce scenarios.")
    add_heading_with_color(doc, "Scenarios", level=2)
    for s, desc in [
        ("W1 — Daily Sales Summary", "3-node workflow: query → LLM transform → log. Finance analyst use case."),
        ("W2 — Inventory Low-Stock Alert", "4-node with branching. Ops monitoring use case."),
        ("W3 — Customer Win-back Email Drafting", "3-node with loop. Marketing use case."),
        ("W4 — Approval-gated Workflow", "4-node with human approval pause. Procurement use case."),
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(s).bold = True
        p.add_run(" — " + desc)

    add_heading_with_color(doc, "How to run", level=2)
    for step in [
        "Open test_human/06_Workflow/workflow_scenario.md for the full step-by-step.",
        "Build each workflow in the UI, save with the suggested name, run it, and verify output.",
        "The W4 approval scenario requires a second action: approve / reject via the Approvals UI to resume the paused workflow.",
        "Delete the four HUMAN-TEST-W* workflows when finished.",
    ]:
        doc.add_paragraph(step, style="List Number")
    p = doc.add_paragraph()
    r = p.add_run("Pass criteria: ")
    r.bold = True
    p.add_run("All four scenarios complete with the expected output.")
    doc.add_page_break()

    # 12. Smoke
    add_heading_with_color(doc, "12. Section: Compliance / Integrations / MCP smoke", level=1)
    doc.add_paragraph(
        "A quick CRUD walk-through of the three remaining major surface "
        "areas. The full checklist is in "
        "test_human/07_Smoke_Tests/compliance_integrations_mcp_smoke.md.")
    add_heading_with_color(doc, "What you'll do", level=2)
    for item in [
        "Compliance: create a retailer, view it, delete it. Confirm the schemas list loads.",
        "Integrations: pick a safe template (Azure Blob Storage), save as draft, delete.",
        "MCP: add an MCP server (an echo / test transport), optionally click 'Test Connection', delete.",
        "Watch the browser DevTools console throughout for unexpected red errors.",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    p = doc.add_paragraph()
    r = p.add_run("Pass criteria: ")
    r.bold = True
    p.add_run("Each module's basic CRUD works; no unexplained 5xx; no console errors beyond expected third-party noise.")
    doc.add_page_break()

    # 13. Consolidated scoring sheet
    add_heading_with_color(doc, "13. Consolidated scoring sheet", level=1)
    doc.add_paragraph(
        "Fill in this table at the end. Use the per-section test_plan.md "
        "files for the question-level scores; this table is the rolled-up "
        "summary that goes to the team.")
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Section", "Questions / Scenarios", "Pass", "Partial", "Fail"]):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in [
        ("Finance",     "23", "", "", ""),
        ("Operations",  "23", "", "", ""),
        ("IT",          "24", "", "", ""),
        ("Planning",    "28", "", "", ""),
        ("Data Assistant", "20", "", "", ""),
        ("Workflow",    "4 scenarios", "", "", ""),
        ("Smoke tests", "4 modules", "", "", ""),
        ("TOTAL",       "—", "", "", ""),
    ]:
        r = table.add_row().cells
        for i, v in enumerate(row):
            r[i].text = v

    doc.add_paragraph("")
    add_heading_with_color(doc, "Tester sign-off", level=2)
    add_kv_table(doc, [
        ("Tester name",         ""),
        ("Date run",            ""),
        ("Total elapsed time",  ""),
        ("AI Hub build / commit",""),
        ("Overall verdict",     "PASS / PARTIAL / FAIL"),
        ("Notes",               ""),
    ])
    doc.add_page_break()

    # 14. Troubleshooting
    add_heading_with_color(doc, "14. What to do when something fails", level=1)
    p = doc.add_paragraph()
    p.add_run("Distinguish three failure types:").bold = True
    for item in [
        "Product bug — the agent gives the wrong answer or refuses incorrectly. File a ticket with the question, the answer, the expected output, and the fixture file involved.",
        "Indexing failure — agent answers 'I don't know' on questions that should clearly hit the uploaded doc. Symptom: check the agent's knowledge view and confirm the chunk count is > 0. If not, the indexer didn't process the doc — check skr_trace.txt and the document API logs.",
        "Environment problem — services down, DB connection invalid, etc. These manifest as 5xx errors or 'Connection refused'. Fix the environment and re-run.",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph(
        "When in doubt: capture a screenshot of the agent's response, "
        "the exact question text, and the fixture filename. File the "
        "ticket in the team's issue tracker tagged 'human-test'.")

    # 15. Glossary
    add_heading_with_color(doc, "15. Glossary & cleanup checklist", level=1)
    add_heading_with_color(doc, "Glossary", level=2)
    add_kv_table(doc, [
        ("Northwind Outdoor Co.", "The fictional retail/wholesale/ecommerce company used across all fixtures."),
        ("Agent",                 "An AI Hub assistant with its own knowledge base and tool set."),
        ("Fixture",               "A document (xlsx/pdf/docx) prepared specifically for this test plan, with known facts."),
        ("Fingerprint fact",      "A specific value (number, name, date) that the agent should retrieve verbatim."),
        ("Channel",               "Sales channel — Retail Stores, Wholesale, or Ecommerce."),
        ("S&OP",                  "Sales & Operations Plan — annual production and demand alignment."),
        ("RMA",                   "Return Merchandise Authorization — the number assigned to an approved return."),
        ("DLQ",                   "Dead-letter queue — where failed messages land in the integration pipeline."),
        ("DC",                    "Distribution Center — Northwind operates Western DC, Central DC, Eastern DC."),
    ])
    add_heading_with_color(doc, "Cleanup checklist", level=2)
    for item in [
        "Delete all HUMAN-TEST-* agents in the Agent Knowledge UI.",
        "Delete all HUMAN-TEST-W* workflows in the Workflow Builder.",
        "Delete the test retailer in Compliance (HUMAN-TEST-Retailer-Acme).",
        "Delete the test integration in Integrations (HUMAN-TEST-Int-Blob).",
        "Delete the test MCP server (HUMAN-TEST-MCP-Echo).",
        "If any error screenshots were saved, move them to the team's ticket attachments.",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph("")
    p = doc.add_paragraph()
    r = p.add_run("End of master test plan.")
    r.italic = True
    r.font.color.rgb = RGBColor(102, 102, 102)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(OUT)
    print(f"wrote {OUT}")


if __name__ == "__main__":
    main()
