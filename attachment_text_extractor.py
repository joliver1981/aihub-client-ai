# attachment_text_extractor.py
"""
Attachment Text Extractor
=========================
Extracts text content from various file types for AI agent consumption.

Supported Formats:
- PDF (via PyMuPDF) - Fast extraction for text-based PDFs
- DOCX (via python-docx) - Word documents
- XLSX/XLS (via openpyxl) - Excel spreadsheets
- TXT, CSV, MD, JSON, HTML - Direct read
- Scanned PDFs - Falls back to existing /document/extract endpoint

Dependencies:
    pip install PyMuPDF python-docx openpyxl

Usage:
    from attachment_text_extractor import extract_text_from_attachment
    
    text = extract_text_from_attachment(
        file_bytes=raw_bytes,
        filename="report.pdf",
        content_type="application/pdf",
        max_chars=50000
    )
"""

import io
import os
import logging
from logging.handlers import WatchedFileHandler
import json
import csv
import re
import requests
from typing import Optional, Dict, Tuple, Any
from html.parser import HTMLParser
from CommonUtils import get_document_api_base_url, rotate_logs_on_startup, get_log_path
import config as cfg


# Configure logging
def setup_logging():
    """Configure logging for the agent API"""
    logger = logging.getLogger("AttachmentExtractor")
    log_level_name = os.getenv('LOG_LEVEL', 'DEBUG')
    log_level = getattr(logging, log_level_name, logging.DEBUG)
    logger.setLevel(log_level)
    formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
    handler = WatchedFileHandler(filename=os.getenv('EMAIL_ATTACHMENT_EXTRACTOR', get_log_path('email_attachment_extractor.txt')), encoding='utf-8')
    handler.setFormatter(formatter)
    logger.addHandler(handler)
    
    return logger

rotate_logs_on_startup(os.getenv('EMAIL_ATTACHMENT_EXTRACTOR', get_log_path('email_attachment_extractor.txt')))

logger = setup_logging()

# Configurable size limit (default 100MB)
MAX_FILE_SIZE_BYTES = int(cfg.MAX_ATTACHMENT_SIZE_MB or 100) * 1024 * 1024

# Default max characters to return
DEFAULT_MAX_CHARS = int(cfg.MAX_ATTACHMENT_CHARS or 500000)

# ============================================================================
# HTML Text Extractor (no dependencies)
# ============================================================================

class HTMLTextExtractor(HTMLParser):
    """Simple HTML to text converter."""
    
    def __init__(self):
        super().__init__()
        self.text_parts = []
        self.skip_tags = {'script', 'style', 'head', 'meta', 'link'}
        self.current_tag = None
        
    def handle_starttag(self, tag, attrs):
        self.current_tag = tag.lower()
        if tag.lower() in ('p', 'div', 'br', 'tr', 'li', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
            self.text_parts.append('\n')
            
    def handle_endtag(self, tag):
        self.current_tag = None
        
    def handle_data(self, data):
        if self.current_tag not in self.skip_tags:
            text = data.strip()
            if text:
                self.text_parts.append(text + ' ')
                
    def get_text(self):
        return ''.join(self.text_parts).strip()


def extract_html_text(html_content: str) -> str:
    """Extract text from HTML content."""
    try:
        parser = HTMLTextExtractor()
        parser.feed(html_content)
        text = parser.get_text()
        # Clean up excessive whitespace
        text = re.sub(r'\n\s*\n', '\n\n', text)
        text = re.sub(r' +', ' ', text)
        return text.strip()
    except Exception as e:
        logger.warning(f"HTML parsing failed, returning raw: {e}")
        # Fallback: strip tags with regex
        return re.sub(r'<[^>]+>', ' ', html_content).strip()


# ============================================================================
# PDF Extraction (PyMuPDF)
# ============================================================================
def detect_pdf_type(pdf_bytes: bytes) -> dict:
    """
    Analyze a PDF to determine if it's scanned, text-based, or mixed.
    
    Returns:
        dict with type, confidence, recommendation, and metadata
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        return {'type': 'unknown', 'confidence': 0.0, 'recommendation': 'direct'}
    
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        page_count = len(doc)
        
        if page_count == 0:
            doc.close()
            return {'type': 'empty', 'confidence': 1.0, 'recommendation': 'none'}
        
        # Sample pages (first, last, middle, quarters) - max 5
        sample_indices = [0]
        if page_count > 1:
            sample_indices.append(page_count - 1)
        if page_count > 2:
            sample_indices.append(page_count // 2)
        if page_count > 6:
            sample_indices.extend([page_count // 4, 3 * page_count // 4])
        sample_indices = sorted(set(sample_indices))[:5]
        
        total_text_chars = 0
        total_fonts = 0
        total_images = 0
        scanned_pages = 0
        
        for page_num in sample_indices:
            page = doc[page_num]
            
            # Get text
            text = page.get_text().strip()
            text_chars = len(text)
            total_text_chars += text_chars
            
            # Get fonts
            try:
                total_fonts += len(page.get_fonts())
            except:
                pass
            
            # Get images and calculate coverage
            try:
                images = page.get_images()
                total_images += len(images)
                
                page_area = page.rect.width * page.rect.height
                image_area = 0
                for img in images:
                    try:
                        for rect in page.get_image_rects(img[0]):
                            image_area += rect.width * rect.height
                    except:
                        pass
                image_coverage = image_area / page_area if page_area > 0 else 0
            except:
                image_coverage = 0
            
            # Is this page scanned?
            if (image_coverage > 0.7 and text_chars < 100) or (text_chars < 50 and total_images > 0):
                scanned_pages += 1
        
        doc.close()
        
        # Calculate metrics
        sampled_count = len(sample_indices)
        scanned_ratio = scanned_pages / sampled_count if sampled_count > 0 else 0
        avg_chars_per_page = total_text_chars / sampled_count if sampled_count > 0 else 0
        
        # Determine type and recommendation
        if scanned_ratio >= 0.7:
            pdf_type, confidence, recommendation = 'scanned', scanned_ratio, 'ocr'
        elif scanned_ratio >= 0.3:
            pdf_type, confidence, recommendation = 'mixed', 0.7, 'ocr'
        elif avg_chars_per_page < 50 and total_images > 0:
            pdf_type, confidence, recommendation = 'scanned', 0.7, 'ocr'
        else:
            pdf_type, confidence, recommendation = 'text', 1 - scanned_ratio, 'direct'
        
        return {
            'type': pdf_type,
            'confidence': round(confidence, 2),
            'avg_chars_per_page': round(avg_chars_per_page, 0),
            'page_count': page_count,
            'sampled_pages': sampled_count,
            'scanned_pages_in_sample': scanned_pages,
            'has_fonts': total_fonts > 0,
            'has_images': total_images > 0,
            'recommendation': recommendation
        }
        
    except Exception as e:
        return {'type': 'unknown', 'confidence': 0.0, 'recommendation': 'direct', 'error': str(e)}
    
def extract_pdf_text_direct(file_bytes: bytes, filename: str = "document.pdf") -> tuple:
    """
    Extract text from PDF using PyMuPDF (no OCR).
    
    Returns:
        Tuple of (extracted_text, page_count)
    """
    import fitz
    
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    text_parts = []
    
    for page_num, page in enumerate(doc, 1):
        page_text = page.get_text("text")
        if page_text.strip():
            text_parts.append(f"--- Page {page_num} ---\n{page_text}")
    
    page_count = len(doc)
    doc.close()
    
    return "\n\n".join(text_parts).strip(), page_count


def extract_pdf_text(file_bytes: bytes, filename: str = "document.pdf", 
                     allow_ocr: bool = True) -> tuple:
    """
    Extract text from PDF with automatic scanned detection.
    
    Returns:
        Tuple of (text, analysis_dict)
    """
    # Step 1: Detect PDF type
    analysis = detect_pdf_type(file_bytes)
    analysis['extraction_method'] = None
    analysis['ocr_used'] = False
    
    recommendation = analysis.get('recommendation', 'direct')
    
    # Step 2: Extract based on recommendation
    if recommendation == 'direct' or not allow_ocr:
        # Text-based PDF: use PyMuPDF directly
        text, page_count = extract_pdf_text_direct(file_bytes, filename)
        analysis['extraction_method'] = 'pymupdf_direct'
        analysis['page_count'] = page_count
        
        # Safety check: if very little text but has images, try OCR
        if len(text.strip()) < 200 and analysis.get('has_images') and allow_ocr:
            try:
                ocr_text = extract_scanned_pdf_via_api(file_bytes, filename)
                if len(ocr_text.strip()) > len(text.strip()):
                    text = ocr_text
                    analysis['extraction_method'] = 'ocr_api'
                    analysis['ocr_used'] = True
            except:
                pass  # Keep direct text
        
        return text, analysis
        
    else:  # recommendation == 'ocr' (scanned or mixed)
        # Use OCR API
        try:
            text = extract_scanned_pdf_via_api(file_bytes, filename)
            analysis['extraction_method'] = 'ocr_api'
            analysis['ocr_used'] = True
            return text, analysis
        except Exception as e:
            # Fallback to direct extraction
            text, page_count = extract_pdf_text_direct(file_bytes, filename)
            analysis['extraction_method'] = 'pymupdf_direct_fallback'
            analysis['ocr_error'] = str(e)
            if not text.strip():
                text = f"[PDF appears scanned but OCR failed: {e}]"
            return text, analysis


# ============================================================================
# Word Document Extraction (python-docx)
# ============================================================================

def extract_docx_text(file_bytes: bytes, filename: str = "document.docx") -> str:
    """Extract text from Word document."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx not installed. Run: pip install python-docx")
    
    try:
        doc = Document(io.BytesIO(file_bytes))
        text_parts = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extract tables
        for table_idx, table in enumerate(doc.tables, 1):
            table_text = [f"\n[Table {table_idx}]"]
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    table_text.append(row_text)
            if len(table_text) > 1:
                text_parts.append("\n".join(table_text))
        
        return "\n\n".join(text_parts).strip()
        
    except Exception as e:
        logger.error(f"DOCX extraction failed for '{filename}': {e}")
        raise


# ============================================================================
# Excel Extraction (openpyxl)
# ============================================================================

def extract_xlsx_text(file_bytes: bytes, filename: str = "spreadsheet.xlsx") -> str:
    """Extract text from Excel spreadsheet as markdown-style tables."""
    try:
        import openpyxl
    except ImportError:
        raise ImportError("openpyxl not installed. Run: pip install openpyxl")
    
    try:
        wb = openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
        text_parts = []
        
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            rows = list(sheet.iter_rows(values_only=True))
            
            if not rows:
                continue
            
            # Filter out completely empty rows
            rows = [r for r in rows if any(cell is not None for cell in r)]
            
            if not rows:
                continue
            
            sheet_text = [f"\n## Sheet: {sheet_name}\n"]
            
            # Find max columns with data
            max_cols = max(len(row) for row in rows) if rows else 0
            
            # Build markdown table
            for row_idx, row in enumerate(rows[:500]):  # Limit to 500 rows
                # Pad row to max_cols
                padded_row = list(row) + [None] * (max_cols - len(row))
                
                # Convert cells to strings
                cell_strs = []
                for cell in padded_row:
                    if cell is None:
                        cell_strs.append("")
                    elif isinstance(cell, (int, float)):
                        cell_strs.append(str(cell))
                    else:
                        cell_strs.append(str(cell).replace("|", "\\|").replace("\n", " "))
                
                row_text = "| " + " | ".join(cell_strs) + " |"
                sheet_text.append(row_text)
                
                # Add header separator after first row
                if row_idx == 0:
                    separator = "| " + " | ".join(["---"] * max_cols) + " |"
                    sheet_text.append(separator)
            
            if len(rows) > 500:
                sheet_text.append(f"\n... (truncated, {len(rows) - 500} more rows)")
            
            text_parts.append("\n".join(sheet_text))
        
        wb.close()
        return "\n".join(text_parts).strip()
        
    except Exception as e:
        logger.error(f"XLSX extraction failed for '{filename}': {e}")
        raise


# ============================================================================
# Plain Text / CSV Extraction
# ============================================================================

def extract_text_file(file_bytes: bytes, filename: str = "file.txt") -> str:
    """Extract text from plain text files."""
    # Try different encodings
    encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252', 'iso-8859-1']
    
    for encoding in encodings:
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    
    # Last resort: decode with errors ignored
    return file_bytes.decode('utf-8', errors='replace')


def extract_csv_text(file_bytes: bytes, filename: str = "file.csv") -> str:
    """Extract CSV as markdown table."""
    try:
        text = extract_text_file(file_bytes, filename)
        lines = text.strip().split('\n')
        
        if not lines:
            return ""
        
        # Detect delimiter
        first_line = lines[0]
        delimiter = ','
        if '\t' in first_line and first_line.count('\t') > first_line.count(','):
            delimiter = '\t'
        elif ';' in first_line and first_line.count(';') > first_line.count(','):
            delimiter = ';'
        
        # Parse CSV
        reader = csv.reader(io.StringIO(text), delimiter=delimiter)
        rows = list(reader)
        
        if not rows:
            return text
        
        # Build markdown table
        output = []
        max_cols = max(len(row) for row in rows)
        
        for row_idx, row in enumerate(rows[:500]):
            padded = row + [''] * (max_cols - len(row))
            cleaned = [cell.replace('|', '\\|').replace('\n', ' ') for cell in padded]
            output.append("| " + " | ".join(cleaned) + " |")
            
            if row_idx == 0:
                output.append("| " + " | ".join(["---"] * max_cols) + " |")
        
        if len(rows) > 500:
            output.append(f"\n... (truncated, {len(rows) - 500} more rows)")
        
        return "\n".join(output)
        
    except Exception as e:
        logger.warning(f"CSV parsing failed, returning raw text: {e}")
        return extract_text_file(file_bytes, filename)


def extract_json_text(file_bytes: bytes, filename: str = "file.json") -> str:
    """Extract JSON with pretty formatting."""
    try:
        text = extract_text_file(file_bytes, filename)
        data = json.loads(text)
        return json.dumps(data, indent=2, ensure_ascii=False)
    except json.JSONDecodeError:
        return extract_text_file(file_bytes, filename)


# ============================================================================
# Scanned PDF Extraction (via existing document API)
# ============================================================================

def extract_scanned_pdf_via_api(file_bytes: bytes, filename: str = "scanned.pdf") -> str:
    """
    Extract text from scanned PDF using the existing /document/extract endpoint.
    This uses OCR and AI extraction.
    
    Note: The document API expects a file path, so we save to a temp file first.
    """
    import tempfile
    
    # Handle None filename
    if not filename:
        filename = "scanned.pdf"
    
    temp_path = None
    try:
        # Get document API base URL
        base_url = get_document_api_base_url()
        if not base_url:
            base_url = "http://localhost:5000"  # Default fallback
        
        endpoint = f"{base_url.rstrip('/')}/document/extract_text"
        
        logger.info(f"Calling document extraction API for scanned PDF: {filename}")
        
        # Save bytes to a temp file (the API expects a file path, not bytes)
        # Use a suffix to preserve the file extension for allowed_file() check
        suffix = os.path.splitext(filename)[1] or '.pdf'
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(file_bytes)
            temp_path = tmp.name
        
        logger.debug(f"Saved PDF to temp file: {temp_path}")
        
        # Prepare the request with file path
        data = {
            'filePath': temp_path,
            'document_type': 'auto',
            'force_ai_extraction': 'true',
            'use_batch_processing': 'true',
            'batch_size': '3',
            'extract_fields': 'false',  # Just get text, not structured data
            'do_not_store': 'true'  # Don't save to documents table
        }
        
        response = requests.post(
            endpoint,
            data=data,
            timeout=240  # 2 minute timeout for OCR
        )
        
        if response.status_code == 200:
            result = response.json()
            if result.get('status') == 'success':
                text = result.get('text', '')
                if not text:
                    text = result.get('document_text', '')
                if not text:
                    # Concatenate all string values
                    text_parts = []
                    for key, value in extracted.items():
                        if isinstance(value, str) and value.strip():
                            text_parts.append(f"{key}: {value}")
                    text = "\n".join(text_parts)
                return text
            else:
                raise Exception(result.get('message', 'Extraction failed'))
        else:
            raise Exception(f"API returned {response.status_code}")
            
    except requests.exceptions.RequestException as e:
        logger.error(f"Document API request failed: {e}")
        raise Exception(f"Could not connect to document extraction service: {e}")
    except Exception as e:
        logger.error(f"Scanned PDF extraction failed: {e}")
        raise
    finally:
        # Clean up temp file
        if temp_path and os.path.exists(temp_path):
            try:
                os.unlink(temp_path)
                logger.debug(f"Cleaned up temp file: {temp_path}")
            except Exception as cleanup_err:
                logger.warning(f"Could not delete temp file {temp_path}: {cleanup_err}")


# ============================================================================
# Main Extraction Function
# ============================================================================

# Map content types to extractors
CONTENT_TYPE_MAP = {
    # PDF
    'application/pdf': 'pdf',
    
    # Word
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
    'application/msword': 'doc',  # Old format - limited support
    
    # Excel
    'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': 'xlsx',
    'application/vnd.ms-excel': 'xls',
    
    # Text
    'text/plain': 'txt',
    'text/csv': 'csv',
    'text/tab-separated-values': 'csv',
    'text/markdown': 'txt',
    'text/html': 'html',
    'application/json': 'json',
    'application/xml': 'txt',
    'text/xml': 'txt',
}

# Map file extensions to extractors (fallback)
EXTENSION_MAP = {
    '.pdf': 'pdf',
    '.docx': 'docx',
    '.doc': 'doc',
    '.xlsx': 'xlsx',
    '.xls': 'xls',
    '.txt': 'txt',
    '.csv': 'csv',
    '.tsv': 'csv',
    '.md': 'txt',
    '.markdown': 'txt',
    '.html': 'html',
    '.htm': 'html',
    '.json': 'json',
    '.xml': 'txt',
    '.log': 'txt',
    '.ini': 'txt',
    '.cfg': 'txt',
    '.yaml': 'txt',
    '.yml': 'txt',
}


def detect_file_type(filename: str, content_type: str = None) -> str:
    """Detect the file type for extraction."""
    # Try content type first
    if content_type:
        content_type_lower = content_type.lower().split(';')[0].strip()
        if content_type_lower in CONTENT_TYPE_MAP:
            return CONTENT_TYPE_MAP[content_type_lower]
    
    # Fall back to extension
    if filename:
        ext = os.path.splitext(filename.lower())[1]
        if ext in EXTENSION_MAP:
            return EXTENSION_MAP[ext]
    
    # Default to text
    return 'txt'


def truncate_text(text: str, max_chars: int, filename: str = None) -> str:
    """Truncate text to max_chars with indicator."""
    if len(text) <= max_chars:
        return text
    
    truncated = text[:max_chars]
    remaining = len(text) - max_chars
    
    indicator = f"\n\n[... Content truncated. {remaining:,} more characters not shown"
    if filename:
        indicator += f" from '{filename}'"
    indicator += ". Use a smaller portion or specific search if needed.]"
    
    return truncated + indicator


def extract_text_from_attachment(
    file_bytes: bytes,
    filename: str,
    content_type: str = None,
    max_chars: int = DEFAULT_MAX_CHARS,
    allow_ocr_fallback: bool = True
) -> Dict[str, Any]:
    """
    Extract text from an attachment.
    
    Args:
        file_bytes: Raw file bytes
        filename: Original filename
        content_type: MIME type (optional)
        max_chars: Maximum characters to return
        allow_ocr_fallback: If True, use OCR API for scanned PDFs
        
    Returns:
        Dict with:
            - success: bool
            - text: Extracted text (if successful)
            - truncated: bool - whether text was truncated
            - original_length: int - original text length before truncation
            - file_type: str - detected file type
            - extraction_method: str - method used (pymupdf, docx, etc.)
            - error: str (if failed)
    """
    result = {
        'success': False,
        'text': '',
        'truncated': False,
        'original_length': 0,
        'file_type': None,
        'extraction_method': None,
        'error': None
    }
    
    try:
        # Check file size
        if len(file_bytes) > MAX_FILE_SIZE_BYTES:
            result['error'] = f"File too large ({len(file_bytes) / 1024 / 1024:.1f}MB). Maximum size is {MAX_FILE_SIZE_BYTES / 1024 / 1024:.0f}MB."
            return result
        
        # Detect file type
        file_type = detect_file_type(filename, content_type)
        result['file_type'] = file_type
        
        text = ""
        
        # Extract based on type
        if file_type == 'pdf':
            text, pdf_analysis = extract_pdf_text(file_bytes, filename, allow_ocr=allow_ocr_fallback)
            result['extraction_method'] = pdf_analysis.get('extraction_method', 'pymupdf')
            result['pdf_analysis'] = pdf_analysis
            
            if pdf_analysis.get('ocr_used'):
                result['ocr_used'] = True
            
            pdf_type = pdf_analysis.get('type', 'unknown')
            if pdf_type in ('scanned', 'mixed'):
                result['note'] = f'PDF detected as {pdf_type}'
                        
        elif file_type == 'docx':
            text = extract_docx_text(file_bytes, filename)
            result['extraction_method'] = 'python-docx'
            
        elif file_type == 'doc':
            # Old Word format - limited support
            result['error'] = "Old .doc format not fully supported. Please convert to .docx for best results."
            result['text'] = "[Old .doc format - conversion to .docx recommended]"
            return result
            
        elif file_type == 'xlsx':
            text = extract_xlsx_text(file_bytes, filename)
            result['extraction_method'] = 'openpyxl'
            
        elif file_type == 'xls':
            # Old Excel format - limited support
            result['error'] = "Old .xls format not fully supported. Please convert to .xlsx for best results."
            result['text'] = "[Old .xls format - conversion to .xlsx recommended]"
            return result
            
        elif file_type == 'csv':
            text = extract_csv_text(file_bytes, filename)
            result['extraction_method'] = 'csv'
            
        elif file_type == 'json':
            text = extract_json_text(file_bytes, filename)
            result['extraction_method'] = 'json'
            
        elif file_type == 'html':
            raw_text = extract_text_file(file_bytes, filename)
            text = extract_html_text(raw_text)
            result['extraction_method'] = 'html'
            
        else:  # txt and fallback
            text = extract_text_file(file_bytes, filename)
            result['extraction_method'] = 'text'
        
        result['original_length'] = len(text)
        
        # Truncate if needed
        if len(text) > max_chars:
            text = truncate_text(text, max_chars, filename)
            result['truncated'] = True
        
        result['text'] = text
        result['success'] = True
        
    except ImportError as e:
        result['error'] = f"Missing dependency: {e}"
        logger.error(f"Import error during extraction: {e}")
        
    except Exception as e:
        result['error'] = f"Extraction failed: {str(e)}"
        logger.error(f"Extraction error for '{filename}': {e}", exc_info=True)
    
    return result


def get_supported_formats() -> Dict[str, str]:
    """Return a dict of supported file formats and their descriptions."""
    return {
        'PDF': 'Adobe PDF documents (text-based and scanned)',
        'DOCX': 'Microsoft Word documents',
        'XLSX': 'Microsoft Excel spreadsheets',
        'CSV': 'Comma-separated values',
        'TXT': 'Plain text files',
        'JSON': 'JSON data files',
        'HTML': 'HTML web pages',
        'MD': 'Markdown files',
        'XML': 'XML documents'
    }


# ============================================================================
# Exports
# ============================================================================

__all__ = [
    'extract_text_from_attachment',
    'detect_file_type',
    'get_supported_formats',
    'DEFAULT_MAX_CHARS',
    'MAX_FILE_SIZE_BYTES'
]
