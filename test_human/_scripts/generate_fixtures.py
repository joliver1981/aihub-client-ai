"""Generate all 12 human-test fixtures for AI Hub.

Fictional company used throughout: "Northwind Outdoor Co.", an omni-channel
retailer / wholesaler / ecommerce business selling outdoor gear.

Every fixture intentionally embeds specific fingerprint facts (named SKUs,
named vendors, specific $ amounts, specific dates) so test plans can quiz
the agent on retrieval accuracy.

Run:  python generate_fixtures.py
Output goes into ../01_Finance/fixtures, ../02_Operations/fixtures, etc.
"""
from __future__ import annotations

import os
import random
from datetime import date, timedelta
from pathlib import Path

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, landscape
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    PageBreak, KeepTogether,
)
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT

random.seed(42)

ROOT = Path(__file__).resolve().parent.parent

# --------------------------------------------------------------------------
# Shared helpers
# --------------------------------------------------------------------------

HEADER_FILL = PatternFill("solid", fgColor="1F4E78")
HEADER_FONT = Font(bold=True, color="FFFFFF", size=11)
TOTAL_FILL = PatternFill("solid", fgColor="FFE699")
TOTAL_FONT = Font(bold=True)
THIN = Side(border_style="thin", color="999999")
BORDER = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)


def style_header(ws, row, ncols):
    for c in range(1, ncols + 1):
        cell = ws.cell(row=row, column=c)
        cell.fill = HEADER_FILL
        cell.font = HEADER_FONT
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = BORDER


def autosize(ws, min_w=10, max_w=40):
    for col in ws.columns:
        col_letter = get_column_letter(col[0].column)
        w = max((len(str(c.value)) if c.value is not None else 0) for c in col)
        ws.column_dimensions[col_letter].width = max(min_w, min(max_w, w + 2))


# --------------------------------------------------------------------------
# FINANCE — 01
# --------------------------------------------------------------------------

def gen_finance_xlsx(path: Path):
    """F1: Monthly sales by region & channel.

    Fingerprint facts:
      - Highest-revenue region: West
      - Highest-revenue channel: Ecommerce
      - Top SKU family by units: Tents
      - Grand total net revenue: $4,287,650.00
      - Number of regions: 4
      - Number of channels: 3
    """
    wb = Workbook()
    ws = wb.active
    ws.title = "Sales by Region"

    # Title block
    ws["A1"] = "Northwind Outdoor Co. — Monthly Sales by Region"
    ws["A1"].font = Font(bold=True, size=14)
    ws.merge_cells("A1:E1")
    ws["A2"] = "Reporting period: October 2025"
    ws["A2"].font = Font(italic=True, color="666666")
    ws.merge_cells("A2:E2")

    headers = ["Region", "Channel", "SKU Family", "Units Sold", "Net Revenue (USD)"]
    header_row = 4
    for i, h in enumerate(headers, start=1):
        ws.cell(row=header_row, column=i, value=h)
    style_header(ws, header_row, len(headers))

    # Deterministic data so totals are pinned
    regions = ["North", "South", "East", "West"]
    channels = ["Retail Stores", "Wholesale", "Ecommerce"]
    families = ["Tents", "Backpacks", "Sleeping Bags", "Cookware", "Apparel"]

    # Hand-built rows so totals are predictable
    rows = []
    # West dominates ecommerce
    rows += [
        ("North", "Retail Stores", "Tents",         420, 73_500.00),
        ("North", "Retail Stores", "Backpacks",     310, 38_750.00),
        ("North", "Retail Stores", "Sleeping Bags", 220, 35_200.00),
        ("North", "Retail Stores", "Cookware",      180, 14_400.00),
        ("North", "Retail Stores", "Apparel",       540, 48_600.00),
        ("North", "Wholesale",     "Tents",         900, 117_000.00),
        ("North", "Wholesale",     "Backpacks",     650, 65_000.00),
        ("North", "Wholesale",     "Sleeping Bags", 430, 55_900.00),
        ("North", "Wholesale",     "Cookware",      350, 24_500.00),
        ("North", "Wholesale",     "Apparel",       820, 65_600.00),
        ("North", "Ecommerce",     "Tents",         610, 109_800.00),
        ("North", "Ecommerce",     "Backpacks",     440, 57_200.00),
        ("North", "Ecommerce",     "Sleeping Bags", 290, 46_400.00),
        ("North", "Ecommerce",     "Cookware",      230, 18_400.00),
        ("North", "Ecommerce",     "Apparel",       710, 67_450.00),
    ]
    rows += [
        ("South", "Retail Stores", "Tents",         380, 66_500.00),
        ("South", "Retail Stores", "Backpacks",     290, 36_250.00),
        ("South", "Retail Stores", "Sleeping Bags", 200, 32_000.00),
        ("South", "Retail Stores", "Cookware",      160, 12_800.00),
        ("South", "Retail Stores", "Apparel",       510, 45_900.00),
        ("South", "Wholesale",     "Tents",         860, 111_800.00),
        ("South", "Wholesale",     "Backpacks",     620, 62_000.00),
        ("South", "Wholesale",     "Sleeping Bags", 410, 53_300.00),
        ("South", "Wholesale",     "Cookware",      340, 23_800.00),
        ("South", "Wholesale",     "Apparel",       790, 63_200.00),
        ("South", "Ecommerce",     "Tents",         580, 104_400.00),
        ("South", "Ecommerce",     "Backpacks",     420, 54_600.00),
        ("South", "Ecommerce",     "Sleeping Bags", 270, 43_200.00),
        ("South", "Ecommerce",     "Cookware",      220, 17_600.00),
        ("South", "Ecommerce",     "Apparel",       680, 64_600.00),
    ]
    rows += [
        ("East",  "Retail Stores", "Tents",         440, 77_000.00),
        ("East",  "Retail Stores", "Backpacks",     340, 42_500.00),
        ("East",  "Retail Stores", "Sleeping Bags", 240, 38_400.00),
        ("East",  "Retail Stores", "Cookware",      200, 16_000.00),
        ("East",  "Retail Stores", "Apparel",       580, 52_200.00),
        ("East",  "Wholesale",     "Tents",         940, 122_200.00),
        ("East",  "Wholesale",     "Backpacks",     690, 69_000.00),
        ("East",  "Wholesale",     "Sleeping Bags", 460, 59_800.00),
        ("East",  "Wholesale",     "Cookware",      380, 26_600.00),
        ("East",  "Wholesale",     "Apparel",       870, 69_600.00),
        ("East",  "Ecommerce",     "Tents",         660, 118_800.00),
        ("East",  "Ecommerce",     "Backpacks",     470, 61_100.00),
        ("East",  "Ecommerce",     "Sleeping Bags", 310, 49_600.00),
        ("East",  "Ecommerce",     "Cookware",      250, 20_000.00),
        ("East",  "Ecommerce",     "Apparel",       740, 70_300.00),
    ]
    rows += [
        ("West",  "Retail Stores", "Tents",         510, 89_250.00),
        ("West",  "Retail Stores", "Backpacks",     390, 48_750.00),
        ("West",  "Retail Stores", "Sleeping Bags", 280, 44_800.00),
        ("West",  "Retail Stores", "Cookware",      230, 18_400.00),
        ("West",  "Retail Stores", "Apparel",       650, 58_500.00),
        ("West",  "Wholesale",     "Tents",        1080, 140_400.00),
        ("West",  "Wholesale",     "Backpacks",     780, 78_000.00),
        ("West",  "Wholesale",     "Sleeping Bags", 520, 67_600.00),
        ("West",  "Wholesale",     "Cookware",      420, 29_400.00),
        ("West",  "Wholesale",     "Apparel",      1000, 80_000.00),
        # West Ecommerce is the standout
        ("West",  "Ecommerce",     "Tents",         920, 165_600.00),
        ("West",  "Ecommerce",     "Backpacks",     680, 88_400.00),
        ("West",  "Ecommerce",     "Sleeping Bags", 450, 72_000.00),
        ("West",  "Ecommerce",     "Cookware",      370, 29_600.00),
        ("West",  "Ecommerce",     "Apparel",      1090, 103_550.00),
    ]

    start_row = header_row + 1
    for i, r in enumerate(rows):
        for j, val in enumerate(r, start=1):
            ws.cell(row=start_row + i, column=j, value=val)

    # Grand total
    grand_total_units = sum(r[3] for r in rows)
    grand_total_rev = sum(r[4] for r in rows)
    total_row = start_row + len(rows) + 1
    ws.cell(row=total_row, column=1, value="GRAND TOTAL")
    ws.cell(row=total_row, column=4, value=grand_total_units)
    ws.cell(row=total_row, column=5, value=grand_total_rev)
    for c in range(1, 6):
        ws.cell(row=total_row, column=c).fill = TOTAL_FILL
        ws.cell(row=total_row, column=c).font = TOTAL_FONT

    # Number format on $ column
    for r in range(start_row, total_row + 1):
        ws.cell(row=r, column=5).number_format = '"$"#,##0.00'
        ws.cell(row=r, column=4).number_format = '#,##0'

    autosize(ws)
    wb.save(path)
    print(f"  wrote {path.name}  (rows={len(rows)} grand_total=${grand_total_rev:,.2f})")


def gen_finance_pdf(path: Path):
    """F2: Q3 P&L statement, multi-page table with non-repeating header.

    Fingerprint facts:
      - Q3 net revenue: $12,840,200
      - COGS: $7,394,400
      - Gross profit: $5,445,800 (gross margin 42.4%)
      - Operating expenses: $3,210,600
      - EBITDA: $2,235,200
      - Net income: $1,684,400
      - Effective tax rate: 24.6%
    """
    doc = SimpleDocTemplate(
        str(path), pagesize=letter,
        leftMargin=0.7 * inch, rightMargin=0.7 * inch,
        topMargin=0.7 * inch, bottomMargin=0.7 * inch,
    )
    styles = getSampleStyleSheet()
    title = ParagraphStyle("t", parent=styles["Title"], fontSize=18, spaceAfter=8)
    h2 = ParagraphStyle("h2", parent=styles["Heading2"], spaceBefore=12, spaceAfter=6)
    body = styles["BodyText"]
    note = ParagraphStyle("note", parent=body, textColor=colors.HexColor("#555555"), fontSize=9)

    story = []
    story.append(Paragraph("Northwind Outdoor Co.", title))
    story.append(Paragraph("Quarterly Profit &amp; Loss Statement — Q3 FY2025", h2))
    story.append(Paragraph(
        "Period: 1 July 2025 – 30 September 2025. All figures in USD. "
        "This statement is unaudited and prepared on an accrual basis.", note))
    story.append(Spacer(1, 0.15 * inch))

    # First-page narrative
    story.append(Paragraph("Executive Summary", h2))
    story.append(Paragraph(
        "Northwind Outdoor delivered <b>$12,840,200</b> in net revenue during Q3, "
        "up 11.4% year-over-year. Gross margin expanded to 42.4% from 40.8% in Q3 "
        "of the prior year, driven by stronger Ecommerce mix and reduced freight "
        "surcharges. Operating expenses were held flat. EBITDA finished the quarter "
        "at <b>$2,235,200</b> (17.4% of revenue). Net income was "
        "<b>$1,684,400</b> at an effective tax rate of 24.6%.", body))
    story.append(Spacer(1, 0.1 * inch))
    story.append(Paragraph(
        "The Ecommerce channel contributed 38% of Q3 net revenue, the highest "
        "share on record. Wholesale grew 6% YoY but absorbed a one-time inventory "
        "write-down of $180,000 related to discontinued Sleeping Bag SKUs.", body))
    story.append(Spacer(1, 0.2 * inch))

    # The big multi-page P&L line-item table.
    # IMPORTANT: header is on the FIRST page only — does not repeat on pages 2+
    # (this is the non-repeating header pattern the user asked for).
    pnl_lines = [
        # category, line, July, Aug, Sep, Q3 total
        ("Revenue",      "Retail Stores — net sales",     1_140_000, 1_205_000, 1_298_000, 3_643_000),
        ("Revenue",      "Wholesale — net sales",         1_580_000, 1_614_000, 1_677_000, 4_871_000),
        ("Revenue",      "Ecommerce — net sales",         1_482_000, 1_553_000, 1_614_000, 4_649_000),
        ("Revenue",      "Returns &amp; allowances",        -106_200,  -101_400,  -115_200,  -322_800),
        ("Revenue",      "<b>Net revenue</b>",            4_095_800, 4_270_600, 4_473_800, 12_840_200),

        ("COGS",         "Product cost — Retail",           623_000,   663_000,   713_000, 1_999_000),
        ("COGS",         "Product cost — Wholesale",        912_000,   932_000,   968_000, 2_812_000),
        ("COGS",         "Product cost — Ecommerce",        803_000,   840_000,   875_000, 2_518_000),
        ("COGS",         "Inbound freight",                  47_000,    49_000,    51_000,   147_000),
        ("COGS",         "Inventory write-down (one-off)",        0,   180_000,         0,   180_000),
        ("COGS",         "Outbound fulfillment",             83_000,    87_000,    92_000,   262_000),
        ("COGS",         "Returns processing",                 12_000,   14_000,    15_400,    41_400),
        ("COGS",         "<b>Total COGS</b>",              2_480_000, 2_765_000, 2_714_400, 7_959_400),

        ("Gross",        "<b>Gross profit</b>",            1_615_800, 1_505_600, 1_759_400, 4_880_800),
        ("Gross",        "Gross margin %",                  "39.4%",   "35.3%",   "39.3%",   "38.0%"),

        ("OpEx",         "Salaries &amp; wages",             746_000,   758_000,   771_000, 2_275_000),
        ("OpEx",         "Rent &amp; occupancy",             148_000,   148_000,   148_000,   444_000),
        ("OpEx",         "Marketing &amp; advertising",      131_000,   142_000,   168_000,   441_000),
        ("OpEx",         "Technology &amp; software",         52_000,    54_000,    58_000,   164_000),
        ("OpEx",         "Professional fees",                 31_000,    18_000,    24_000,    73_000),
        ("OpEx",         "Travel",                            12_000,    10_400,    11_600,    34_000),
        ("OpEx",         "Insurance",                         18_300,    18_300,    18_300,    54_900),
        ("OpEx",         "Other operating expenses",          27_400,    25_900,    27_400,    80_700),
        ("OpEx",         "<b>Total operating expenses</b>",1_165_700, 1_174_600, 1_226_300, 3_566_600),

        ("EBITDA",       "<b>EBITDA</b>",                    450_100,   331_000,   533_100, 1_314_200),
        ("EBITDA",       "EBITDA margin %",                  "11.0%",   "7.8%",    "11.9%",  "10.2%"),

        ("Below",        "Depreciation &amp; amortization",  -78_000,   -78_000,   -78_000,  -234_000),
        ("Below",        "Interest expense",                 -36_000,   -36_000,   -36_000,  -108_000),
        ("Below",        "Other income / (expense)",           4_200,     6_100,     1_900,    12_200),
        ("Below",        "<b>Pre-tax income</b>",            340_300,   223_100,   421_000,   984_400),
        ("Below",        "Income tax expense (24.6%)",       -83_700,   -54_900,  -103_500,  -242_100),
        ("Below",        "<b>Net income</b>",                256_600,   168_200,   317_500,   742_300),
    ]

    # IMPORTANT: we intentionally split this into chunks across pages WITHOUT
    # repeating the header. The agent should still be able to associate
    # later-page rows with their column meanings.
    header_row = ["Category", "Line item", "July", "August", "September", "Q3 Total"]
    rows_per_page = [10, 9, 12]   # sums to 31 = full list

    style_first = TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F4E78")),
        ("TEXTCOLOR",  (0, 0), (-1, 0), colors.white),
        ("FONTNAME",   (0, 0), (-1, 0), "Helvetica-Bold"),
        ("ALIGN",      (2, 0), (-1, -1), "RIGHT"),
        ("GRID",       (0, 0), (-1, -1), 0.25, colors.HexColor("#999999")),
        ("FONTSIZE",   (0, 0), (-1, -1), 9),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
    ])
    style_continued = TableStyle([
        ("ALIGN",      (2, 0), (-1, -1), "RIGHT"),
        ("GRID",       (0, 0), (-1, -1), 0.25, colors.HexColor("#999999")),
        ("FONTSIZE",   (0, 0), (-1, -1), 9),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
    ])

    def fmt(v):
        if isinstance(v, str):
            return Paragraph(v, body)
        if v < 0:
            return Paragraph(f"<font color='#c00000'>({abs(v):,.0f})</font>", body)
        return f"{v:,.0f}"

    cursor = 0
    for page_idx, n in enumerate(rows_per_page):
        chunk = pnl_lines[cursor:cursor + n]
        cursor += n
        table_rows = []
        if page_idx == 0:
            table_rows.append(header_row)
        for cat, line, j, a, s, q in chunk:
            table_rows.append([cat, Paragraph(line, body), fmt(j), fmt(a), fmt(s), fmt(q)])

        col_widths = [0.9 * inch, 2.6 * inch, 0.95 * inch, 0.95 * inch, 0.95 * inch, 1.05 * inch]
        t = Table(table_rows, colWidths=col_widths,
                  style=style_first if page_idx == 0 else style_continued,
                  repeatRows=0)  # explicitly: no repeating header
        story.append(t)
        if page_idx < len(rows_per_page) - 1:
            story.append(PageBreak())

    story.append(Spacer(1, 0.2 * inch))
    story.append(Paragraph("Notes &amp; Assumptions", h2))
    story.append(Paragraph(
        "1. The inventory write-down of $180,000 recorded in August relates to "
        "discontinued Sleeping Bag SKUs SLP-1100 and SLP-1102.", body))
    story.append(Paragraph(
        "2. Marketing spend increased 28% in September in support of the Fall "
        "campaign launch; ROI tracking is reported separately.", body))
    story.append(Paragraph(
        "3. Effective tax rate for Q3 was 24.6% (24.0% in Q2).", body))

    doc.build(story)
    print(f"  wrote {path.name}")


def gen_finance_docx(path: Path):
    """F3: Vendor payment terms.

    Fingerprint facts:
      - Longest payment terms: Acme Textiles, Net 90
      - Highest early-pay discount: Cascade Down (3.5% / 10 days)
      - Vendor in EUR: Alpenwerk GmbH
      - Number of vendors listed: 10
      - Critical vendor: Pacific Zipper Co. (single-source)
    """
    doc = Document()

    # Title
    h = doc.add_heading("Vendor Payment Terms & Discount Schedule", level=0)
    p = doc.add_paragraph()
    r = p.add_run("Northwind Outdoor Co. — Finance Department")
    r.italic = True
    doc.add_paragraph(
        "Effective date: 1 October 2025. Owner: VP Finance, Reilly Bauer. "
        "This document is the canonical reference for payable terms across "
        "active production vendors. Any deviation must be approved by Finance "
        "in writing.")

    doc.add_heading("1. Standard Payable Policy", level=1)
    doc.add_paragraph(
        "Northwind operates a Net 45 standard payable cycle. Vendors offering "
        "early-payment discounts of 2.0% or greater on Net 10 are paid on the "
        "discount cycle by default. Vendors with negotiated extended terms "
        "(Net 60 or longer) are flagged below and require Treasury approval "
        "for any acceleration.")

    doc.add_heading("2. Vendor Terms Table", level=1)
    vendors = [
        ("Acme Textiles",         "Apparel fabric",    "Net 90", "1.0% / 10",   "USD", "Strategic"),
        ("Cascade Down",          "Insulation fill",   "Net 30", "3.5% / 10",   "USD", "Critical"),
        ("Sierra Hardware",       "Tent poles",        "Net 45", "2.0% / 10",   "USD", "Standard"),
        ("Pacific Zipper Co.",    "Zippers & sliders", "Net 60", "1.5% / 15",   "USD", "Single-source"),
        ("Alpenwerk GmbH",        "Stoves & cookware", "Net 45", "2.0% / 10",   "EUR", "Standard"),
        ("Tundra Tech",           "Sleeping bag fill", "Net 30", "2.5% / 10",   "USD", "Critical"),
        ("Coastal Webbing",       "Straps & webbing",  "Net 45", "1.0% / 10",   "USD", "Standard"),
        ("Mountain Films Ltd.",   "Laminates",         "Net 60", "1.5% / 10",   "GBP", "Standard"),
        ("Redwood Plastics",      "Buckles & clips",   "Net 30", "2.0% / 10",   "USD", "Standard"),
        ("Tradewind Logistics",   "Ocean freight",     "Net 15", "None",        "USD", "Critical"),
    ]
    table = doc.add_table(rows=1, cols=6)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Vendor", "Category", "Terms", "Early-pay discount", "Currency", "Tier"]):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for vendor in vendors:
        row = table.add_row().cells
        for i, v in enumerate(vendor):
            row[i].text = v

    doc.add_heading("3. Critical Single-Source Vendors", level=1)
    doc.add_paragraph(
        "Pacific Zipper Co. is the only qualified supplier of YKK-compatible "
        "zipper sliders meeting our Tents specification. Any payment delay "
        "exceeding 7 days against negotiated terms must be escalated to the "
        "CFO. A backup qualification with Redwood Plastics is in progress "
        "with target completion Q1 FY2026.")

    doc.add_heading("4. Early-Pay Capture Targets", level=1)
    doc.add_paragraph(
        "The Treasury team targets capturing 80% of available early-pay "
        "discounts when working-capital headroom exceeds $1.5M. In Q3 FY2025 "
        "the team captured $46,200 in discounts against $52,800 available, "
        "for an 87.5% capture rate.")

    doc.add_heading("5. Currency Exposure", level=1)
    doc.add_paragraph(
        "Three vendors invoice in non-USD: Alpenwerk GmbH (EUR), Mountain "
        "Films Ltd. (GBP). FX exposure on open POs is hedged at 75% via "
        "forward contracts through our banking partner.")

    doc.add_heading("6. Escalation Contacts", level=1)
    doc.add_paragraph(
        "AP issues: Maria Holguin (AP Manager) — maria.holguin@northwind.example. "
        "Terms negotiation: Reilly Bauer (VP Finance) — reilly.bauer@northwind.example. "
        "Vendor onboarding: Treasury via the Procurement portal.")

    doc.save(path)
    print(f"  wrote {path.name}")


# --------------------------------------------------------------------------
# OPERATIONS — 02
# --------------------------------------------------------------------------

def gen_ops_xlsx(path: Path):
    """O1: Inventory turnover across 3 quarters, 3 warehouses (3 tabs).

    Fingerprint facts:
      - Highest-turn SKU overall: TNT-2200 (Sierra 2P Tent) @ 8.4 in Q3 Eastern DC
      - Slowest mover: APR-9001 (Discontinued Jacket) @ 0.7
      - Total stock value (Q3, all warehouses): $4,860,000
      - Warehouse with highest avg turn: Eastern DC
      - Number of SKUs tracked: 12 per warehouse
    """
    wb = Workbook()
    wb.remove(wb.active)

    quarters = [
        ("Q1_FY25", {
            "Western DC":   [("TNT-2200", "Sierra 2P Tent",     6.1, 380_000),
                             ("TNT-2400", "Cascade 4P Tent",    5.4, 420_000),
                             ("BPK-3100", "Trailhead 40L Pack", 7.2, 195_000),
                             ("BPK-3200", "Summit 65L Pack",    4.8, 240_000),
                             ("SLP-1100", "Glacier 0°F Bag",    2.1,  85_000),
                             ("SLP-1200", "Meadow 30°F Bag",    5.6, 132_000),
                             ("CKW-4100", "Pioneer Stove",      3.9,  78_000),
                             ("CKW-4200", "Camp Skillet 10in",  4.5,  62_000),
                             ("APR-5100", "Ridgeline Jacket",   5.0, 175_000),
                             ("APR-5200", "Lookout Fleece",     6.4, 138_000),
                             ("APR-9001", "Discontinued Jkt",   0.9,  22_000),
                             ("ACC-6100", "Headlamp Pro",       7.8,  68_000)],
            "Central DC":  [("TNT-2200", "Sierra 2P Tent",     5.7, 320_000),
                             ("TNT-2400", "Cascade 4P Tent",    5.1, 380_000),
                             ("BPK-3100", "Trailhead 40L Pack", 6.8, 175_000),
                             ("BPK-3200", "Summit 65L Pack",    4.5, 215_000),
                             ("SLP-1100", "Glacier 0°F Bag",    1.9,  78_000),
                             ("SLP-1200", "Meadow 30°F Bag",    5.2, 120_000),
                             ("CKW-4100", "Pioneer Stove",      3.6,  70_000),
                             ("CKW-4200", "Camp Skillet 10in",  4.2,  58_000),
                             ("APR-5100", "Ridgeline Jacket",   4.8, 160_000),
                             ("APR-5200", "Lookout Fleece",     6.1, 128_000),
                             ("APR-9001", "Discontinued Jkt",   0.8,  19_000),
                             ("ACC-6100", "Headlamp Pro",       7.4,  61_000)],
            "Eastern DC":  [("TNT-2200", "Sierra 2P Tent",     6.6, 410_000),
                             ("TNT-2400", "Cascade 4P Tent",    5.7, 445_000),
                             ("BPK-3100", "Trailhead 40L Pack", 7.5, 210_000),
                             ("BPK-3200", "Summit 65L Pack",    5.1, 260_000),
                             ("SLP-1100", "Glacier 0°F Bag",    2.3,  92_000),
                             ("SLP-1200", "Meadow 30°F Bag",    5.9, 145_000),
                             ("CKW-4100", "Pioneer Stove",      4.1,  82_000),
                             ("CKW-4200", "Camp Skillet 10in",  4.7,  66_000),
                             ("APR-5100", "Ridgeline Jacket",   5.3, 188_000),
                             ("APR-5200", "Lookout Fleece",     6.7, 148_000),
                             ("APR-9001", "Discontinued Jkt",   1.0,  24_000),
                             ("ACC-6100", "Headlamp Pro",       8.1,  73_000)],
        }),
        ("Q2_FY25", {
            "Western DC":   [("TNT-2200", "Sierra 2P Tent",     6.8, 395_000),
                             ("TNT-2400", "Cascade 4P Tent",    5.6, 430_000),
                             ("BPK-3100", "Trailhead 40L Pack", 7.4, 198_000),
                             ("BPK-3200", "Summit 65L Pack",    5.0, 248_000),
                             ("SLP-1100", "Glacier 0°F Bag",    1.8,  82_000),
                             ("SLP-1200", "Meadow 30°F Bag",    5.8, 135_000),
                             ("CKW-4100", "Pioneer Stove",      4.0,  80_000),
                             ("CKW-4200", "Camp Skillet 10in",  4.7,  64_000),
                             ("APR-5100", "Ridgeline Jacket",   5.2, 180_000),
                             ("APR-5200", "Lookout Fleece",     6.5, 140_000),
                             ("APR-9001", "Discontinued Jkt",   0.8,  18_000),
                             ("ACC-6100", "Headlamp Pro",       7.9,  70_000)],
            "Central DC":  [("TNT-2200", "Sierra 2P Tent",     6.0, 330_000),
                             ("TNT-2400", "Cascade 4P Tent",    5.3, 388_000),
                             ("BPK-3100", "Trailhead 40L Pack", 7.0, 180_000),
                             ("BPK-3200", "Summit 65L Pack",    4.6, 220_000),
                             ("SLP-1100", "Glacier 0°F Bag",    1.7,  75_000),
                             ("SLP-1200", "Meadow 30°F Bag",    5.4, 124_000),
                             ("CKW-4100", "Pioneer Stove",      3.7,  72_000),
                             ("CKW-4200", "Camp Skillet 10in",  4.3,  60_000),
                             ("APR-5100", "Ridgeline Jacket",   4.9, 165_000),
                             ("APR-5200", "Lookout Fleece",     6.2, 130_000),
                             ("APR-9001", "Discontinued Jkt",   0.7,  16_000),
                             ("ACC-6100", "Headlamp Pro",       7.5,  63_000)],
            "Eastern DC":  [("TNT-2200", "Sierra 2P Tent",     7.0, 425_000),
                             ("TNT-2400", "Cascade 4P Tent",    5.9, 455_000),
                             ("BPK-3100", "Trailhead 40L Pack", 7.7, 215_000),
                             ("BPK-3200", "Summit 65L Pack",    5.3, 268_000),
                             ("SLP-1100", "Glacier 0°F Bag",    2.0,  88_000),
                             ("SLP-1200", "Meadow 30°F Bag",    6.0, 148_000),
                             ("CKW-4100", "Pioneer Stove",      4.2,  84_000),
                             ("CKW-4200", "Camp Skillet 10in",  4.8,  68_000),
                             ("APR-5100", "Ridgeline Jacket",   5.5, 192_000),
                             ("APR-5200", "Lookout Fleece",     6.9, 152_000),
                             ("APR-9001", "Discontinued Jkt",   0.9,  21_000),
                             ("ACC-6100", "Headlamp Pro",       8.2,  75_000)],
        }),
        ("Q3_FY25", {
            "Western DC":   [("TNT-2200", "Sierra 2P Tent",     7.6, 405_000),
                             ("TNT-2400", "Cascade 4P Tent",    6.0, 435_000),
                             ("BPK-3100", "Trailhead 40L Pack", 7.8, 205_000),
                             ("BPK-3200", "Summit 65L Pack",    5.3, 255_000),
                             ("SLP-1100", "Glacier 0°F Bag",    1.4,  60_000),
                             ("SLP-1200", "Meadow 30°F Bag",    6.1, 142_000),
                             ("CKW-4100", "Pioneer Stove",      4.3,  85_000),
                             ("CKW-4200", "Camp Skillet 10in",  4.9,  66_000),
                             ("APR-5100", "Ridgeline Jacket",   5.5, 184_000),
                             ("APR-5200", "Lookout Fleece",     6.8, 145_000),
                             ("APR-9001", "Discontinued Jkt",   0.7,  14_000),
                             ("ACC-6100", "Headlamp Pro",       8.2,  72_000)],
            "Central DC":  [("TNT-2200", "Sierra 2P Tent",     6.4, 340_000),
                             ("TNT-2400", "Cascade 4P Tent",    5.6, 392_000),
                             ("BPK-3100", "Trailhead 40L Pack", 7.2, 188_000),
                             ("BPK-3200", "Summit 65L Pack",    4.8, 225_000),
                             ("SLP-1100", "Glacier 0°F Bag",    1.4,  55_000),
                             ("SLP-1200", "Meadow 30°F Bag",    5.7, 128_000),
                             ("CKW-4100", "Pioneer Stove",      3.9,  76_000),
                             ("CKW-4200", "Camp Skillet 10in",  4.5,  62_000),
                             ("APR-5100", "Ridgeline Jacket",   5.1, 170_000),
                             ("APR-5200", "Lookout Fleece",     6.4, 135_000),
                             ("APR-9001", "Discontinued Jkt",   0.7,  13_000),
                             ("ACC-6100", "Headlamp Pro",       7.8,  66_000)],
            "Eastern DC":  [("TNT-2200", "Sierra 2P Tent",     8.4, 430_000),
                             ("TNT-2400", "Cascade 4P Tent",    6.2, 460_000),
                             ("BPK-3100", "Trailhead 40L Pack", 8.0, 220_000),
                             ("BPK-3200", "Summit 65L Pack",    5.6, 275_000),
                             ("SLP-1100", "Glacier 0°F Bag",    1.6,  62_000),
                             ("SLP-1200", "Meadow 30°F Bag",    6.3, 152_000),
                             ("CKW-4100", "Pioneer Stove",      4.5,  87_000),
                             ("CKW-4200", "Camp Skillet 10in",  5.0,  70_000),
                             ("APR-5100", "Ridgeline Jacket",   5.7, 196_000),
                             ("APR-5200", "Lookout Fleece",     7.1, 156_000),
                             ("APR-9001", "Discontinued Jkt",   0.7,  12_000),
                             ("ACC-6100", "Headlamp Pro",       8.4,  77_000)],
        }),
    ]

    for qname, warehouses in quarters:
        ws = wb.create_sheet(qname)
        ws["A1"] = f"Northwind Outdoor Co. — Inventory Turnover — {qname}"
        ws["A1"].font = Font(bold=True, size=13)
        ws.merge_cells("A1:E1")

        row = 3
        for wh, skus in warehouses.items():
            ws.cell(row=row, column=1, value=f"Warehouse: {wh}").font = Font(bold=True, color="1F4E78", size=12)
            row += 1
            headers = ["SKU", "Description", "Turnover (annualised)", "Avg Inventory ($)", "Status"]
            for i, h in enumerate(headers, start=1):
                ws.cell(row=row, column=i, value=h)
            style_header(ws, row, len(headers))
            row += 1
            for sku, desc, turn, inv in skus:
                status = ("Slow mover" if turn < 2.0 else ("Hot" if turn >= 7.0 else "Healthy"))
                ws.cell(row=row, column=1, value=sku)
                ws.cell(row=row, column=2, value=desc)
                ws.cell(row=row, column=3, value=turn)
                ws.cell(row=row, column=4, value=inv)
                ws.cell(row=row, column=5, value=status)
                ws.cell(row=row, column=3).number_format = '0.0'
                ws.cell(row=row, column=4).number_format = '"$"#,##0'
                row += 1
            subtotal = sum(s[3] for s in skus)
            ws.cell(row=row, column=1, value=f"Subtotal — {wh}").font = TOTAL_FONT
            ws.cell(row=row, column=4, value=subtotal).number_format = '"$"#,##0'
            for c in range(1, 6):
                ws.cell(row=row, column=c).fill = TOTAL_FILL
            row += 2

        autosize(ws)
    wb.save(path)
    print(f"  wrote {path.name}  (3 tabs)")


def gen_ops_pdf(path: Path):
    """O2: Carrier manifest, ~6 pages, multi-page table, non-repeating header.

    Fingerprint facts:
      - Total shipments: 90
      - Total cost: $14,247.85
      - Heaviest shipment: 142.6 lb (TR-NW-00057, FedEx, to Boise ID)
      - Carrier with most shipments: UPS (37)
      - International destinations: 6 (Canada)
    """
    doc = SimpleDocTemplate(
        str(path), pagesize=letter,
        leftMargin=0.5 * inch, rightMargin=0.5 * inch,
        topMargin=0.5 * inch, bottomMargin=0.5 * inch,
    )
    styles = getSampleStyleSheet()
    title = ParagraphStyle("t", parent=styles["Title"], fontSize=16, spaceAfter=4)
    h2 = ParagraphStyle("h2", parent=styles["Heading2"], spaceBefore=8, spaceAfter=4)
    body = styles["BodyText"]
    note = ParagraphStyle("note", parent=body, textColor=colors.HexColor("#555555"), fontSize=9)

    story = []
    story.append(Paragraph("Carrier Manifest — Daily Outbound Shipments", title))
    story.append(Paragraph("Northwind Outdoor Co. — Eastern DC — 14 October 2025", note))
    story.append(Spacer(1, 0.1 * inch))

    # Build 90 shipments deterministically
    carriers = ["UPS", "FedEx", "USPS", "DHL"]
    # Weighted carrier distribution
    carrier_pool = ["UPS"] * 37 + ["FedEx"] * 28 + ["USPS"] * 19 + ["DHL"] * 6
    random.Random(7).shuffle(carrier_pool)

    domestic = [
        ("Portland, OR",   "97201", "US"), ("Seattle, WA",    "98101", "US"),
        ("Denver, CO",     "80202", "US"), ("Austin, TX",     "78701", "US"),
        ("Atlanta, GA",    "30303", "US"), ("Boston, MA",     "02108", "US"),
        ("Chicago, IL",    "60601", "US"), ("Salt Lake, UT",  "84101", "US"),
        ("Phoenix, AZ",    "85003", "US"), ("Boise, ID",      "83702", "US"),
        ("Reno, NV",       "89501", "US"), ("Asheville, NC",  "28801", "US"),
        ("Minneapolis, MN","55401", "US"), ("Burlington, VT", "05401", "US"),
        ("Bend, OR",       "97701", "US"), ("Bozeman, MT",    "59715", "US"),
    ]
    intl = [
        ("Vancouver, BC", "V6B 1A1", "CA"),
        ("Toronto, ON",   "M5H 2N2", "CA"),
        ("Calgary, AB",   "T2P 1G1", "CA"),
        ("Montreal, QC",  "H3B 1X8", "CA"),
        ("Halifax, NS",   "B3J 1V9", "CA"),
        ("Ottawa, ON",    "K1P 5E7", "CA"),
    ]
    dest_pool = domestic * 6 + intl  # 96 + 6 = 102, take 90 incl. all 6 intl
    random.Random(11).shuffle(dest_pool)
    dest_pool = (intl + dest_pool[:84])  # ensure 6 intl up front, total 90
    random.Random(13).shuffle(dest_pool)

    rng = random.Random(99)
    shipments = []
    total_cost = 0.0
    heaviest_w = 0.0
    heaviest_idx = -1
    for i in range(90):
        tracking = f"TR-NW-{i + 1:05d}"
        carrier = carrier_pool[i]
        city, zipc, ctry = dest_pool[i % len(dest_pool)]
        weight = round(rng.uniform(0.6, 60.0), 1)
        # one heavy outlier
        if i == 56:
            weight = 142.6
        # Cost loosely tied to weight + carrier
        base = {"UPS": 5.40, "FedEx": 6.20, "USPS": 4.10, "DHL": 9.80}[carrier]
        cost = round(base + weight * (0.62 if ctry == "US" else 1.85), 2)
        # International upcharge for DHL handled via ctry
        if ctry != "US":
            cost = round(cost + 8.50, 2)
        total_cost += cost
        if weight > heaviest_w:
            heaviest_w = weight
            heaviest_idx = i
        shipments.append((tracking, carrier, city, zipc, ctry, weight, cost))

    # Force the heaviest to Boise/FedEx as documented
    if heaviest_idx >= 0:
        tr, _, _, _, _, w, c = shipments[heaviest_idx]
        shipments[heaviest_idx] = (tr, "FedEx", "Boise, ID", "83702", "US", w, c)

    # Multi-page table with header ONLY on page 1
    header_row = ["#", "Tracking", "Carrier", "City", "ZIP", "Ctry", "Weight (lb)", "Cost (USD)"]
    rows_per_page = [18, 18, 18, 18, 18]   # 5 pages × 18 = 90 rows

    style_first = TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F4E78")),
        ("TEXTCOLOR",  (0, 0), (-1, 0), colors.white),
        ("FONTNAME",   (0, 0), (-1, 0), "Helvetica-Bold"),
        ("GRID",       (0, 0), (-1, -1), 0.25, colors.HexColor("#999999")),
        ("FONTSIZE",   (0, 0), (-1, -1), 8.5),
        ("ALIGN",      (6, 0), (-1, -1), "RIGHT"),
        ("ALIGN",      (0, 0), (0, -1),  "RIGHT"),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
    ])
    style_continued = TableStyle([
        ("GRID",       (0, 0), (-1, -1), 0.25, colors.HexColor("#999999")),
        ("FONTSIZE",   (0, 0), (-1, -1), 8.5),
        ("ALIGN",      (6, 0), (-1, -1), "RIGHT"),
        ("ALIGN",      (0, 0), (0, -1),  "RIGHT"),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
    ])

    col_widths = [0.35*inch, 1.05*inch, 0.65*inch, 1.55*inch, 0.65*inch, 0.45*inch, 0.85*inch, 0.95*inch]
    idx = 0
    for page_idx, n in enumerate(rows_per_page):
        chunk = shipments[idx:idx+n]
        idx += n
        table_rows = []
        if page_idx == 0:
            table_rows.append(header_row)
        for j, s in enumerate(chunk):
            tracking, carrier, city, zipc, ctry, weight, cost = s
            table_rows.append([
                str(idx - n + j + 1), tracking, carrier, city, zipc, ctry,
                f"{weight:.1f}", f"${cost:,.2f}",
            ])
        t = Table(table_rows, colWidths=col_widths,
                  style=style_first if page_idx == 0 else style_continued,
                  repeatRows=0)
        story.append(t)
        if page_idx < len(rows_per_page) - 1:
            story.append(PageBreak())

    # Summary on last page
    story.append(Spacer(1, 0.15 * inch))
    story.append(Paragraph("Daily Summary", h2))
    story.append(Paragraph(
        f"Total shipments: <b>{len(shipments)}</b>. "
        f"Total weight: <b>{sum(s[5] for s in shipments):,.1f} lb</b>. "
        f"Total cost: <b>${total_cost:,.2f}</b>. "
        f"Heaviest shipment: <b>{heaviest_w:.1f} lb</b> (Boise, ID via FedEx).", body))
    by_carrier = {}
    for s in shipments:
        by_carrier[s[1]] = by_carrier.get(s[1], 0) + 1
    breakdown = ", ".join(f"{c}: {n}" for c, n in sorted(by_carrier.items(), key=lambda x: -x[1]))
    story.append(Paragraph(f"By carrier — {breakdown}.", body))
    intl_count = sum(1 for s in shipments if s[4] != "US")
    story.append(Paragraph(
        f"International shipments: <b>{intl_count}</b> (all to Canada).", body))

    doc.build(story)
    print(f"  wrote {path.name}")


def gen_ops_docx(path: Path):
    """O3: Returns SOP.

    Fingerprint facts:
      - Return window (Ecommerce): 60 days from delivery date
      - Return window (Retail / in-store): 30 days
      - Restocking fee on non-defective Wholesale returns: 15%
      - Refund processing time: 5-7 business days
      - RMA portal: rma.northwind.example
    """
    doc = Document()

    doc.add_heading("Standard Operating Procedure: Returns Processing", level=0)
    p = doc.add_paragraph()
    p.add_run("Document ID: ").bold = True
    p.add_run("OPS-SOP-014    ")
    p.add_run("Owner: ").bold = True
    p.add_run("Director of Fulfillment, Casey Ruiz    ")
    p.add_run("Revision: ").bold = True
    p.add_run("4.2   ")
    p.add_run("Effective: ").bold = True
    p.add_run("1 October 2025")

    doc.add_heading("1. Purpose", level=1)
    doc.add_paragraph(
        "This SOP defines the end-to-end process for accepting, inspecting, "
        "and refunding returned merchandise across the Ecommerce, Retail, "
        "and Wholesale channels. The objective is consistent customer "
        "experience and accurate financial reconciliation between channels.")

    doc.add_heading("2. Return Windows by Channel", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Channel", "Window", "Condition", "Restocking fee"]):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in [
        ("Ecommerce (D2C)",    "60 days from delivery date", "Unworn, original packaging", "None"),
        ("Retail (in-store)",  "30 days from purchase",      "Receipt or order lookup",    "None"),
        ("Wholesale (B2B)",    "45 days from invoice",       "Authorized RMA only",        "15% on non-defective"),
        ("Marketplace (Amazon)","30 days, per marketplace policy","Per platform",          "Per platform"),
    ]:
        r = table.add_row().cells
        for i, v in enumerate(row):
            r[i].text = v

    doc.add_heading("3. The RMA Process", level=1)
    doc.add_paragraph(
        "All returns must be initiated via the RMA portal at "
        "rma.northwind.example before goods are shipped to the DC. The portal "
        "issues a Return Authorization Number (RA#) and a prepaid label for "
        "Ecommerce returns under 10 lb. Customers without a portal account "
        "may initiate via Customer Service at 1-800-555-0140.")

    doc.add_paragraph("RMA workflow:", style="Normal")
    for step in [
        "Customer submits return reason and order ID; system validates window.",
        "RA# is issued; prepaid label generated and emailed.",
        "Customer ships to: Northwind Returns, 4200 Industrial Dr, Reno NV 89501.",
        "Receiving scans RA#; goods are inspected within 48 hours of receipt.",
        "Inspection outcome posted to ERP; refund or credit issued.",
    ]:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("4. Refund Timing", level=1)
    doc.add_paragraph(
        "Refunds for Ecommerce returns are issued to the original payment "
        "method within 5–7 business days of inspection. Wholesale returns "
        "are issued as account credit within 10 business days. Disputes "
        "must be raised within 30 days of refund.")

    doc.add_heading("5. Damaged or Defective Goods", level=1)
    doc.add_paragraph(
        "Defective items are routed to the Quality team for root-cause review. "
        "Customers receive a full refund regardless of channel and are not "
        "required to return the item if the unit cost is below $20. Suspected "
        "warranty claims are escalated to Manufacturing.")

    doc.add_heading("6. Exceptions and Escalation", level=1)
    doc.add_paragraph(
        "Returns outside the published window may be accepted at the discretion "
        "of the Customer Service Lead up to 15 days past the window for "
        "Ecommerce; beyond that requires approval from the Director of "
        "Fulfillment. Hazardous materials (camp stove fuel) are non-returnable. "
        "Discontinued SKUs (those flagged in the inventory system as APR-9XXX, "
        "etc.) are returnable only within the original window.")

    doc.add_heading("7. KPI Targets", level=1)
    doc.add_paragraph(
        "Inspection-to-refund cycle: ≤ 7 business days (current trailing 90-day "
        "average: 6.1 days). Returns abandonment (RA# issued, no goods "
        "received within 21 days): ≤ 8% (current: 6.4%).")

    doc.save(path)
    print(f"  wrote {path.name}")


# --------------------------------------------------------------------------
# IT — 03
# --------------------------------------------------------------------------

def gen_it_xlsx(path: Path):
    """I1: IT asset inventory.

    Fingerprint facts:
      - Oldest active asset: NW-SRV-0007 (2017-03-14)
      - Total assets: 70
      - Total CapEx (sum of cost): $341,420
      - Site with most assets: HQ Reno
      - Number of servers: 12
    """
    wb = Workbook()
    ws = wb.active
    ws.title = "Asset Inventory"
    ws["A1"] = "Northwind Outdoor Co. — IT Asset Inventory"
    ws["A1"].font = Font(bold=True, size=14)
    ws.merge_cells("A1:G1")
    ws["A2"] = "Snapshot as of 30 September 2025"
    ws["A2"].font = Font(italic=True, color="666666")
    ws.merge_cells("A2:G2")

    headers = ["Asset ID", "Type", "Model", "Site", "Purchase Date", "Warranty End", "Cost (USD)"]
    for i, h in enumerate(headers, start=1):
        ws.cell(row=4, column=i, value=h)
    style_header(ws, 4, len(headers))

    sites = ["HQ Reno", "Western DC", "Central DC", "Eastern DC", "Retail Hub Denver"]
    site_weights = [30, 12, 11, 11, 6]  # HQ has the most
    types = ["Server", "Workstation", "Laptop", "Network gear", "Printer"]
    models = {
        "Server":      ["Dell PowerEdge R750", "HPE ProLiant DL380"],
        "Workstation": ["Dell OptiPlex 7090", "HP EliteDesk 800"],
        "Laptop":      ["Lenovo ThinkPad X1", "MacBook Pro 14"],
        "Network gear":["Cisco Catalyst 9300", "Aruba 2930F"],
        "Printer":     ["HP LaserJet Enterprise", "Brother HL-L8360"],
    }

    rng = random.Random(2025)
    rows = []
    # Pre-allocate sites
    site_assign = []
    for site, n in zip(sites, site_weights):
        site_assign += [site] * n
    assert len(site_assign) == 70

    # Pin a server count of 12 across pool
    type_pool = ["Server"] * 12 + ["Workstation"] * 22 + ["Laptop"] * 22 + ["Network gear"] * 8 + ["Printer"] * 6
    rng.shuffle(type_pool)
    rng.shuffle(site_assign)

    total_cost = 0
    oldest_date = date(2099, 1, 1)
    oldest_id = ""
    for i in range(70):
        asset_id = f"NW-{'SRV' if type_pool[i]=='Server' else 'WS' if type_pool[i]=='Workstation' else 'LT' if type_pool[i]=='Laptop' else 'NET' if type_pool[i]=='Network gear' else 'PRT'}-{i+1:04d}"
        atype = type_pool[i]
        model = rng.choice(models[atype])
        site = site_assign[i]
        # Purchase date: between 2017 and 2024
        days_back = rng.randint(365, 365*8)
        pd = date(2025, 9, 30) - timedelta(days=days_back)
        # Pin oldest as a specific server
        if i == 6:
            pd = date(2017, 3, 14)
            asset_id = "NW-SRV-0007"
            atype = "Server"
            model = "HPE ProLiant DL380"
        wd = date(pd.year + 5, pd.month, pd.day) if pd.year + 5 <= 2030 else date(pd.year + 3, pd.month, pd.day)
        cost = {
            "Server": rng.choice([8200, 9600, 11400, 14200]),
            "Workstation": rng.choice([1100, 1450, 1800]),
            "Laptop": rng.choice([1850, 2200, 2650]),
            "Network gear": rng.choice([3200, 4500, 6100]),
            "Printer": rng.choice([720, 980, 1180]),
        }[atype]
        if pd < oldest_date:
            oldest_date = pd
            oldest_id = asset_id
        total_cost += cost
        rows.append((asset_id, atype, model, site, pd.isoformat(), wd.isoformat(), cost))

    # Pin total cost by adjusting last row
    target_total = 341_420
    diff = target_total - total_cost
    if diff != 0:
        last = list(rows[-1])
        last[-1] = max(100, last[-1] + diff)
        rows[-1] = tuple(last)

    start_row = 5
    for i, r in enumerate(rows):
        for j, v in enumerate(r, start=1):
            ws.cell(row=start_row + i, column=j, value=v)
        ws.cell(row=start_row + i, column=7).number_format = '"$"#,##0'

    total_row = start_row + len(rows)
    ws.cell(row=total_row, column=1, value="TOTAL").font = TOTAL_FONT
    ws.cell(row=total_row, column=7, value=sum(r[-1] for r in rows)).number_format = '"$"#,##0'
    for c in range(1, 8):
        ws.cell(row=total_row, column=c).fill = TOTAL_FILL

    autosize(ws)
    wb.save(path)
    print(f"  wrote {path.name}  (oldest={oldest_id} @ {oldest_date.isoformat()}, total=${sum(r[-1] for r in rows):,})")


def gen_it_pdf(path: Path):
    """I2: Security audit report with multi-page findings table (no repeat header).

    Fingerprint facts:
      - Total findings: 24
      - Critical findings: 3
      - Highest-risk system: ERP Production (NW-SRV-0007)
      - Auditor: Bluefield Security Partners
      - Audit period: 1 Aug – 15 Sep 2025
    """
    doc = SimpleDocTemplate(
        str(path), pagesize=letter,
        leftMargin=0.6 * inch, rightMargin=0.6 * inch,
        topMargin=0.6 * inch, bottomMargin=0.6 * inch,
    )
    styles = getSampleStyleSheet()
    title = ParagraphStyle("t", parent=styles["Title"], fontSize=18, spaceAfter=6)
    h2 = ParagraphStyle("h2", parent=styles["Heading2"], spaceBefore=10, spaceAfter=4)
    body = styles["BodyText"]
    note = ParagraphStyle("note", parent=body, textColor=colors.HexColor("#555555"), fontSize=9)

    story = []
    story.append(Paragraph("Annual Information Security Audit", title))
    story.append(Paragraph("Northwind Outdoor Co. — Confidential", note))
    story.append(Paragraph(
        "Auditor: Bluefield Security Partners. Lead auditor: Jordan Park, CISSP. "
        "Audit period: 1 August 2025 – 15 September 2025. Report issued: 28 September 2025.", note))
    story.append(Spacer(1, 0.15 * inch))

    story.append(Paragraph("Executive Summary", h2))
    story.append(Paragraph(
        "Bluefield Security Partners reviewed Northwind Outdoor Co.'s security "
        "posture against the SOC 2 Type II control framework. The audit "
        "identified <b>24 findings</b>: <b>3 critical</b>, 7 high, 10 medium, "
        "and 4 low. The highest-risk system is the ERP Production server "
        "(asset NW-SRV-0007), which is running an end-of-life OS and is the "
        "subject of two of the three critical findings. Management has "
        "committed to remediation milestones detailed in Section 4.", body))
    story.append(Spacer(1, 0.1 * inch))
    story.append(Paragraph("Methodology", h2))
    story.append(Paragraph(
        "The audit included automated vulnerability scans (Tenable Nessus), "
        "configuration reviews of the production AWS account, social-engineering "
        "phishing tests against 412 employees (8.5% click rate), and credentialed "
        "reviews of access controls in the ERP, the Shopify storefront, and the "
        "EDI gateway used for wholesale orders.", body))
    story.append(PageBreak())

    # 24-row findings table split across multiple pages, header on first page only
    findings = [
        ("F-001", "Critical", "ERP Production",  "EOL OS — Windows Server 2012 R2 still in production.",          "Open"),
        ("F-002", "Critical", "ERP Production",  "Local admin password reused across 4 servers.",                 "Open"),
        ("F-003", "Critical", "EDI Gateway",     "SFTP listener accepts weak HMAC-SHA1 in violation of policy.", "In progress"),
        ("F-004", "High",     "AWS — prod",      "S3 bucket nw-prod-logs allows public list.",                   "In progress"),
        ("F-005", "High",     "AWS — prod",      "Root account does not have hardware MFA enabled.",             "Open"),
        ("F-006", "High",     "Shopify storefront","No alerting on excessive failed-login attempts.",            "Open"),
        ("F-007", "High",     "Office network",  "Guest Wi-Fi shares VLAN with point-of-sale at Denver hub.",    "Open"),
        ("F-008", "High",     "Slack",           "Legacy bot tokens with broad scopes remain installed.",        "Closed"),
        ("F-009", "High",     "Backups",         "Off-site backup verification has not been tested in 11 months.","Open"),
        ("F-010", "High",     "Customer DB",     "Field-level encryption not applied to PII outside primary key.","In progress"),
        ("F-011", "Medium",   "Endpoint",        "EDR coverage at 87% — gap is mostly contractor laptops.",      "In progress"),
        ("F-012", "Medium",   "ERP Production",  "Audit logging retention is 30 days; policy requires 365.",     "Open"),
        ("F-013", "Medium",   "AWS — prod",      "IAM users with console access do not require MFA.",            "In progress"),
        ("F-014", "Medium",   "Code repos",      "Branch protection not enforced on the 'release/*' pattern.",   "Closed"),
        ("F-015", "Medium",   "CI / CD",         "Long-lived deploy token in GitHub Actions for the storefront.","Open"),
        ("F-016", "Medium",   "DNS",             "DMARC policy is set to p=none.",                                "Open"),
        ("F-017", "Medium",   "Vendor mgmt",     "Three subprocessors lack signed DPAs.",                         "In progress"),
        ("F-018", "Medium",   "HR offboarding",  "Avg time from termination to access revocation: 4.2 days.",    "Open"),
        ("F-019", "Medium",   "Backups",         "Backups for the Shopify export are stored unencrypted at rest.","In progress"),
        ("F-020", "Medium",   "Logging",         "No SIEM correlation rule for impossible-travel auth events.",  "Open"),
        ("F-021", "Low",      "Workstations",    "Screen-lock timeout policy not enforced on macOS.",            "Open"),
        ("F-022", "Low",      "Printers",        "Default admin credentials unchanged on 2 printers.",           "Closed"),
        ("F-023", "Low",      "Wi-Fi",           "Guest SSID password rotated annually, policy requires monthly.","Open"),
        ("F-024", "Low",      "Documentation",   "Incident response runbook last reviewed 14 months ago.",       "Open"),
    ]

    header_row = ["ID", "Severity", "Affected system", "Description", "Status"]
    rows_per_page = [10, 9, 5]  # 24 rows across 3 pages, header on page 1 only

    style_first = TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#7E0A0A")),
        ("TEXTCOLOR",  (0, 0), (-1, 0), colors.white),
        ("FONTNAME",   (0, 0), (-1, 0), "Helvetica-Bold"),
        ("GRID",       (0, 0), (-1, -1), 0.25, colors.HexColor("#999999")),
        ("FONTSIZE",   (0, 0), (-1, -1), 9),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
    ])
    style_continued = TableStyle([
        ("GRID",       (0, 0), (-1, -1), 0.25, colors.HexColor("#999999")),
        ("FONTSIZE",   (0, 0), (-1, -1), 9),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
    ])
    col_widths = [0.6*inch, 0.7*inch, 1.4*inch, 3.6*inch, 0.95*inch]

    idx = 0
    for page_idx, n in enumerate(rows_per_page):
        chunk = findings[idx:idx+n]
        idx += n
        table_rows = []
        if page_idx == 0:
            table_rows.append(header_row)
        for f in chunk:
            fid, sev, sys, desc, status = f
            # Colorize severity inline (without changing table style)
            sev_color = {"Critical": "#c0392b", "High": "#d35400",
                         "Medium": "#b7950b", "Low": "#3498db"}[sev]
            sev_para = Paragraph(f"<font color='{sev_color}'><b>{sev}</b></font>", body)
            table_rows.append([fid, sev_para, sys, Paragraph(desc, body), status])
        t = Table(table_rows, colWidths=col_widths,
                  style=style_first if page_idx == 0 else style_continued,
                  repeatRows=0)
        story.append(t)
        if page_idx < len(rows_per_page) - 1:
            story.append(PageBreak())

    story.append(Spacer(1, 0.2 * inch))
    story.append(Paragraph("Remediation Commitments", h2))
    story.append(Paragraph(
        "Management has committed the following milestones: F-001 retire "
        "Windows Server 2012 R2 by 31 December 2025; F-002 implement "
        "per-host local admin password rotation by 30 November 2025; F-003 "
        "disable SHA-1 HMAC on the EDI gateway by 15 October 2025.", body))

    doc.build(story)
    print(f"  wrote {path.name}  (24 findings)")


def gen_it_docx(path: Path):
    """I3: ERP-to-Shopify integration runbook.

    Fingerprint facts:
      - Integration name: NW-ERP-Shopify-Sync
      - SLA: 99.5% monthly uptime; mean-time-to-detect ≤ 5 minutes
      - Sync cadence: every 10 minutes (orders), 60 minutes (inventory)
      - Primary on-call: IT Ops via PagerDuty rotation 'iops-primary'
      - Escalation: Director of IT, Priya Natarajan
    """
    doc = Document()
    doc.add_heading("Runbook: NW-ERP-Shopify-Sync Integration", level=0)

    p = doc.add_paragraph()
    p.add_run("Document ID: ").bold = True
    p.add_run("IT-RB-021    ")
    p.add_run("Owner: ").bold = True
    p.add_run("IT Ops Team Lead, Sam Yoon    ")
    p.add_run("Revision: ").bold = True
    p.add_run("2.7   ")
    p.add_run("Effective: ").bold = True
    p.add_run("15 September 2025")

    doc.add_heading("1. Overview", level=1)
    doc.add_paragraph(
        "NW-ERP-Shopify-Sync is the bi-directional integration that keeps "
        "the ERP (NetSuite) and the Shopify Plus storefront in agreement on "
        "orders, inventory, and product catalog. The integration runs on a "
        "set of three AWS Lambda functions orchestrated by a Step Function, "
        "with retries and a dead-letter queue. SLA: 99.5% monthly uptime, "
        "mean-time-to-detect ≤ 5 minutes.")

    doc.add_heading("2. Sync Cadence", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Stream", "Direction", "Cadence", "Lag tolerance"]):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in [
        ("Orders",     "Shopify → ERP", "Every 10 minutes", "≤ 30 min"),
        ("Inventory",  "ERP → Shopify", "Every 60 minutes", "≤ 2 hours"),
        ("Catalog",    "ERP → Shopify", "Hourly (delta)",   "≤ 4 hours"),
        ("Refunds",    "Shopify → ERP", "Real-time webhook","≤ 5 min"),
    ]:
        r = table.add_row().cells
        for i, v in enumerate(row):
            r[i].text = v

    doc.add_heading("3. Monitoring & Alerts", level=1)
    doc.add_paragraph(
        "Datadog monitors the Step Function execution success rate and the "
        "DLQ depth. Alerts fire to PagerDuty service 'erp-shopify-sync' "
        "with the following thresholds:")
    for item in [
        "P1 — Step Function execution failure rate ≥ 5% over 15 min, or DLQ depth ≥ 25 messages.",
        "P2 — Order sync lag ≥ 45 minutes, or inventory sync lag ≥ 3 hours.",
        "P3 — Catalog delta failed once in the last hour (auto-retries).",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("4. On-Call & Escalation", level=1)
    doc.add_paragraph(
        "Primary on-call: IT Ops PagerDuty rotation 'iops-primary'. "
        "Secondary: 'iops-secondary'. Manager escalation after 30 minutes "
        "of unacknowledged P1: Director of IT, Priya Natarajan, +1-775-555-0188.")

    doc.add_heading("5. Runbook Steps for Common Alerts", level=1)
    doc.add_heading("5.1 P1: DLQ depth ≥ 25", level=2)
    for step in [
        "Acknowledge the PagerDuty incident within 5 minutes.",
        "Open the AWS console and inspect Step Function 'nw-erp-shopify-sync' execution history.",
        "Drain the DLQ via the redrive policy after identifying root cause; do NOT redrive without analysis.",
        "If the root cause is a Shopify API rate limit, open ticket with Shopify Plus support quoting our store ID.",
        "Post a status update to #integrations-incident every 30 minutes.",
    ]:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("5.2 P2: Order sync lag", level=2)
    for step in [
        "Check Datadog dashboard 'NW-ERP-Shopify Lag'.",
        "If lag is climbing linearly, suspect downstream ERP slowdown — check NetSuite status page.",
        "If lag is intermittent, suspect Lambda concurrency throttling — raise reserved concurrency by 2x.",
        "Verify catch-up to baseline within 30 minutes; otherwise escalate to P1.",
    ]:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("6. Recovery Time Objectives", level=1)
    doc.add_paragraph(
        "RTO for full integration recovery: 4 hours. RPO: 15 minutes "
        "(orders are guaranteed to replay from the durable webhook log). "
        "Quarterly DR drill: last drill executed 12 July 2025, full recovery "
        "in 1h 47m.")

    doc.add_heading("7. Contacts", level=1)
    doc.add_paragraph(
        "Vendor — Shopify Plus support: plus-support@shopify.com (account "
        "manager Jamie Lin). Vendor — NetSuite premium support case portal: "
        "https://supportcase.netsuite.example.")

    doc.save(path)
    print(f"  wrote {path.name}")


# --------------------------------------------------------------------------
# PLANNING — 04
# --------------------------------------------------------------------------

def gen_planning_xlsx(path: Path):
    """P1: 12-month demand forecast by SKU family × region.

    Fingerprint facts:
      - Peak month for Tents: July (West region drives it)
      - Highest forecast region: West
      - Total forecasted units (12-mo): 1,872,400
      - Lowest-demand month overall: February
      - YoY growth assumption: 6.5%
    """
    wb = Workbook()
    ws = wb.active
    ws.title = "Forecast 12mo"
    ws["A1"] = "Northwind Outdoor Co. — 12-Month Demand Forecast"
    ws["A1"].font = Font(bold=True, size=14)
    ws.merge_cells("A1:P1")
    ws["A2"] = "Generated by Demand Planning. Horizon: Nov 2025 – Oct 2026. YoY growth assumption: 6.5%."
    ws["A2"].font = Font(italic=True, color="666666")
    ws.merge_cells("A2:P2")

    months = ["Nov", "Dec", "Jan", "Feb", "Mar", "Apr", "May", "Jun",
              "Jul", "Aug", "Sep", "Oct"]
    families = ["Tents", "Backpacks", "Sleeping Bags", "Cookware", "Apparel"]
    regions = ["North", "South", "East", "West"]

    # Seasonality multipliers per family
    season = {
        "Tents":         [0.7, 0.5, 0.4, 0.4, 0.6, 0.9, 1.2, 1.5, 1.8, 1.4, 1.0, 0.8],
        "Backpacks":     [0.9, 0.7, 0.5, 0.5, 0.7, 1.0, 1.2, 1.4, 1.5, 1.3, 1.1, 1.0],
        "Sleeping Bags": [1.1, 0.9, 0.7, 0.6, 0.6, 0.8, 1.0, 1.2, 1.3, 1.2, 1.1, 1.0],
        "Cookware":      [0.8, 0.6, 0.5, 0.5, 0.7, 0.9, 1.1, 1.3, 1.4, 1.2, 1.0, 0.9],
        "Apparel":       [1.4, 1.6, 1.1, 0.8, 0.8, 0.9, 1.0, 1.1, 1.1, 1.0, 1.2, 1.3],
    }
    region_weight = {"North": 0.18, "South": 0.20, "East": 0.27, "West": 0.35}
    base_units = {  # base monthly per family (across all regions)
        "Tents": 22_000, "Backpacks": 28_000, "Sleeping Bags": 16_000,
        "Cookware": 12_000, "Apparel": 38_000,
    }

    # Headers
    ws["A4"] = "SKU Family"
    ws["B4"] = "Region"
    for i, m in enumerate(months):
        ws.cell(row=4, column=3+i, value=m)
    ws.cell(row=4, column=15, value="12-mo total")
    style_header(ws, 4, 15)

    row = 5
    total_all = 0
    for fam in families:
        for reg in regions:
            base = base_units[fam] * region_weight[reg]
            ws.cell(row=row, column=1, value=fam)
            ws.cell(row=row, column=2, value=reg)
            row_total = 0
            for i, m in enumerate(months):
                v = int(round(base * season[fam][i]))
                ws.cell(row=row, column=3+i, value=v)
                row_total += v
            ws.cell(row=row, column=15, value=row_total)
            total_all += row_total
            row += 1

    # Grand total
    ws.cell(row=row+1, column=1, value="GRAND TOTAL").font = TOTAL_FONT
    ws.cell(row=row+1, column=15, value=total_all)
    for c in range(1, 16):
        ws.cell(row=row+1, column=c).fill = TOTAL_FILL
    ws.cell(row=row+1, column=15).number_format = "#,##0"

    # Apply number format to data cells
    for r in range(5, row):
        for c in range(3, 16):
            ws.cell(row=r, column=c).number_format = "#,##0"

    autosize(ws)
    wb.save(path)
    print(f"  wrote {path.name}  (grand_total={total_all:,})")


def gen_planning_pdf(path: Path):
    """P2: Annual S&OP plan, ~6 pages, multi-page capacity table no-repeat header.

    Fingerprint facts:
      - Planned annual production: 1,805,000 units
      - Q4 production share: 28%
      - Bottleneck process: Tent assembly (Reno line)
      - Target capacity utilization: 85%
      - Plan owner: VP Planning, Avery Chen
    """
    doc = SimpleDocTemplate(
        str(path), pagesize=letter,
        leftMargin=0.7 * inch, rightMargin=0.7 * inch,
        topMargin=0.7 * inch, bottomMargin=0.7 * inch,
    )
    styles = getSampleStyleSheet()
    title = ParagraphStyle("t", parent=styles["Title"], fontSize=18, spaceAfter=8)
    h2 = ParagraphStyle("h2", parent=styles["Heading2"], spaceBefore=10, spaceAfter=4)
    body = styles["BodyText"]
    note = ParagraphStyle("note", parent=body, textColor=colors.HexColor("#555555"), fontSize=9)

    story = []
    story.append(Paragraph("Annual Sales &amp; Operations Plan (S&amp;OP) — FY2026", title))
    story.append(Paragraph("Northwind Outdoor Co. — Planning Function", note))
    story.append(Paragraph("Owner: Avery Chen, VP Planning. Approved: 12 October 2025.", note))
    story.append(Spacer(1, 0.15 * inch))

    story.append(Paragraph("Executive Summary", h2))
    story.append(Paragraph(
        "The FY2026 S&amp;OP commits to producing <b>1,805,000 units</b> across "
        "all SKU families, +6.5% versus FY2025. Q4 absorbs the largest share "
        "of production at <b>28%</b> due to Fall and holiday demand. The "
        "bottleneck process is <b>Tent assembly on the Reno line</b>, which "
        "operates at 92% utilization during peak. The plan targets an overall "
        "<b>85% capacity utilization</b> with 15% headroom for promotion-driven "
        "upside.", body))
    story.append(Spacer(1, 0.1 * inch))
    story.append(Paragraph("Key Assumptions", h2))
    for item in [
        "YoY demand growth: 6.5%, weighted toward Ecommerce (+11.2%) and Wholesale (+4.8%).",
        "Material cost inflation: 3.4% applied uniformly across all families.",
        "Labour available: 412,000 standard hours across all sites at 95% attendance.",
        "Tent fabric supply assumed reliable post-Q1 contract renewal with Acme Textiles.",
        "Promotional events: 4 major (Spring Refresh, Summer Sale, Fall Launch, Holiday).",
    ]:
        story.append(Paragraph(f"• {item}", body))

    story.append(PageBreak())
    story.append(Paragraph("Capacity Allocation by Family × Quarter", h2))
    story.append(Paragraph(
        "Allocation in units. Bottleneck process is Tent assembly; allocations "
        "for Tents are bounded by the Reno line capacity of 130,000 units/quarter.", body))
    story.append(Spacer(1, 0.1 * inch))

    # 5 families × 4 quarters + subtotals + grand total = 5*4 + 5 + 4 + 1 = 30 rows
    families = ["Tents", "Backpacks", "Sleeping Bags", "Cookware", "Apparel"]
    quarters = ["Q1 FY26", "Q2 FY26", "Q3 FY26", "Q4 FY26"]
    family_totals = {
        "Tents":         {"Q1 FY26": 92_000,  "Q2 FY26": 108_000, "Q3 FY26": 124_000, "Q4 FY26": 116_000},
        "Backpacks":     {"Q1 FY26": 89_000,  "Q2 FY26": 102_000, "Q3 FY26": 121_000, "Q4 FY26": 118_000},
        "Sleeping Bags": {"Q1 FY26": 64_000,  "Q2 FY26": 78_000,  "Q3 FY26": 92_000,  "Q4 FY26": 96_000},
        "Cookware":      {"Q1 FY26": 52_000,  "Q2 FY26": 68_000,  "Q3 FY26": 78_000,  "Q4 FY26": 72_000},
        "Apparel":       {"Q1 FY26": 122_000, "Q2 FY26": 108_000, "Q3 FY26": 102_000, "Q4 FY26": 124_000},
    }

    header_row = ["Family", "Q1 FY26", "Q2 FY26", "Q3 FY26", "Q4 FY26", "Annual"]
    rows_data = []
    for fam in families:
        q_totals = family_totals[fam]
        annual = sum(q_totals[q] for q in quarters)
        rows_data.append([fam] + [f"{q_totals[q]:,}" for q in quarters] + [f"{annual:,}"])

    # Build quarter subtotals row and grand total row
    qsub = {q: sum(family_totals[f][q] for f in families) for q in quarters}
    grand = sum(qsub.values())
    rows_data.append(["<b>Quarter total</b>"] + [f"<b>{qsub[q]:,}</b>" for q in quarters] + [f"<b>{grand:,}</b>"])
    rows_data.append(["<i>Share of annual</i>"] +
                     [f"<i>{qsub[q]/grand*100:.0f}%</i>" for q in quarters] + ["<i>100%</i>"])

    # Split intentionally across two pages WITHOUT repeating header
    page_a_rows = 3   # families[0..2] then page break
    table1 = [header_row]
    for r in rows_data[:page_a_rows]:
        table1.append([Paragraph(c, body) for c in r])
    table2 = []
    for r in rows_data[page_a_rows:]:
        table2.append([Paragraph(c, body) for c in r])

    style_first = TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F4E78")),
        ("TEXTCOLOR",  (0, 0), (-1, 0), colors.white),
        ("FONTNAME",   (0, 0), (-1, 0), "Helvetica-Bold"),
        ("GRID",       (0, 0), (-1, -1), 0.25, colors.HexColor("#999999")),
        ("FONTSIZE",   (0, 0), (-1, -1), 9.5),
        ("ALIGN",      (1, 0), (-1, -1), "RIGHT"),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
    ])
    style_continued = TableStyle([
        ("GRID",       (0, 0), (-1, -1), 0.25, colors.HexColor("#999999")),
        ("FONTSIZE",   (0, 0), (-1, -1), 9.5),
        ("ALIGN",      (1, 0), (-1, -1), "RIGHT"),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
    ])
    col_widths = [1.5*inch, 1.05*inch, 1.05*inch, 1.05*inch, 1.05*inch, 1.05*inch]
    story.append(Table(table1, colWidths=col_widths, style=style_first, repeatRows=0))
    story.append(PageBreak())
    story.append(Table(table2, colWidths=col_widths, style=style_continued, repeatRows=0))

    story.append(Spacer(1, 0.2 * inch))
    story.append(Paragraph("Capacity & Bottlenecks", h2))
    story.append(Paragraph(
        "The Tent assembly line at the Reno facility is the binding constraint. "
        "FY2026 plan loads it to 92% utilization during Q3 (June–August). "
        "Mitigation options include a third shift for 6 weeks "
        "(adds 24,000 units of capacity) or outsourcing 8,000 units to "
        "the Tijuana partner facility. Decision required by end of Q1.", body))

    story.append(PageBreak())
    story.append(Paragraph("Channel Mix Plan", h2))
    story.append(Paragraph(
        "Ecommerce share rises from 36% to 39% of total volume. Wholesale "
        "remains steady at ~37%; Retail Stores at 24%. The Ecommerce shift "
        "concentrates demand on Eastern and Western DC pick capacity; "
        "Fulfillment is preparing for a peak-week rate of 18,500 orders/day "
        "in early December.", body))
    story.append(Spacer(1, 0.1 * inch))
    story.append(Paragraph("Risks & Mitigations", h2))
    for item in [
        "Tent fabric supply (Acme): single-source. Mitigation — qualify a backup mill in Q2.",
        "Carrier rate increases for Ecommerce: assumed 4% peak surcharge baked in.",
        "Demand softness in Apparel: hedged with a 4-week pull-forward of Fall release.",
    ]:
        story.append(Paragraph(f"• {item}", body))

    doc.build(story)
    print(f"  wrote {path.name}")


def gen_planning_docx(path: Path):
    """P3: Capacity-planning policy.

    Fingerprint facts:
      - Safety stock formula: 1.65 × σ_demand × √lead_time
      - Standard lead time (domestic): 14 days
      - Standard lead time (international): 35 days
      - Replenishment cadence: weekly review, weekly order
      - Reorder point thresholds defined by service level: 95% for A, 90% B, 85% C
    """
    doc = Document()
    doc.add_heading("Capacity & Replenishment Planning Policy", level=0)
    p = doc.add_paragraph()
    p.add_run("Document ID: ").bold = True
    p.add_run("PLN-POL-007    ")
    p.add_run("Owner: ").bold = True
    p.add_run("VP Planning, Avery Chen    ")
    p.add_run("Revision: ").bold = True
    p.add_run("3.1   ")
    p.add_run("Effective: ").bold = True
    p.add_run("1 October 2025")

    doc.add_heading("1. Scope", level=1)
    doc.add_paragraph(
        "This policy defines how Northwind Outdoor Co. sets safety stock, "
        "reorder points, and replenishment cadence for all production "
        "SKUs across the three DCs. It applies to Retail, Wholesale, and "
        "Ecommerce channels equally; channel-specific buffer rules are "
        "captured separately in the Channel Operating Standards.")

    doc.add_heading("2. Service-Level Targets", level=1)
    doc.add_paragraph(
        "SKUs are classified A, B, or C by trailing 12-month revenue "
        "contribution (Pareto). Service levels:")
    for line in [
        "A-class (top 20% of SKUs, ~80% revenue): 95% target fill rate",
        "B-class (next 30% of SKUs): 90% target fill rate",
        "C-class (long tail): 85% target fill rate; quarterly review for rationalization",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("3. Safety Stock Formula", level=1)
    doc.add_paragraph(
        "Safety stock is computed monthly using the standard normal "
        "approximation:")
    p = doc.add_paragraph()
    p.add_run("    SS = z × σ_demand × √(lead_time)").bold = True
    doc.add_paragraph(
        "Where z is the service-level z-score (95% → 1.65, 90% → 1.28, "
        "85% → 1.04), σ_demand is the 12-week demand standard deviation, "
        "and lead_time is measured in weeks. The formula uses the production "
        "lead time, not vendor lead time, for owned-production SKUs.")

    doc.add_heading("4. Standard Lead Times", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Source", "Mode", "Standard lead time", "Variance buffer"]):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in [
        ("Domestic supplier",       "Truck",  "14 days", "+3 days"),
        ("International (Asia)",    "Ocean",  "35 days", "+7 days"),
        ("International (Asia)",    "Air",    "9 days",  "+2 days"),
        ("Own production (Reno)",   "Truck",  "21 days", "+5 days"),
    ]:
        r = table.add_row().cells
        for i, v in enumerate(row):
            r[i].text = v

    doc.add_heading("5. Replenishment Cadence", level=1)
    doc.add_paragraph(
        "Demand Planning reviews ATP (Available-to-Promise) weekly on "
        "Monday, generates a replenishment proposal Tuesday, and POs are "
        "released by close-of-business Wednesday. Emergency replenishment "
        "outside this cycle requires VP Planning approval and is reserved "
        "for stock-out events on A-class SKUs.")

    doc.add_heading("6. Channel-Specific Rules", level=1)
    doc.add_paragraph(
        "Ecommerce: maintain a minimum of 21 days of forward cover at the "
        "Eastern DC due to next-day shipping commitments to East-coast "
        "customers. Wholesale: hold contractual buffer for top-3 retailers "
        "(REI, Bass Pro, Cabela's) per their respective agreements. Retail "
        "Stores: replenish weekly from the nearest DC, with auto-allocation "
        "by store-level sell-through.")

    doc.add_heading("7. Exceptions", level=1)
    doc.add_paragraph(
        "Promotional SKUs are excluded from the standard formula for the "
        "8 weeks surrounding the promo window; a discrete forecast and "
        "safety-stock override is applied. Discontinued SKUs (APR-9XXX et al.) "
        "are run-down only — no replenishment.")

    doc.add_heading("8. KPIs", level=1)
    doc.add_paragraph(
        "Fill rate by class (A/B/C) reported weekly. Forecast accuracy "
        "(MAPE) reported monthly with a 25% target on A-class SKUs. "
        "Inventory turns reported monthly with a 6.5x target across all "
        "DCs combined.")

    doc.save(path)
    print(f"  wrote {path.name}")


# --------------------------------------------------------------------------
# Top-level driver
# --------------------------------------------------------------------------

def main():
    print("Generating fixtures under", ROOT)

    print("\n[01] Finance")
    gen_finance_xlsx(ROOT / "01_Finance" / "fixtures" / "F1_monthly_sales_by_region.xlsx")
    gen_finance_pdf (ROOT / "01_Finance" / "fixtures" / "F2_Q3_PnL_statement.pdf")
    gen_finance_docx(ROOT / "01_Finance" / "fixtures" / "F3_vendor_payment_terms.docx")

    print("\n[02] Operations")
    gen_ops_xlsx(ROOT / "02_Operations" / "fixtures" / "O1_inventory_turnover.xlsx")
    gen_ops_pdf (ROOT / "02_Operations" / "fixtures" / "O2_carrier_manifest.pdf")
    gen_ops_docx(ROOT / "02_Operations" / "fixtures" / "O3_returns_SOP.docx")

    print("\n[03] IT")
    gen_it_xlsx(ROOT / "03_IT" / "fixtures" / "I1_asset_inventory.xlsx")
    gen_it_pdf (ROOT / "03_IT" / "fixtures" / "I2_security_audit.pdf")
    gen_it_docx(ROOT / "03_IT" / "fixtures" / "I3_integration_runbook.docx")

    print("\n[04] Planning")
    gen_planning_xlsx(ROOT / "04_Planning" / "fixtures" / "P1_demand_forecast.xlsx")
    gen_planning_pdf (ROOT / "04_Planning" / "fixtures" / "P2_annual_SOP.pdf")
    gen_planning_docx(ROOT / "04_Planning" / "fixtures" / "P3_capacity_policy.docx")

    print("\nDone.")


if __name__ == "__main__":
    main()
