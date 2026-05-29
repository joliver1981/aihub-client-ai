"""
Tool-level smoke tests for the four artifact-producing tools added to GeneralAgent:
    create_csv, create_excel, create_text_file, create_word_doc

These are DETERMINISTIC tests — they exercise the tool bodies directly,
not the LLM. They cover:
    - Small / medium / large input sizes
    - Empty arrays
    - Single-row inputs
    - Missing keys in some rows
    - Special characters (commas, quotes, newlines)
    - Unicode (non-ASCII names + emoji)
    - Multi-sheet Excel
    - Bad input (wrong types, malformed JSON)

Run:
    python tests_v2/file_tools_smoke.py
"""

from __future__ import annotations

import io
import json
import shutil
import sys
import traceback
from pathlib import Path

REPO_ROOT = Path(__file__).parent.parent
sys.path.insert(0, str(REPO_ROOT))

CONV_ID = "test_file_tools_smoke"

import active_chat_context
import chat_file_manager
from GeneralAgent import (  # noqa: E402
    create_csv, create_excel, create_text_file, create_word_doc,
)

PASS, FAIL = 0, 0


def _result(name: str, ok: bool, detail: str = ""):
    global PASS, FAIL
    marker = "PASS" if ok else "FAIL"
    if ok:
        PASS += 1
    else:
        FAIL += 1
    print(f"[{marker}] {name}" + (f" — {detail}" if detail else ""))


def _invoke(tool, **kwargs) -> str:
    """Invoke a LangChain @tool and return its string result."""
    return tool.invoke(kwargs)


def _parse_blocks(result_str: str):
    """Parse a tool's JSON-block return; returns list-of-blocks or raises."""
    return json.loads(result_str)


def _read_artifact(block: dict) -> bytes:
    path = chat_file_manager.get_output_path(CONV_ID, block["artifact_id"])
    assert path is not None, f"artifact not on disk: {block}"
    return path.read_bytes()


def _setup():
    chat_file_manager.delete_conversation(CONV_ID)


def _teardown():
    chat_file_manager.delete_conversation(CONV_ID)


# ─── create_csv ──────────────────────────────────────────────────────────

def test_csv_small():
    rows = [
        {"date": "2025-Q1", "region": "North", "revenue": 12000},
        {"date": "2025-Q1", "region": "South", "revenue": 8500},
    ]
    result = _invoke(create_csv, name="q1_revenue", rows=json.dumps(rows))
    blocks = _parse_blocks(result)
    body = _read_artifact(blocks[0]).decode("utf-8-sig")
    lines = body.strip().splitlines()
    assert lines[0] == "date,region,revenue", lines[0]
    assert len(lines) == 3  # header + 2 rows
    _result("csv: small (2 rows)", True)


def test_csv_medium():
    rows = [{"id": i, "name": f"user_{i}", "score": i * 3} for i in range(100)]
    result = _invoke(create_csv, name="medium", rows=json.dumps(rows))
    blocks = _parse_blocks(result)
    body = _read_artifact(blocks[0]).decode("utf-8-sig")
    assert body.count("\n") >= 100
    _result("csv: medium (100 rows)", True)


def test_csv_large():
    rows = [{"id": i, "value": f"v{i}"} for i in range(1000)]
    result = _invoke(create_csv, name="big", rows=json.dumps(rows))
    blocks = _parse_blocks(result)
    body = _read_artifact(blocks[0]).decode("utf-8-sig")
    assert body.count("\n") >= 1000
    _result("csv: large (1000 rows)", True)


def test_csv_empty():
    result = _invoke(create_csv, name="empty", rows="[]")
    ok = "empty" in result.lower()
    _result("csv: empty array returns clean error", ok, detail=result[:80])


def test_csv_single_row():
    result = _invoke(create_csv, name="one", rows=json.dumps([{"x": 1}]))
    blocks = _parse_blocks(result)
    body = _read_artifact(blocks[0]).decode("utf-8-sig").strip().splitlines()
    assert body == ["x", "1"]
    _result("csv: single row", True)


def test_csv_missing_keys():
    rows = [
        {"a": 1, "b": 2, "c": 3},
        {"a": 10},          # missing b, c
        {"b": 20, "c": 30}  # missing a
    ]
    result = _invoke(create_csv, name="sparse", rows=json.dumps(rows))
    blocks = _parse_blocks(result)
    lines = _read_artifact(blocks[0]).decode("utf-8-sig").strip().splitlines()
    # Missing keys should yield empty cells
    assert lines[0] == "a,b,c"
    assert lines[2] == "10,,"
    assert lines[3] == ",20,30"
    _result("csv: missing keys -> empty cells", True)


def test_csv_special_chars():
    rows = [
        {"col": 'has "quotes" inside'},
        {"col": "has,comma,inside"},
        {"col": "has\nnewline"},
    ]
    result = _invoke(create_csv, name="special", rows=json.dumps(rows))
    blocks = _parse_blocks(result)
    body = _read_artifact(blocks[0]).decode("utf-8-sig")
    # Each problematic cell must be properly quoted
    import csv
    reader = csv.DictReader(io.StringIO(body))
    parsed = list(reader)
    assert parsed[0]["col"] == 'has "quotes" inside'
    assert parsed[1]["col"] == "has,comma,inside"
    assert parsed[2]["col"] == "has\nnewline"
    _result("csv: special chars (quotes, commas, newlines)", True)


def test_csv_unicode():
    rows = [{"name": "Zoë", "emoji": "🎉"}, {"name": "François", "emoji": "✨"}]
    result = _invoke(create_csv, name="unicode", rows=json.dumps(rows))
    blocks = _parse_blocks(result)
    body = _read_artifact(blocks[0]).decode("utf-8-sig")
    assert "Zoë" in body and "🎉" in body and "François" in body
    _result("csv: unicode + emoji", True)


def test_csv_bad_json():
    result = _invoke(create_csv, name="bad", rows="this is not json")
    ok = "invalid json" in result.lower()
    _result("csv: malformed JSON -> clean error", ok, detail=result[:80])


def test_csv_non_array():
    result = _invoke(create_csv, name="bad", rows='{"not": "array"}')
    ok = "array" in result.lower()
    _result("csv: non-array -> clean error", ok, detail=result[:80])


# ─── create_excel ────────────────────────────────────────────────────────

def test_excel_single_sheet():
    sheets = [{"name": "Sales", "rows": [
        {"region": "North", "revenue": 12000},
        {"region": "South", "revenue": 8500},
    ]}]
    result = _invoke(create_excel, name="single", sheets=json.dumps(sheets))
    blocks = _parse_blocks(result)
    raw = _read_artifact(blocks[0])
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(raw))
    assert wb.sheetnames == ["Sales"]
    ws = wb["Sales"]
    assert ws.cell(1, 1).value == "region"
    assert ws.cell(1, 1).font.bold
    assert ws.freeze_panes == "A2"
    assert ws.cell(2, 2).value == 12000
    _result("excel: single sheet + header style + freeze", True)


def test_excel_multi_sheet():
    sheets = [
        {"name": "Q1", "rows": [{"region": "N", "rev": 12}, {"region": "S", "rev": 8}]},
        {"name": "Q2", "rows": [{"region": "N", "rev": 15}, {"region": "S", "rev": 9}]},
        {"name": "Q3", "rows": [{"region": "N", "rev": 18}]},
    ]
    result = _invoke(create_excel, name="multi", sheets=json.dumps(sheets))
    blocks = _parse_blocks(result)
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(_read_artifact(blocks[0])))
    assert wb.sheetnames == ["Q1", "Q2", "Q3"]
    assert wb["Q1"].cell(2, 2).value == 12
    assert wb["Q3"].cell(2, 2).value == 18
    _result("excel: multi-sheet (3 sheets)", True)


def test_excel_long_sheet_name():
    long_name = "This is a very long sheet name that exceeds Excel limit"
    sheets = [{"name": long_name, "rows": [{"x": 1}]}]
    result = _invoke(create_excel, name="long", sheets=json.dumps(sheets))
    blocks = _parse_blocks(result)
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(_read_artifact(blocks[0])))
    assert all(len(s) <= 31 for s in wb.sheetnames)
    _result("excel: sheet name truncated to 31 chars", True)


def test_excel_missing_keys():
    sheets = [{"name": "Sparse", "rows": [
        {"a": 1, "b": 2}, {"a": 3}, {"b": 5}
    ]}]
    result = _invoke(create_excel, name="sparse_xl", sheets=json.dumps(sheets))
    blocks = _parse_blocks(result)
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(_read_artifact(blocks[0])))
    ws = wb["Sparse"]
    assert ws.cell(2, 1).value == 1 and ws.cell(2, 2).value == 2
    assert ws.cell(3, 1).value == 3 and ws.cell(3, 2).value in (None, "")
    assert ws.cell(4, 1).value in (None, "") and ws.cell(4, 2).value == 5
    _result("excel: missing keys -> empty cells", True)


def test_excel_unicode():
    sheets = [{"name": "Unicode", "rows": [
        {"name": "Zoë", "emoji": "🎉"}, {"name": "Mañana", "emoji": "✨"}
    ]}]
    result = _invoke(create_excel, name="unicode_xl", sheets=json.dumps(sheets))
    blocks = _parse_blocks(result)
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(_read_artifact(blocks[0])))
    ws = wb["Unicode"]
    assert ws.cell(2, 1).value == "Zoë"
    assert ws.cell(2, 2).value == "🎉"
    _result("excel: unicode + emoji", True)


def test_excel_empty_sheets():
    result = _invoke(create_excel, name="empty", sheets="[]")
    ok = "non-empty" in result.lower()
    _result("excel: empty array -> clean error", ok, detail=result[:80])


# ─── create_text_file ────────────────────────────────────────────────────

def test_text_txt():
    result = _invoke(create_text_file, name="notes", content="Hello world.\nLine 2.", format="txt")
    blocks = _parse_blocks(result)
    assert blocks[0]["name"].endswith(".txt")
    assert _read_artifact(blocks[0]).decode("utf-8") == "Hello world.\nLine 2."
    _result("text: txt", True)


def test_text_markdown():
    md = "# Title\n\n- bullet one\n- bullet two\n"
    result = _invoke(create_text_file, name="readme", content=md, format="md")
    blocks = _parse_blocks(result)
    assert blocks[0]["name"].endswith(".md")
    assert _read_artifact(blocks[0]).decode("utf-8") == md
    _result("text: markdown", True)


def test_text_json():
    payload = '{"k": "v", "list": [1, 2, 3]}'
    result = _invoke(create_text_file, name="cfg", content=payload, format="json")
    blocks = _parse_blocks(result)
    assert blocks[0]["name"].endswith(".json")
    assert _read_artifact(blocks[0]).decode("utf-8") == payload
    _result("text: json", True)


def test_text_html():
    html = "<html><body><h1>Hi</h1></body></html>"
    result = _invoke(create_text_file, name="page", content=html, format="html")
    blocks = _parse_blocks(result)
    assert blocks[0]["name"].endswith(".html")
    _result("text: html", True)


def test_text_unicode():
    result = _invoke(create_text_file, name="emoji", content="Zoë says 🎉", format="txt")
    blocks = _parse_blocks(result)
    assert _read_artifact(blocks[0]).decode("utf-8") == "Zoë says 🎉"
    _result("text: unicode", True)


def test_text_bad_format():
    result = _invoke(create_text_file, name="x", content="anything", format="exe")
    ok = "format must be" in result.lower()
    _result("text: bad format -> clean error", ok, detail=result[:80])


# ─── create_word_doc ─────────────────────────────────────────────────────

def test_word_basic():
    sections = [
        {"heading": "Intro", "paragraphs": ["First paragraph.", "Second paragraph."]},
        {"heading": "Details", "paragraphs": ["More content here."]},
    ]
    result = _invoke(create_word_doc, name="report",
                     title="My Report", sections=json.dumps(sections))
    blocks = _parse_blocks(result)
    raw = _read_artifact(blocks[0])
    from docx import Document
    doc = Document(io.BytesIO(raw))
    texts = [p.text for p in doc.paragraphs]
    assert "My Report" in texts
    assert "Intro" in texts
    assert "First paragraph." in texts
    assert "Details" in texts
    _result("word: title + sections + paragraphs", True)


def test_word_with_table():
    sections = [{
        "heading": "Quarterly Numbers",
        "paragraphs": ["Summary below."],
        "table": {
            "headers": ["Quarter", "Revenue", "Units"],
            "rows": [["Q1", "12k", "340"], ["Q2", "15k", "410"]],
        }
    }]
    result = _invoke(create_word_doc, name="numbers", title="Q1-Q2",
                     sections=json.dumps(sections))
    blocks = _parse_blocks(result)
    raw = _read_artifact(blocks[0])
    from docx import Document
    doc = Document(io.BytesIO(raw))
    assert len(doc.tables) == 1
    t = doc.tables[0]
    assert t.rows[0].cells[0].text == "Quarter"
    assert t.rows[1].cells[1].text == "12k"
    _result("word: with table", True)


def test_word_unicode():
    sections = [{"heading": "Café", "paragraphs": ["Zoë 🎉 says hi."]}]
    result = _invoke(create_word_doc, name="uni", title="Étude",
                     sections=json.dumps(sections))
    blocks = _parse_blocks(result)
    from docx import Document
    doc = Document(io.BytesIO(_read_artifact(blocks[0])))
    texts = [p.text for p in doc.paragraphs]
    assert "Étude" in texts and "Café" in texts and "Zoë 🎉 says hi." in texts
    _result("word: unicode", True)


def test_word_empty_sections():
    result = _invoke(create_word_doc, name="empty", title="Just Title",
                     sections="[]")
    blocks = _parse_blocks(result)
    from docx import Document
    doc = Document(io.BytesIO(_read_artifact(blocks[0])))
    texts = [p.text for p in doc.paragraphs]
    assert "Just Title" in texts
    _result("word: title-only (no sections)", True)


def test_word_bad_json():
    result = _invoke(create_word_doc, name="bad", title="x", sections="not json")
    ok = "invalid json" in result.lower()
    _result("word: malformed JSON -> clean error", ok, detail=result[:80])


# ─── Conversation-context guard ──────────────────────────────────────────

def test_no_conversation_context():
    """Without a bound conversation, the tool should refuse cleanly."""
    # We're outside the bind_active_chat block here.
    result = create_csv.invoke({"name": "x", "rows": '[{"a":1}]'})
    ok = "no active conversation" in result.lower()
    _result("guard: no active conversation -> clean error", ok, detail=result[:80])


# ─── Runner ──────────────────────────────────────────────────────────────

def main():
    _setup()
    try:
        with active_chat_context.bind_active_chat(CONV_ID, user_id=42, agent_id=9999):
            for test in (
                test_csv_small, test_csv_medium, test_csv_large, test_csv_empty,
                test_csv_single_row, test_csv_missing_keys, test_csv_special_chars,
                test_csv_unicode, test_csv_bad_json, test_csv_non_array,
                test_excel_single_sheet, test_excel_multi_sheet,
                test_excel_long_sheet_name, test_excel_missing_keys,
                test_excel_unicode, test_excel_empty_sheets,
                test_text_txt, test_text_markdown, test_text_json,
                test_text_html, test_text_unicode, test_text_bad_format,
                test_word_basic, test_word_with_table, test_word_unicode,
                test_word_empty_sections, test_word_bad_json,
            ):
                try:
                    test()
                except Exception as e:
                    _result(test.__name__, False, detail=f"{type(e).__name__}: {e}")
                    traceback.print_exc()
        # Run the no-context guard test OUTSIDE the bind block
        try:
            test_no_conversation_context()
        except Exception as e:
            _result("test_no_conversation_context", False, detail=str(e))
    finally:
        _teardown()

    print(f"\n{PASS} passed, {FAIL} failed")
    sys.exit(0 if FAIL == 0 else 1)


if __name__ == "__main__":
    main()
