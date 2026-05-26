"""Generate Word .docx fixture files for the agent-knowledge Word COMPETENCY suite.

Each fixture probes a different extraction / retrieval / reasoning surface:

   01_clean_handbook.docx          baseline body + headings + bullets
   02_tables_heavy_report.docx     facts buried in many tables
   03_embedded_charts_kpis.docx    text + matplotlib chart PNGs (caption survives?)
   04_tracked_changes_contract.docx doc with tracked insertions/deletions
   05_long_doc_toc_footnotes.docx  30+ pages, TOC, footnotes, cross-refs

Every fixture has at least one **fingerprint** value the test can ask about.

Run:
    C:\\Users\\james\\miniconda3\\envs\\aihub2.1\\python.exe _generate.py
"""
from __future__ import annotations

import io
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUT_DIR = Path(__file__).resolve().parent


# ============================================================================
# Helpers
# ============================================================================

def _set_default_font(doc, name="Calibri", size=11):
    style = doc.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(size)


def _heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def _table(doc, headers, rows, style="Light Grid Accent 1"):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = style
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
    return t


# ============================================================================
# 01 — Clean handbook (baseline)
# ============================================================================

def fixture_clean_handbook():
    doc = Document()
    _set_default_font(doc)

    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("VERIDIAN LABS — Employee Handbook 2026")
    r.font.size = Pt(22)
    r.font.bold = True

    doc.add_paragraph(
        "Welcome to Veridian Labs. This handbook summarises our policies, "
        "benefits and code of conduct. Effective date: April 1, 2026. "
        "Owner: People Operations (lead: Sasha Kovalenko)."
    )

    doc.add_page_break()
    _heading(doc, "1. Company Overview", level=1)
    doc.add_paragraph(
        "Veridian Labs was founded in March 2014 by Dr. Anika Vorhees "
        "and Marcus Holloway in Cambridge, Massachusetts. The company "
        "designs precision optical sensors for industrial automation and "
        "ships to 38 countries. As of January 2026 the company employs "
        "412 people across 5 sites."
    )

    _heading(doc, "1.1 Sites", level=2)
    sites = [
        "Cambridge, MA — Headquarters (engineering + HQ functions)",
        "Eindhoven, Netherlands — European R&D and customer support",
        "Penang, Malaysia — Manufacturing site #1 (optical assemblies)",
        "Bangalore, India — Software platform and DevOps",
        "Austin, TX — North-American sales hub",
    ]
    for s in sites:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(s)

    _heading(doc, "2. Vacation & Leave", level=1)
    doc.add_paragraph(
        "All full-time staff accrue paid time off based on tenure. "
        "The tenure brackets and corresponding annual days are:"
    )
    bullets = [
        "0–2 years of service: 18 days per year",
        "3–5 years of service: 22 days per year",
        "6–10 years of service: 25 days per year",
        "11+ years of service: 28 days per year",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_paragraph(
        "Unused vacation rolls over up to a cap of 10 days per year. "
        "Any amount above the cap is forfeit on December 31."
    )

    _heading(doc, "3. Travel & Per Diem", level=1)
    doc.add_paragraph(
        "Domestic travel is approved by line managers. International "
        "travel requires VP-level approval. Per diem amounts (USD) by "
        "city tier are:"
    )
    bullets = [
        "Tier 1 (NYC, SF, London, Tokyo, Singapore): $115/day",
        "Tier 2 (most major capitals): $90/day",
        "Tier 3 (everywhere else): $70/day",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    _heading(doc, "4. Performance Reviews", level=1)
    doc.add_paragraph(
        "Performance reviews are conducted semi-annually, in April and "
        "October. Compensation adjustments take effect on May 1 (Spring "
        "cycle) and November 1 (Fall cycle). The review uses a 5-point "
        "scale: Exceeds, Strong, Meets, Developing, Not Meeting."
    )

    _heading(doc, "5. Conduct & Reporting", level=1)
    doc.add_paragraph(
        "Concerns may be reported anonymously to ethics@veridianlabs.com. "
        "Reports are reviewed by the People Operations lead, Sasha "
        "Kovalenko, and the General Counsel, Bao Tran. Retaliation "
        "against reporters is grounds for termination."
    )

    # Fingerprint anchor
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("ANCHOR — Veridian Labs founded March 2014 in Cambridge, MA "
              "by Anika Vorhees and Marcus Holloway. 412 employees at "
              "5 sites as of January 2026.").italic = True

    out = OUT_DIR / "01_clean_handbook.docx"
    doc.save(out)
    print(f"Wrote {out}")


# ============================================================================
# 02 — Tables-heavy report
# ============================================================================

def fixture_tables_heavy_report():
    doc = Document()
    _set_default_font(doc)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Eldoria Logistics — Q1 2026 Operations Report")
    r.font.size = Pt(20)
    r.font.bold = True

    doc.add_paragraph(
        "Prepared by COO Office, April 18, 2026. Distribution: Executive "
        "team + Board. All values in USD unless noted."
    )

    _heading(doc, "1. Hub Throughput (parcels / day)", level=1)
    _table(doc,
           ["Hub", "Jan 2026", "Feb 2026", "Mar 2026", "Q1 avg"],
           [
               ("Memphis",      48200, 47800, 51300, 49100),
               ("Louisville",   31500, 32100, 34200, 32600),
               ("Cologne",      28900, 29400, 30100, 29467),
               ("Rotterdam",    22100, 22500, 23800, 22800),
               ("Singapore",    19400, 20200, 21800, 20467),
               ("Dubai",        12800, 13400, 14900, 13700),
           ])

    _heading(doc, "2. On-Time Delivery Rate by Service Tier", level=1)
    _table(doc,
           ["Tier", "Target SLA", "Q1 actual", "Variance"],
           [
               ("Priority overnight", "99.0%", "98.6%", "-0.4 pp"),
               ("Express 2-day",      "97.5%", "97.9%", "+0.4 pp"),
               ("Standard 4-day",     "95.0%", "96.2%", "+1.2 pp"),
               ("Economy 7-day",      "92.0%", "93.5%", "+1.5 pp"),
           ])

    _heading(doc, "3. Top 10 Lanes by Revenue", level=1)
    _table(doc,
           ["Rank", "Origin", "Destination", "Q1 revenue ($K)",
            "Parcels (K)"],
           [
               (1,  "Memphis",    "Atlanta",     2840, 1240),
               (2,  "Memphis",    "Dallas",      2210,  985),
               (3,  "Louisville", "Chicago",     2105,  890),
               (4,  "Cologne",    "Berlin",      1980,  810),
               (5,  "Memphis",    "Miami",       1810,  720),
               (6,  "Cologne",    "Paris",       1620,  640),
               (7,  "Rotterdam",  "Amsterdam",   1480,  595),
               (8,  "Singapore",  "Kuala Lumpur",1310,  528),
               (9,  "Louisville", "Indianapolis",1185,  475),
               (10, "Dubai",      "Riyadh",       980,  390),
           ])

    _heading(doc, "4. Damage & Loss Claims", level=1)
    _table(doc,
           ["Hub", "Claims filed", "Avg claim ($)", "Total paid ($K)"],
           [
               ("Memphis",    412, 187, 77.0),
               ("Louisville", 198, 162, 32.1),
               ("Cologne",    142, 215, 30.5),
               ("Rotterdam",   88, 198, 17.4),
               ("Singapore",   62, 224, 13.9),
               ("Dubai",       38, 256,  9.7),
           ])

    _heading(doc, "5. Anchor Facts", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "Memphis was Q1's highest-throughput hub at 49,100 parcels/day. "
        "Priority Overnight was the only service tier to miss its SLA "
        "(98.6% vs 99.0%, -0.4 pp). The Memphis→Atlanta lane generated "
        "$2,840K in Q1, our top revenue lane. Cologne paid the highest "
        "average damage claim at $215 per claim."
    ).italic = True

    out = OUT_DIR / "02_tables_heavy_report.docx"
    doc.save(out)
    print(f"Wrote {out}")


# ============================================================================
# 03 — Embedded charts + KPI commentary
# ============================================================================

def _chart_kpi_trend() -> bytes:
    """Line chart: monthly KPI trend with labeled values."""
    fig, ax = plt.subplots(figsize=(6, 3), dpi=150)
    months = ["Jul", "Aug", "Sep", "Oct", "Nov", "Dec", "Jan", "Feb", "Mar"]
    values = [82.4, 83.1, 84.0, 85.2, 86.0, 86.4, 87.2, 88.1, 89.7]
    ax.plot(months, values, marker="o", linewidth=2, color="#1F3864")
    ax.set_ylabel("Active accounts (thousands)")
    ax.set_title("Atlas Networks — Monthly Active Accounts (last 9 months)")
    for x, y in zip(months, values):
        ax.annotate(f"{y:.1f}K", (x, y), textcoords="offset points",
                    xytext=(0, 8), ha="center", fontsize=8)
    ax.grid(axis="y", linestyle=":", alpha=0.4)
    fig.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png")
    plt.close(fig)
    return buf.getvalue()


def _chart_arr_breakdown() -> bytes:
    """Pie chart of ARR by product line."""
    fig, ax = plt.subplots(figsize=(5, 4), dpi=150)
    labels = ["Atlas Core", "Atlas Pro", "Atlas Edge", "Add-ons",
              "Professional Services"]
    sizes = [42, 28, 14, 9, 7]
    colors = ["#1F3864", "#2E74B5", "#5B9BD5", "#9DC3E6", "#BDD7EE"]
    ax.pie(sizes, labels=labels, colors=colors, autopct="%1.0f%%",
           startangle=90, wedgeprops=dict(edgecolor="white"))
    ax.set_title("Atlas Networks — FY2026 ARR by Product Line ($248M total)")
    fig.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png")
    plt.close(fig)
    return buf.getvalue()


def fixture_embedded_charts_kpis():
    doc = Document()
    _set_default_font(doc)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Atlas Networks — FY2026 Board Update")
    r.font.size = Pt(22)
    r.font.bold = True

    doc.add_paragraph(
        "Prepared by the CEO's office for the May 18, 2026 board meeting. "
        "Atlas Networks operates a fleet-management SaaS platform "
        "serving 4,200 customers across 23 countries. CEO: Larissa Verma. "
        "CFO: Hideo Tanaka."
    )

    _heading(doc, "1. Active Accounts Trend", level=1)
    doc.add_paragraph(
        "Active accounts (paid logos) grew from 82,400 in July 2025 to "
        "89,700 in March 2026 — a 9-month sequential growth of 8.9%. "
        "The chart below shows the trajectory."
    )
    doc.add_picture(io.BytesIO(_chart_kpi_trend()), width=Inches(5.8))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Figure 1 — Monthly active accounts. Jul 2025: 82.4K. "
              "Mar 2026: 89.7K.").italic = True

    _heading(doc, "2. ARR Composition", level=1)
    doc.add_paragraph(
        "FY2026 ARR totals $248M, up 18.4% year-on-year. The split by "
        "product line is shown in Figure 2."
    )
    doc.add_picture(io.BytesIO(_chart_arr_breakdown()), width=Inches(4.2))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Figure 2 — Atlas Core remains the largest line at 42% of "
              "ARR. Professional Services contributes the smallest share at 7%."
              ).italic = True

    _heading(doc, "3. Anchor Numbers", level=1)
    doc.add_paragraph(
        "Active accounts: 89,700 (March 2026). Total ARR: $248M. "
        "ARR YoY growth: +18.4%. Atlas Core ARR share: 42%. "
        "Customer count: 4,200 across 23 countries."
    ).italic = True

    out = OUT_DIR / "03_embedded_charts_kpis.docx"
    doc.save(out)
    print(f"Wrote {out}")


# ============================================================================
# 04 — Tracked changes contract
# ============================================================================

def fixture_tracked_changes_contract():
    """A simulated contract draft that has 'tracked changes' embedded
    using python-docx's revision XML. Real Word would render these as
    insertions/deletions. The competency question is: does the AI Hub
    extractor see the CURRENT state (insertions accepted, deletions
    removed) or BOTH states (mashed together)?"""

    doc = Document()
    _set_default_font(doc)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("MASTER SERVICES AGREEMENT — Draft, Round 3")
    r.font.size = Pt(18)
    r.font.bold = True

    doc.add_paragraph(
        "Between Cyrene Industries (the \"Client\") and Westbrook & Vale "
        "LLP (the \"Provider\"). Effective date: June 1, 2026."
    )

    _heading(doc, "1. Term and Termination", level=1)
    doc.add_paragraph(
        "The initial term of this Agreement is twenty-four (24) months, "
        "commencing on the Effective Date. Either party may terminate "
        "for convenience with sixty (60) days' written notice."
    )

    _heading(doc, "2. Fees and Payment", level=1)
    doc.add_paragraph(
        "Client shall pay Provider a monthly fee of $48,500 USD. "
        "Invoices are due net-30. Late payments accrue interest at 1.5% per month."
    )

    _heading(doc, "3. Confidentiality", level=1)
    doc.add_paragraph(
        "Both parties agree to keep Confidential Information secret for "
        "a period of five (5) years following termination. Permitted "
        "disclosures include legal compulsion and disclosure to "
        "professional advisors under similar confidentiality obligations."
    )

    _heading(doc, "4. Limitation of Liability", level=1)
    doc.add_paragraph(
        "Provider's aggregate liability under this Agreement shall not "
        "exceed the fees paid by Client in the twelve (12) months "
        "preceding the claim, except for liability arising from gross "
        "negligence, wilful misconduct, or breach of confidentiality."
    )

    # Insert a "tracked change" — python-docx doesn't directly support
    # revisions, so we inject the XML manually so the document genuinely
    # contains <w:ins> and <w:del> elements that competent extractors
    # should ignore (rejecting deletions; accepting insertions is the
    # "current state" view).

    # Add a paragraph then inject ins/del runs into its XML.
    p = doc.add_paragraph()
    p_xml = p._p

    def _run_xml(text, color="000000"):
        return (
            f'<w:r xmlns:w="http://schemas.openxmlformats.org/'
            f'wordprocessingml/2006/main">'
            f'<w:rPr><w:color w:val="{color}"/></w:rPr>'
            f'<w:t xml:space="preserve">{text}</w:t></w:r>'
        )

    # Composite: "The penalty for late delivery is [DELETED $5,000] [INSERTED $12,500] per day."
    ins_xml = (
        '<w:ins xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'w:id="1" w:author="Cyrene Legal" w:date="2026-05-12T10:00:00Z">'
        + _run_xml("$12,500 USD")
        + '</w:ins>'
    )
    del_xml = (
        '<w:del xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'w:id="2" w:author="Westbrook Legal" w:date="2026-05-10T09:00:00Z">'
        '<w:r><w:rPr><w:color w:val="000000"/></w:rPr>'
        '<w:delText xml:space="preserve">$5,000 USD</w:delText>'
        '</w:r></w:del>'
    )

    from lxml import etree
    intro = etree.fromstring(_run_xml("Section 5. Service Levels. The penalty for late delivery is "))
    closing = etree.fromstring(_run_xml(" per business day, capped at thirty (30) days."))
    p_xml.append(intro)
    p_xml.append(etree.fromstring(del_xml))
    p_xml.append(etree.fromstring(ins_xml))
    p_xml.append(closing)

    _heading(doc, "6. Anchor — Current state (post-revisions)", level=1)
    doc.add_paragraph(
        "ANCHOR: After accepting all tracked changes, Section 5 reads: "
        "'The penalty for late delivery is $12,500 USD per business day, "
        "capped at thirty (30) days.' The previous round had $5,000 "
        "(removed by tracked deletion)."
    ).italic = True

    out = OUT_DIR / "04_tracked_changes_contract.docx"
    doc.save(out)
    print(f"Wrote {out}")


# ============================================================================
# 05 — Long doc with footnotes + TOC
# ============================================================================

def _add_footnote_ref(paragraph, note_id):
    """Append a footnote reference marker to a paragraph. Word renders
    these as superscript numbers."""
    r = paragraph.add_run()
    r.font.superscript = True
    r.add_text(f"[fn{note_id}]")


def fixture_long_doc_footnotes():
    doc = Document()
    _set_default_font(doc)

    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Hyperion Logistics Network — Technical Reference (v4.2)")
    r.font.size = Pt(20)
    r.font.bold = True

    doc.add_paragraph(
        "April 2026 release. This reference describes the architecture, "
        "operation and governance of the Hyperion Logistics Network "
        "(HLN). Maintainer: Platform Engineering, lead: Mei-Ling Park."
    )

    # Simulated TOC (real TOCs require Word to update; we just stub
    # plain text so the doc looks like a long reference)
    _heading(doc, "Table of Contents", level=1)
    toc_items = [
        ("1. Introduction",                   3),
        ("2. Network Topology",               5),
        ("3. Routing Subsystem",              9),
        ("4. Identity & Access",             14),
        ("5. Data Plane",                    18),
        ("6. Control Plane",                 22),
        ("7. Observability",                 26),
        ("8. Disaster Recovery",             30),
        ("9. Glossary",                      33),
        ("10. Anchor Fingerprint Facts",     35),
    ]
    for title, page in toc_items:
        p = doc.add_paragraph()
        p.add_run(f"{title} ........ {page}")
    doc.add_page_break()

    # Pad with many sections so the file is genuinely "long"
    sections = [
        ("1. Introduction",
         "Hyperion Logistics Network is a multi-region, multi-tenant "
         "fleet routing platform deployed across AWS and GCP. It is "
         "designed for high-throughput dispatch (peak 28,000 events/sec) "
         "with sub-second routing decisions."),
        ("2. Network Topology",
         "HLN is built around 3 primary regions: us-east-1, eu-central-1, "
         "and ap-southeast-2. Each region has 4 availability zones with "
         "asymmetric capacity. Cross-region traffic uses dedicated peering."),
        ("3. Routing Subsystem",
         "The routing layer is implemented as a Rust service named "
         "router-core, deployed as a stateless Kubernetes Deployment "
         "with HPA min=12 pods/region. Latency p95 = 84 ms end-to-end."),
        ("4. Identity & Access",
         "Authentication uses OIDC against Okta. Authorization is policy-as-code "
         "in OPA/Rego, distributed via a sidecar pattern. Service-to-service "
         "uses SPIFFE/SPIRE-issued mTLS certificates rotated every 12 hours."),
        ("5. Data Plane",
         "The data plane consists of Apache Kafka clusters (6 brokers/region) "
         "and PostgreSQL 16 with Patroni HA. Event retention is 7 days. "
         "Total storage footprint as of April 2026 is 184 TB."),
        ("6. Control Plane",
         "The control plane is a Go service named hln-ctrl deployed in a "
         "single region (us-east-1) with hot standby in eu-central-1. "
         "All configuration changes flow through a GitOps pipeline gated "
         "by an automated policy reviewer."),
        ("7. Observability",
         "Logs and metrics are aggregated in Grafana Cloud. Trace sampling "
         "is 100% for error paths and 5% baseline for happy paths. "
         "Active SLO burn-rate alerts: 14."),
        ("8. Disaster Recovery",
         "Recovery Point Objective (RPO) is 60 seconds for the data plane "
         "and 5 minutes for control plane state. Recovery Time Objective "
         "(RTO) is 30 minutes for full regional failover. Last successful "
         "DR drill: February 14, 2026."),
        ("9. Glossary",
         "HLN: Hyperion Logistics Network. SPIFFE: Secure Production "
         "Identity Framework for Everyone. SPIRE: SPIFFE Runtime "
         "Environment. OPA: Open Policy Agent."),
    ]
    for hdr, body in sections:
        _heading(doc, hdr, level=1)
        # Make each section span multiple paragraphs to inflate length
        doc.add_paragraph(body)
        # Add a footnote-like reference in one section
        if hdr.startswith("3."):
            p = doc.add_paragraph(
                "The routing subsystem follows the design described in "
                "the internal RFC 'RFC-HLN-031'"
            )
            _add_footnote_ref(p, 1)
            p.add_run(". This RFC was authored by Mei-Ling Park and "
                      "approved on October 18, 2025.")
        if hdr.startswith("8."):
            p = doc.add_paragraph(
                "Failover is automated via Argo Workflows and triggers from "
                "the SLO burn-rate alarms documented in section 7"
            )
            _add_footnote_ref(p, 2)
            p.add_run(".")
        # Page break between sections to make the doc legitimately long
        doc.add_page_break()

    # Footnotes section (simulated)
    _heading(doc, "Footnotes", level=1)
    doc.add_paragraph(
        "[fn1] RFC-HLN-031: 'Stateless Multi-Region Router' — published "
        "October 18, 2025 by Mei-Ling Park. Approved by Platform "
        "Engineering Council on October 25, 2025."
    )
    doc.add_paragraph(
        "[fn2] SLO burn-rate alarms are defined in the file "
        "hln/slo/burn_rate.yaml in the hln-config repository. "
        "Reviewed quarterly; next review: July 2026."
    )

    # Anchor facts (so the test can reliably ground-truth)
    _heading(doc, "10. Anchor Fingerprint Facts", level=1)
    anchors = [
        "Hyperion Logistics Network (HLN) peak throughput: 28,000 events/sec.",
        "HLN deploys to 3 regions: us-east-1, eu-central-1, ap-southeast-2.",
        "router-core service is written in Rust, p95 = 84 ms.",
        "Data plane storage as of April 2026: 184 TB.",
        "Last successful DR drill: February 14, 2026. RPO 60 s. RTO 30 min.",
        "RFC-HLN-031 was authored by Mei-Ling Park, approved October 18, 2025.",
        "Number of active SLO burn-rate alerts: 14.",
    ]
    for a in anchors:
        doc.add_paragraph(a, style="List Bullet")

    out = OUT_DIR / "05_long_doc_toc_footnotes.docx"
    doc.save(out)
    print(f"Wrote {out}")


# ============================================================================

if __name__ == "__main__":
    fixture_clean_handbook()
    fixture_tables_heavy_report()
    fixture_embedded_charts_kpis()
    fixture_tracked_changes_contract()
    fixture_long_doc_footnotes()
    print("All Word competency fixtures generated.")
