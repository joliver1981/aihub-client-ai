"""Format renderers for the pack-23 corpus.

Every document is authored ONCE as a list of page strings, then rendered into one of the
formats `config.DOC_ALLOWED_EXTENSIONS` accepts. Format is an attribute of a document, not
a separate tier -- the point is that the same content must survive every ingest path.

  pdf    native text PDF (reportlab)          -- the common case
  docx   Word (python-docx)                   -- explicit page breaks
  xlsx   Excel (openpyxl)                     -- tabular documents only
  csv    comma-separated                      -- tabular documents only
  txt    form-feed separated pages            -- cheapest ingest
  scan   rasterised image-PDF, 200dpi, skewed -- the OCR path, untested until now
  jpg    single photographed page             -- the OCR path, image upload

Renderers return the ACTUAL physical page count so ground truth records what was written,
not what was intended.
"""
import io
import os

from PIL import Image, ImageFilter
from reportlab import rl_config
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas as rl_canvas

# Without this, reportlab stamps a wall-clock /CreationDate into every PDF and the corpus is
# byte-different on every run even though the content is identical. The corpus lives outside
# git and is regenerated on demand, so reproducibility is the whole guarantee -- an operator
# has to be able to prove the files they are testing are the files ground truth describes.
rl_config.invariant = 1

PAGE_W, PAGE_H = letter
MARGIN = 72.0
FONT, FONT_SIZE, LEADING = "Helvetica", 9.5, 13.0
WRAP_COLS = 88
LINES_PER_PAGE = int((PAGE_H - 2 * MARGIN) / LEADING)  # 49


def _wrap(text):
    """Wrap a logical page into display lines, preserving blank lines and indentation."""
    out = []
    for raw in text.split("\n"):
        if not raw.strip():
            out.append("")
            continue
        indent = len(raw) - len(raw.lstrip())
        pad = " " * indent
        words, line = raw.split(), ""
        for w in words:
            trial = w if not line else line + " " + w
            if len(trial) + indent > WRAP_COLS and line:
                out.append(pad + line)
                line = w
            else:
                line = trial
        out.append(pad + line)
    return out


def _paginate(pages):
    """Logical pages -> physical pages, splitting anything that overflows."""
    physical = []
    for p in pages:
        lines = _wrap(p)
        while lines:
            physical.append(lines[:LINES_PER_PAGE])
            lines = lines[LINES_PER_PAGE:]
    return physical


def render_pdf(path, pages, title=""):
    physical = _paginate(pages)
    c = rl_canvas.Canvas(path, pagesize=letter)
    c.setTitle(title or os.path.basename(path))
    for i, lines in enumerate(physical):
        c.setFont(FONT, FONT_SIZE)
        y = PAGE_H - MARGIN
        for ln in lines:
            c.drawString(MARGIN, y, ln)
            y -= LEADING
        c.setFont(FONT, 7.5)
        c.drawRightString(PAGE_W - MARGIN, MARGIN * 0.55, f"Page {i + 1} of {len(physical)}")
        c.showPage()
    c.save()
    return len(physical)


def render_txt(path, pages, title=""):
    physical = _paginate(pages)
    with open(path, "w", encoding="utf-8") as fh:
        fh.write("\f".join("\n".join(lines) for lines in physical))
    return len(physical)


def render_docx(path, pages, title=""):
    from docx import Document
    from docx.enum.text import WD_BREAK
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    physical = _paginate(pages)
    for i, lines in enumerate(physical):
        for ln in lines:
            para = doc.add_paragraph(ln)
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.line_spacing = 1.0
        if i < len(physical) - 1:
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    doc.save(path)
    return len(physical)


def render_xlsx(path, rows, sheet_title="Sheet1", header_rows=1):
    """`rows` is a list of lists. Tabular documents only."""
    from openpyxl import Workbook
    from openpyxl.styles import Font

    wb = Workbook()
    ws = wb.active
    ws.title = sheet_title[:31]
    for r in rows:
        ws.append(list(r))
    for r in range(1, header_rows + 1):
        for cell in ws[r]:
            cell.font = Font(bold=True)
    for col in ws.columns:
        width = max((len(str(c.value)) for c in col if c.value is not None), default=8)
        ws.column_dimensions[col[0].column_letter].width = min(max(width + 2, 10), 55)
    wb.save(path)
    # Excel ingest counts one "page" per ~50 rows in the doc engine's chunker; report the
    # sheet as the row-derived estimate so page budgets stay honest.
    return max(1, (len(rows) + 49) // 50)


def render_csv(path, rows):
    import csv as _csv
    with open(path, "w", newline="", encoding="utf-8") as fh:
        _csv.writer(fh).writerows(rows)
    return max(1, (len(rows) + 49) // 50)


def _rasterise(pages, dpi, skew_seq, title=""):
    """Render to PDF in memory, then rasterise each page to a scanned-looking PIL image."""
    import fitz

    buf = io.BytesIO()
    physical = _paginate(pages)
    c = rl_canvas.Canvas(buf, pagesize=letter)
    c.setTitle(title)
    for lines in physical:
        c.setFont(FONT, FONT_SIZE)
        y = PAGE_H - MARGIN
        for ln in lines:
            c.drawString(MARGIN, y, ln)
            y -= LEADING
        c.showPage()
    c.save()
    buf.seek(0)

    doc = fitz.open(stream=buf.read(), filetype="pdf")
    images = []
    for i, page in enumerate(doc):
        pix = page.get_pixmap(dpi=dpi)
        img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
        angle = skew_seq[i % len(skew_seq)]
        img = img.rotate(angle, resample=Image.BICUBIC, fillcolor=(255, 255, 255), expand=False)
        img = img.filter(ImageFilter.GaussianBlur(radius=0.4))
        img = img.convert("L").point(lambda v: min(255, int(v * 0.94 + 12))).convert("RGB")
        images.append(img)
    doc.close()
    return images


def render_scan(path, pages, title="", dpi=200, skew_seq=(0.6, -0.4, 0.35, -0.7, 0.5)):
    """Scanned-document PDF: rasterised, slightly skewed, softened. Forces the OCR path."""
    images = _rasterise(pages, dpi, list(skew_seq), title)
    images[0].save(path, "PDF", resolution=float(dpi), save_all=True, append_images=images[1:])
    return len(images)


def render_jpg(path, pages, title="", dpi=150, skew=1.1):
    """A single photographed page -- e.g. a signature page sent from someone's phone."""
    images = _rasterise(pages[:1], dpi, [skew], title)
    images[0].save(path, "JPEG", quality=72)
    return 1


TEXT_RENDERERS = {
    "pdf": render_pdf,
    "txt": render_txt,
    "docx": render_docx,
    "scan": render_scan,   # written with a .pdf extension
    "jpg": render_jpg,
}

EXT_FOR_FORMAT = {"pdf": "pdf", "txt": "txt", "docx": "docx", "scan": "pdf",
                  "jpg": "jpg", "xlsx": "xlsx", "csv": "csv"}
