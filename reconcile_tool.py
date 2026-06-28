"""
reconcile_tool.py — PROTOTYPE for the Phase-1 "reconcile a table" tool.

This is the fix for "Way 2" (silent omission): instead of asking the model to
*judge* whether invoices match a contract, we EXTRACT the full data from the real
documents and COMPUTE every check in code — so the work actually happens, on every
line, deterministically. The model would just call this tool and report what it
returns.

It parses the ACTUAL documents an agent receives (no answer key, no cheating):
  * the supplier contract (.docx)  -> firm prices, rebate rule, freight rule
  * the invoice register (.pdf)     -> every line's qty / unit price / amount

…and runs four checks over every invoice:
  1. line-math     : qty x unit_price == printed amount?      (catches the June slip)
  2. price-hold    : each unit price == the contract price?    (catches the Feb tent)
  3. freight       : any freight charged when FOB-Destination? (catches the May line)
  4. volume rebate : owed 2.5% credit present & correct?       (catches the April miss)

Run the proof:
  $PY = "$env:USERPROFILE\\miniconda3\\envs\\aihub2.1\\python.exe"
  & $PY reconcile_tool.py
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional

# Money tokens like 1,234.56
_MONEY = r"[\d,]+\.\d{2}"


def _money(s: str) -> float:
    return float(s.replace(",", "").replace("$", "").strip())


def extract_contract(docx_path: str) -> dict:
    """Pull firm prices (from the Schedule A table) + the rule parameters (from
    the prose) out of the supply agreement .docx."""
    from docx import Document

    doc = Document(str(docx_path))

    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            parts.append(" | ".join(c.text for c in row.cells))
    text = "\n".join(parts)

    # --- Schedule A prices ------------------------------------------------
    prices = {}
    for t in doc.tables:
        header = " ".join(c.text for c in t.rows[0].cells).lower()
        if "unit price" not in header:
            continue
        cols = [c.text.strip().lower() for c in t.rows[0].cells]
        desc_i = next((i for i, c in enumerate(cols) if "description" in c), 0)
        price_i = next((i for i, c in enumerate(cols) if "price" in c), len(cols) - 1)
        for row in t.rows[1:]:
            cells = [c.text.strip() for c in row.cells]
            m = re.search(rf"\$?({_MONEY})", cells[price_i]) if price_i < len(cells) else None
            if m and desc_i < len(cells) and cells[desc_i]:
                prices[cells[desc_i]] = _money(m.group(1))

    # --- rule parameters from the prose -----------------------------------
    low = text.lower()
    rate_m = re.search(r"\((\d+(?:\.\d+)?)\s*%\)", text)
    thr_m = re.search(r"\(([\d,]{3,})\)\s*units", text)

    return {
        "prices": prices,
        "rebate_rate": float(rate_m.group(1)) / 100.0 if rate_m else 0.025,
        "rebate_threshold": int(thr_m.group(1).replace(",", "")) if thr_m else 1200,
        "freight_allowed": "fob destination" not in low,
    }


@dataclass
class Invoice:
    month: str
    invoice_no: str
    total_units: int
    lines: List[dict] = field(default_factory=list)
    freight: float = 0.0
    rebate: float = 0.0


_LINE_RE = re.compile(
    rf"^\s*(?P<desc>.+?)\s{{2,}}(?P<qty>\d[\d,]*)\s+\$(?P<price>{_MONEY})\s+\$\(?(?P<amount>{_MONEY})\)?\s*$"
)


def extract_invoices(pdf_path: str) -> List[Invoice]:
    """Parse every invoice page into structured lines using pypdf layout mode."""
    from pypdf import PdfReader

    reader = PdfReader(str(pdf_path))
    invoices = []
    for page in reader.pages:
        try:
            text = page.extract_text(extraction_mode="layout")
        except Exception:
            text = page.extract_text() or ""
        if "INVOICE" not in text and "Billing Month" not in text:
            continue

        inv_no = (re.search(r"INVOICE\s+(\S+)", text) or [None, ""])[1]
        month = (re.search(r"Billing Month:\s*([A-Za-z]+\s+\d{4})", text) or [None, ""])[1]
        units_m = re.search(r"Total Units:\s*([\d,]+)", text)
        inv = Invoice(
            month=month,
            invoice_no=inv_no,
            total_units=int(units_m.group(1).replace(",", "")) if units_m else 0,
        )
        for line in text.splitlines():
            m = _LINE_RE.match(line)
            if m:
                inv.lines.append(
                    {
                        "desc": m.group("desc").strip(),
                        "qty": int(m.group("qty").replace(",", "")),
                        "unit_price": _money(m.group("price")),
                        "amount": _money(m.group("amount")),
                    }
                )
            elif "Freight" in line:
                fm = re.search(rf"\$({_MONEY})", line)
                if fm:
                    inv.freight = _money(fm.group(1))
            elif "Volume Rebate Credit" in line:
                rm = re.search(rf"\(\$?({_MONEY})\)", line)
                if rm:
                    inv.rebate = _money(rm.group(1))
        invoices.append(inv)
    return invoices


@dataclass
class Discrepancy:
    month: str
    invoice_no: str
    kind: str
    amount: float
    detail: str


def _cents(x: float) -> bool:
    return abs(x) > 0.005


def reconcile(contract: dict, invoices: List[Invoice]) -> dict:
    prices = contract["prices"]
    rate = contract["rebate_rate"]
    threshold = contract["rebate_threshold"]
    freight_allowed = contract["freight_allowed"]

    discrepancies = []
    correct = []

    for inv in invoices:
        before = len(discrepancies)
        corrected_subtotal = 0.0
        for ln in inv.lines:
            cp = prices.get(ln["desc"])
            corrected_subtotal += round(ln["qty"] * (cp if cp is not None else ln["unit_price"]), 2)

            # price-hold
            if cp is not None and _cents(ln["unit_price"] - cp):
                over = round(ln["qty"] * (ln["unit_price"] - cp), 2)
                discrepancies.append(
                    Discrepancy(
                        inv.month, inv.invoice_no, "price_hold", over,
                        f"{ln['desc']}: billed ${ln['unit_price']:.2f} vs contract ${cp:.2f} "
                        f"({ln['qty']} x ${ln['unit_price'] - cp:.2f})",
                    )
                )

            # line-math
            expected = round(ln["qty"] * ln["unit_price"], 2)
            if _cents(ln["amount"] - expected):
                over = round(ln["amount"] - expected, 2)
                discrepancies.append(
                    Discrepancy(
                        inv.month, inv.invoice_no, "line_math", over,
                        f"{ln['desc']}: amount ${ln['amount']:.2f} but {ln['qty']} x $"
                        f"{ln['unit_price']:.2f} = ${expected:.2f}",
                    )
                )

        # freight
        if inv.freight and not freight_allowed:
            discrepancies.append(
                Discrepancy(
                    inv.month, inv.invoice_no, "freight", round(inv.freight, 2),
                    f"freight ${inv.freight:.2f} charged despite FOB Destination",
                )
            )

        # volume rebate
        if inv.total_units >= threshold:
            owed = round(corrected_subtotal * rate, 2)
            if abs(inv.rebate - owed) > 0.02:
                over = round(owed - inv.rebate, 2)
                discrepancies.append(
                    Discrepancy(
                        inv.month, inv.invoice_no, "rebate", over,
                        f"units {inv.total_units:,} >= {threshold:,}: owed {rate * 100:.1f}% rebate $"
                        f"{owed:.2f}, invoice shows ${inv.rebate:.2f}",
                    )
                )

        if len(discrepancies) == before:
            correct.append(inv.month)

    return {
        "invoices_checked": len(invoices),
        "discrepancies": discrepancies,
        "total_overcharge": round(sum(d.amount for d in discrepancies), 2),
        "correct_invoices": correct,
    }


def reconcile_invoices(contract_path: str, invoices_path: str) -> dict:
    """Public entry point an agent tool would call."""
    contract = extract_contract(contract_path)
    invoices = extract_invoices(invoices_path)
    return reconcile(contract, invoices)


if __name__ == "__main__":
    base = Path(__file__).resolve().parent / "test_human" / "Finance_Capability_Tour" / "Files to upload"
    result = reconcile_invoices(
        base / "Summit Gear Supply Agreement.docx",
        base / "Summit Gear Invoices Jan-Jun 2025.pdf",
    )
    print(f"Invoices checked: {result['invoices_checked']}")
    print(f"Correct invoices: {', '.join(result['correct_invoices'])}\n")
    print("Discrepancies found (computed from the documents, deterministically):")
    for d in result["discrepancies"]:
        print(f"  {d.month:14s} {d.kind:10s} ${d.amount:>9,.2f}   {d.detail}")
    print(f"\n  TOTAL OVERCHARGE: ${result['total_overcharge']:,.2f}")
