"""Generate additional knowledge-test fixtures of DIFFERENT file types
(Excel + Word-with-embedded-charts) to exercise other upload code paths
beyond plain .docx.

Output:
  04_aurora_quarterly_financials_q1_2026.xlsx   — multi-sheet Excel
  05_zenith_v3_architecture_with_diagrams.docx  — Word doc with embedded
                                                  matplotlib chart PNGs
"""
import io
import os
from pathlib import Path

import matplotlib
matplotlib.use("Agg")  # headless backend
import matplotlib.pyplot as plt
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL


OUT_DIR = Path(__file__).resolve().parent


# ════════════════════════════════════════════════════════════════════════
# DOCUMENT 4: Aurora Bioplastics Q1 2026 Financial Report (.xlsx)
# ════════════════════════════════════════════════════════════════════════

def build_excel():
    """Multi-sheet financial workbook with formulas, formatting, and unique
    facts."""
    wb = openpyxl.Workbook()

    # Styling helpers
    hdr_font = Font(name="Calibri", size=12, bold=True, color="FFFFFF")
    hdr_fill = PatternFill("solid", fgColor="1F3864")
    subtotal_fill = PatternFill("solid", fgColor="D9E1F2")
    money_fmt = '"$"#,##0.00'
    int_fmt = "#,##0"
    pct_fmt = "0.0%"
    border = Border(*[Side(style="thin", color="BFBFBF")] * 4)

    def header_row(ws, row, headers):
        for col, val in enumerate(headers, 1):
            cell = ws.cell(row=row, column=col, value=val)
            cell.font = hdr_font
            cell.fill = hdr_fill
            cell.alignment = Alignment(horizontal="center", vertical="center")
            cell.border = border
        ws.row_dimensions[row].height = 22

    def style_data_row(ws, row, ncols, fmt=None):
        for col in range(1, ncols + 1):
            cell = ws.cell(row=row, column=col)
            cell.border = border
            if fmt and col > 1:
                cell.number_format = fmt

    def auto_width(ws, widths):
        for col, w in enumerate(widths, 1):
            ws.column_dimensions[get_column_letter(col)].width = w

    # ── Sheet 1: Profile + Highlights ─────────────────────────────────
    ws = wb.active
    ws.title = "Company Profile"

    ws["A1"] = "AURORA BIOPLASTICS, INC."
    ws["A1"].font = Font(name="Calibri", size=20, bold=True, color="1F3864")
    ws.merge_cells("A1:D1")

    ws["A3"] = "Quarterly Financial Report — Q1 2026"
    ws["A3"].font = Font(name="Calibri", size=14, italic=True)
    ws.merge_cells("A3:D3")

    profile = [
        ("Company name",      "Aurora Bioplastics, Inc."),
        ("Headquarters",      "Munich, Germany"),
        ("Founded",           "April 2014"),
        ("CEO",               "Dr. Hanna Engström"),
        ("CFO",               "Theo Brandt"),
        ("Ticker symbol",     "ABPL (Frankfurt Stock Exchange)"),
        ("Fiscal year start", "January 1"),
        ("Total employees",   312),
        ("Manufacturing sites", 4),
        ("Reporting currency","EUR (figures in USD-equivalent)"),
        ("Report date",       "April 28, 2026"),
        ("Auditor",           "Steinmetz & Partners GmbH"),
    ]
    for i, (k, v) in enumerate(profile, 5):
        ws.cell(row=i, column=1, value=k).font = Font(bold=True)
        ws.cell(row=i, column=2, value=v)
    auto_width(ws, [24, 38])

    # ── Sheet 2: P&L Statement ────────────────────────────────────────
    ws = wb.create_sheet("P&L Statement")
    ws["A1"] = "Profit & Loss Statement — Q1 2026 (USD)"
    ws["A1"].font = Font(size=14, bold=True, color="1F3864")
    ws.merge_cells("A1:E1")
    header_row(ws, 3, ["Line item", "Jan 2026", "Feb 2026", "Mar 2026", "Q1 Total"])

    revenue_rows = [
        ("Bioplastic resin sales",    4_120_000.00, 4_405_000.00, 4_835_500.00),
        ("Compostable film sales",      820_000.00,   875_000.00,   915_000.00),
        ("Engineering services",        285_000.00,   295_000.00,   310_000.00),
        ("Licensing revenue",            85_000.00,    87_500.00,    87_500.00),
    ]
    row = 4
    for name, j, f, m in revenue_rows:
        ws.cell(row=row, column=1, value=name)
        ws.cell(row=row, column=2, value=j)
        ws.cell(row=row, column=3, value=f)
        ws.cell(row=row, column=4, value=m)
        ws.cell(row=row, column=5, value=f"=SUM(B{row}:D{row})")
        style_data_row(ws, row, 5, money_fmt)
        row += 1
    # Revenue subtotal
    ws.cell(row=row, column=1, value="Total Revenue").font = Font(bold=True)
    for c in range(2, 6):
        ws.cell(row=row, column=c,
                value=f"=SUM({get_column_letter(c)}4:{get_column_letter(c)}{row-1})")
        ws.cell(row=row, column=c).font = Font(bold=True)
        ws.cell(row=row, column=c).fill = subtotal_fill
        ws.cell(row=row, column=c).number_format = money_fmt
        ws.cell(row=row, column=c).border = border
    ws.cell(row=row, column=1).fill = subtotal_fill
    ws.cell(row=row, column=1).border = border
    revenue_total_row = row
    row += 2

    cost_rows = [
        ("Raw materials (bio-feedstock)", -1_840_000.00, -1_960_000.00, -2_140_500.00),
        ("Manufacturing labor",            -730_000.00,   -745_000.00,   -780_000.00),
        ("Utilities (Munich + 3 plants)",  -210_000.00,   -218_500.00,   -225_000.00),
        ("Sales & marketing",              -385_000.00,   -395_000.00,   -402_000.00),
        ("R&D",                            -545_000.00,   -560_000.00,   -595_000.00),
        ("G&A",                            -385_000.00,   -390_000.00,   -398_000.00),
        ("Depreciation & amortization",    -185_000.00,   -185_000.00,   -185_000.00),
    ]
    for name, j, f, m in cost_rows:
        ws.cell(row=row, column=1, value=name)
        ws.cell(row=row, column=2, value=j)
        ws.cell(row=row, column=3, value=f)
        ws.cell(row=row, column=4, value=m)
        ws.cell(row=row, column=5, value=f"=SUM(B{row}:D{row})")
        style_data_row(ws, row, 5, money_fmt)
        row += 1
    # Net income (revenue + sum of negative costs)
    ws.cell(row=row, column=1, value="Net Income").font = Font(bold=True, color="1F3864")
    for c in range(2, 6):
        # Net = Total revenue + sum of cost rows
        cell = ws.cell(row=row, column=c)
        cost_start = revenue_total_row + 2
        cost_end = row - 1
        cell.value = (
            f"={get_column_letter(c)}{revenue_total_row}"
            f"+SUM({get_column_letter(c)}{cost_start}:{get_column_letter(c)}{cost_end})"
        )
        cell.font = Font(bold=True, color="1F3864")
        cell.fill = subtotal_fill
        cell.number_format = money_fmt
        cell.border = border
    ws.cell(row=row, column=1).fill = subtotal_fill
    ws.cell(row=row, column=1).border = border
    # Sanity-anchor the fingerprinted total in a non-formula cell so the
    # test can ground-truth against the values the SUM formulas produce.
    # Q1 revenue = 5,310,000 + 5,662,500 + 6,148,000 = $17,120,500
    # Q1 net income = revenue + sum(costs) = $3,661,500
    ws.cell(row=row + 3, column=1,
            value="Anchor — confirmed Q1 totals (for audit cross-check):").font = Font(italic=True)
    ws.cell(row=row + 4, column=1, value="Q1 2026 Total Revenue").font = Font(bold=True)
    ws.cell(row=row + 4, column=2, value=17_120_500.00).number_format = money_fmt
    ws.cell(row=row + 5, column=1, value="Q1 2026 Net Income").font = Font(bold=True)
    ws.cell(row=row + 5, column=2, value=3_661_500.00).number_format = money_fmt
    auto_width(ws, [34, 16, 16, 16, 16])

    # ── Sheet 3: Department Headcount ─────────────────────────────────
    ws = wb.create_sheet("Headcount by Dept")
    ws["A1"] = "Headcount by Department — As of March 31, 2026"
    ws["A1"].font = Font(size=14, bold=True, color="1F3864")
    ws.merge_cells("A1:D1")
    header_row(ws, 3, ["Department", "Munich HQ", "Other Sites", "Total"])
    rows = [
        ("Manufacturing", 84, 64,  148),
        ("R&D",            42,  6,   48),
        ("Engineering",    35,  4,   39),
        ("Sales & Marketing", 24, 5,  29),
        ("G&A (Finance, HR, Legal)", 18, 2, 20),
        ("Quality & Compliance", 12, 4, 16),
        ("Customer Success", 8, 4,  12),
    ]
    r = 4
    for dept, hq, other, total in rows:
        ws.cell(row=r, column=1, value=dept)
        ws.cell(row=r, column=2, value=hq)
        ws.cell(row=r, column=3, value=other)
        ws.cell(row=r, column=4, value=f"=SUM(B{r}:C{r})")
        style_data_row(ws, r, 4, int_fmt)
        r += 1
    ws.cell(row=r, column=1, value="Company total").font = Font(bold=True)
    for c in range(2, 5):
        ws.cell(row=r, column=c, value=f"=SUM({get_column_letter(c)}4:{get_column_letter(c)}{r-1})")
        ws.cell(row=r, column=c).font = Font(bold=True)
        ws.cell(row=r, column=c).fill = subtotal_fill
        ws.cell(row=r, column=c).number_format = int_fmt
        ws.cell(row=r, column=c).border = border
    ws.cell(row=r, column=1).fill = subtotal_fill
    ws.cell(row=r, column=1).border = border
    # Anchor
    ws.cell(row=r + 3, column=1,
            value="Total Aurora Bioplastics employees as of March 31, 2026: 312").font = Font(italic=True, bold=True)
    auto_width(ws, [34, 14, 14, 12])

    # ── Sheet 4: Customer Concentration ───────────────────────────────
    ws = wb.create_sheet("Customer Concentration")
    ws["A1"] = "Top 10 Customers by Q1 2026 Revenue Contribution"
    ws["A1"].font = Font(size=14, bold=True, color="1F3864")
    ws.merge_cells("A1:D1")
    header_row(ws, 3, ["Rank", "Customer", "Revenue (USD)", "% of Q1 Revenue"])
    customers = [
        ("Vellichor Industries",       3_260_510, 0.220),
        ("Bramble & Forest Co.",       1_823_322, 0.123),
        ("PolarKraft Packaging GmbH",  1_408_948, 0.095),
        ("Sundial Foods Cooperative",  1_171_820, 0.079),
        ("Tessuto Holdings",             926_281, 0.063),
        ("Greenline Distributors",       785_487, 0.053),
        ("Cypress Container Corp",       637_281, 0.043),
        ("Mira Pharmaceuticals AG",      518_718, 0.035),
        ("Aurora Bioplastics — Internal",385_333, 0.026),
        ("Other (38 customers, none > 2%)", 3_902_800, 0.263),
    ]
    r = 4
    for i, (name, rev, pct) in enumerate(customers, 1):
        ws.cell(row=r, column=1, value=i)
        ws.cell(row=r, column=2, value=name)
        ws.cell(row=r, column=3, value=rev).number_format = money_fmt
        ws.cell(row=r, column=4, value=pct).number_format = pct_fmt
        for c in range(1, 5):
            ws.cell(row=r, column=c).border = border
        r += 1
    ws.cell(row=r + 2, column=1,
            value="Top customer: Vellichor Industries — 22.0% of Q1 2026 revenue").font = Font(italic=True, bold=True)
    auto_width(ws, [6, 38, 16, 16])

    # ── Sheet 5: Cash Flow ────────────────────────────────────────────
    ws = wb.create_sheet("Cash Flow Summary")
    ws["A1"] = "Cash Flow Summary — Q1 2026"
    ws["A1"].font = Font(size=14, bold=True, color="1F3864")
    ws.merge_cells("A1:C1")
    header_row(ws, 3, ["Item", "Amount (USD)", "Notes"])
    cf_rows = [
        ("Opening cash balance (Jan 1, 2026)",    6_948_200, "Audited closing balance Q4 2025"),
        ("Net cash from operating activities",    2_482_300, "Excludes deferred tax timing"),
        ("Net cash used in investing activities", -1_270_500,"Capacity expansion at Plant #3 (Bremen)"),
        ("Net cash used in financing activities", -1_080_000,"Term-loan principal + Q1 dividend"),
        ("Foreign exchange effects",                221_000, "EUR/USD favorable movement"),
        ("Closing cash balance (Mar 31, 2026)",   8_301_000, "Increase of $1.353M from open"),
    ]
    for i, (name, amt, note) in enumerate(cf_rows, 4):
        ws.cell(row=i, column=1, value=name).font = Font(bold=("balance" in name))
        ws.cell(row=i, column=2, value=amt).number_format = money_fmt
        ws.cell(row=i, column=3, value=note)
        for c in range(1, 4):
            ws.cell(row=i, column=c).border = border
    ws.cell(row=10, column=1).fill = subtotal_fill
    ws.cell(row=10, column=2).fill = subtotal_fill
    ws.cell(row=10, column=3).fill = subtotal_fill
    auto_width(ws, [38, 16, 44])

    out = OUT_DIR / "04_aurora_quarterly_financials_q1_2026.xlsx"
    wb.save(out)
    print(f"Wrote {out} ({out.stat().st_size} bytes)")


# ════════════════════════════════════════════════════════════════════════
# DOCUMENT 5: Zenith v3 Architecture (.docx with embedded charts)
# ════════════════════════════════════════════════════════════════════════

def _chart_latency_distribution() -> bytes:
    """Generate a latency-distribution bar chart as PNG bytes."""
    fig, ax = plt.subplots(figsize=(6, 3), dpi=150)
    buckets = ["p50", "p75", "p90", "p95", "p99", "p99.9"]
    values = [47, 78, 132, 180, 290, 615]
    bars = ax.bar(buckets, values, color="#1F3864")
    ax.set_ylabel("Latency (ms)")
    ax.set_title("Zenith v3 — End-to-end Request Latency (production, May 2026)")
    for bar, val in zip(bars, values):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 8,
                str(val), ha="center", fontsize=8)
    ax.grid(axis="y", linestyle=":", alpha=0.4)
    fig.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png")
    plt.close(fig)
    return buf.getvalue()


def _chart_service_breakdown() -> bytes:
    """Pie chart of microservice categories."""
    fig, ax = plt.subplots(figsize=(5, 4), dpi=150)
    labels = ["Edge / API", "Order pipeline", "Identity & RBAC",
              "Reporting & BI", "Notifications", "Internal tooling"]
    sizes = [5, 7, 3, 4, 2, 2]
    colors = ["#1F3864", "#2E74B5", "#5B9BD5", "#9DC3E6", "#BDD7EE", "#DEEBF7"]
    ax.pie(sizes, labels=labels, colors=colors, autopct="%1.0f%%",
           startangle=90, wedgeprops=dict(edgecolor="white"))
    ax.set_title("Zenith v3 — 23 Microservices by Category")
    fig.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png")
    plt.close(fig)
    return buf.getvalue()


def build_zenith_docx():
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ── Title page ────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Zenith v3")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Architecture Reference — Production v3.0")
    run.font.size = Pt(20)
    run.font.bold = True

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run("Document classification: Internal Engineering\n").italic = True
    info.add_run("Owner: Sebastian Vogel, Director of Platform\n").italic = True
    info.add_run("Effective: May 1, 2026\n").italic = True
    info.add_run("Revision: 3.0 (supersedes 2.7)").italic = True
    doc.add_page_break()

    # ── 1. Overview ───────────────────────────────────────────────────
    doc.add_heading("1. System Overview", level=1)
    doc.add_paragraph(
        "Zenith v3 is the production order-processing platform for our European "
        "and North American operations. It replaces Zenith v2 (decommissioned "
        "March 2026) and consolidates 23 microservices behind a single GraphQL "
        "edge layer hosted on AWS region us-east-2."
    )
    doc.add_paragraph(
        "Engineering ownership is shared across 4 teams: Platform (lead: "
        "Sebastian Vogel), Order Pipeline (lead: Akiko Tamura), Identity "
        "(lead: Devon Cole), and Reporting (lead: Mira Chen). The combined "
        "team headcount is 38 engineers as of May 2026."
    )

    doc.add_heading("1.1 Key Facts", level=2)
    facts = doc.add_paragraph()
    facts.add_run("Headline metrics for Zenith v3:\n").bold = True
    bullets = [
        ("23 microservices in production",
         "Up from 17 in v2; split for clearer ownership"),
        ("AWS us-east-2 (Ohio)",
         "Single-region deployment; multi-AZ within the region"),
        ("PostgreSQL 16 + TimescaleDB",
         "Operational store; TimescaleDB extension for time-series telemetry"),
        ("Apache Kafka",
         "Event streaming backbone; 4 brokers, 96 partitions per topic typical"),
        ("End-to-end p50 latency 47 ms",
         "Measured at edge, May 2026"),
        ("p95 latency 180 ms",
         "SLO target: 200 ms"),
        ("99.97 % availability over past 90 days",
         "Above our 99.95 % SLO"),
    ]
    for headline, detail in bullets:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(headline + " — ").bold = True
        p.add_run(detail)

    doc.add_page_break()

    # ── 2. Diagram — Service Breakdown ────────────────────────────────
    doc.add_heading("2. Service Categories", level=1)
    doc.add_paragraph(
        "Of the 23 microservices, the largest category by count is the order "
        "pipeline (7 services), reflecting the complexity of multi-tenant "
        "order routing, fulfillment, and reconciliation. The figure below "
        "breaks down service ownership by functional area."
    )
    doc.add_picture(io.BytesIO(_chart_service_breakdown()), width=Inches(5.5))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Figure 1 — Microservice distribution by category. ").italic = True
    p.add_run("Order pipeline (7) is the largest; reporting (4) is second.").italic = True

    doc.add_page_break()

    # ── 3. Data architecture ──────────────────────────────────────────
    doc.add_heading("3. Data Architecture", level=1)
    doc.add_paragraph(
        "The primary operational data store is PostgreSQL 16 with the "
        "TimescaleDB 2.14 extension. PostgreSQL handles transactional order "
        "data; TimescaleDB hosts hypertables for telemetry, audit, and metric "
        "ingestion at roughly 280,000 rows per minute peak."
    )

    doc.add_heading("3.1 Schema Layout", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Schema"
    hdr[1].text = "Purpose"
    hdr[2].text = "Owning team"
    rows_data = [
        ("orders",     "Order header + line items + status", "Order Pipeline"),
        ("identity",   "Users, roles, tenants, sessions",    "Identity"),
        ("billing",    "Invoices, payment events, refunds",  "Order Pipeline"),
        ("inventory",  "Stock levels, reservations, holds",  "Order Pipeline"),
        ("reporting",  "Materialised views, dashboards",     "Reporting"),
        ("telemetry",  "Time-series TimescaleDB hypertables","Platform"),
        ("audit",      "Append-only event log",              "Platform"),
    ]
    for s, p_, o in rows_data:
        r = table.add_row().cells
        r[0].text = s
        r[1].text = p_
        r[2].text = o

    doc.add_paragraph()
    doc.add_paragraph(
        "All schemas are owned by exactly one team. Cross-schema queries are "
        "permitted only via materialised views in the reporting schema, "
        "which are refreshed every 5 minutes by an Airflow DAG named "
        "'zenith_reporting_refresh_v3'."
    )

    doc.add_page_break()

    # ── 4. Latency chart ──────────────────────────────────────────────
    doc.add_heading("4. Latency Profile", level=1)
    doc.add_paragraph(
        "Measured end-to-end at the GraphQL edge, from request ingestion "
        "through service composition to response. The chart below shows the "
        "full distribution captured during the week of May 5–11, 2026 from "
        "production traffic (12.4 million requests)."
    )
    doc.add_picture(io.BytesIO(_chart_latency_distribution()), width=Inches(5.8))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Figure 2 — Latency distribution. p50 47 ms; p95 180 ms; "
              "p99 290 ms.").italic = True

    doc.add_heading("4.1 SLO Status", level=2)
    doc.add_paragraph(
        "Our service-level objectives are: p50 ≤ 50 ms, p95 ≤ 200 ms, "
        "p99 ≤ 350 ms. All three are met for the reporting period. The next "
        "review is scheduled for July 14, 2026."
    )

    doc.add_page_break()

    # ── 5. Code samples ───────────────────────────────────────────────
    doc.add_heading("5. Reference Code", level=1)
    doc.add_paragraph(
        "Canonical example of how a service consumes order events from Kafka. "
        "Topic is 'orders.committed.v3'; consumer group naming convention is "
        "'<service-name>.orders.committed.v3'."
    )

    code1 = doc.add_paragraph()
    cr = code1.add_run(
        "from kafka import KafkaConsumer\n"
        "import json\n\n"
        "consumer = KafkaConsumer(\n"
        "    'orders.committed.v3',\n"
        "    bootstrap_servers='kafka-prod-01.zenith.local:9092',\n"
        "    group_id='reporting.orders.committed.v3',\n"
        "    enable_auto_commit=False,\n"
        "    auto_offset_reset='earliest',\n"
        "    value_deserializer=lambda b: json.loads(b.decode('utf-8'))\n"
        ")\n\n"
        "for msg in consumer:\n"
        "    handle_order(msg.value)\n"
        "    consumer.commit()  # at-least-once delivery\n"
    )
    cr.font.name = "Consolas"
    cr.font.size = Pt(9)

    doc.add_heading("5.1 GraphQL query example", level=2)
    code2 = doc.add_paragraph()
    cr = code2.add_run(
        "query OrderById($id: ID!) {\n"
        "  order(id: $id) {\n"
        "    id\n"
        "    status\n"
        "    customer { name email tenantId }\n"
        "    lineItems {\n"
        "      sku\n"
        "      quantity\n"
        "      unitPriceCents\n"
        "    }\n"
        "    createdAt\n"
        "    fulfilledAt\n"
        "  }\n"
        "}\n"
    )
    cr.font.name = "Consolas"
    cr.font.size = Pt(9)

    doc.add_page_break()

    # ── 6. Deployment ─────────────────────────────────────────────────
    doc.add_heading("6. Deployment Topology", level=1)
    doc.add_paragraph(
        "Zenith v3 deploys to Amazon EKS in region us-east-2, using two "
        "EKS clusters: 'zenith-prod-blue' and 'zenith-prod-green', behind "
        "an ALB that switches traffic during rolling deploys. Cutover "
        "between blue and green is gated by a synthetic health probe; the "
        "switch completes in under 90 seconds when no canary alarms trip."
    )

    doc.add_heading("6.1 Instance Footprint", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Node group"
    hdr[1].text = "Instance type"
    hdr[2].text = "Count (peak)"
    inst_rows = [
        ("Edge / API",        "m6i.2xlarge", "8"),
        ("Order pipeline",    "c6i.4xlarge", "14"),
        ("Identity & RBAC",   "m6i.xlarge",  "6"),
        ("Reporting & BI",    "r6i.2xlarge", "8"),
        ("Notifications",     "m6i.large",   "4"),
        ("Internal tooling",  "m6i.large",   "4"),
        ("Kafka brokers",     "m6i.4xlarge", "4"),
        ("PostgreSQL primary","r6i.4xlarge", "1"),
        ("PostgreSQL replicas","r6i.4xlarge","2"),
    ]
    for n, t, c in inst_rows:
        r = table.add_row().cells
        r[0].text = n
        r[1].text = t
        r[2].text = c

    doc.add_page_break()

    # ── 7. Security ───────────────────────────────────────────────────
    doc.add_heading("7. Security Posture", level=1)
    doc.add_paragraph(
        "All inter-service communication uses mTLS issued by an internal "
        "Vault PKI engine. Certificate rotation is automated every 14 days, "
        "with a 48-hour overlap window for the previous certificate."
    )
    doc.add_paragraph(
        "External API authentication uses OAuth 2.1 with PKCE; bearer tokens "
        "are JWTs signed by Auth0, validated at the GraphQL edge with a "
        "5-minute key cache."
    )
    doc.add_heading("7.1 Secrets Management", level=2)
    doc.add_paragraph(
        "Application secrets live in HashiCorp Vault, with namespace per "
        "tenant. The Vault leader's address is internal-only: "
        "vault-prod.zenith.local:8200. Production access requires hardware "
        "MFA (YubiKey)."
    )

    doc.add_page_break()

    # ── 8. Observability ──────────────────────────────────────────────
    doc.add_heading("8. Observability", level=1)
    doc.add_paragraph(
        "Logs route to Datadog via the Fluent Bit DaemonSet, partitioned by "
        "service name and severity. Traces use OpenTelemetry with the "
        "Datadog exporter; sample rate is 5% baseline, 100% for any request "
        "tagged with error=true or trace-debug=true."
    )

    doc.add_heading("8.1 Key dashboards", level=2)
    dashes = [
        "ZENITH-OVERVIEW (high-level platform health, ~32 panels)",
        "ZENITH-ORDERS-FLOW (end-to-end order journey)",
        "ZENITH-DB-PERFORMANCE (PostgreSQL + TimescaleDB)",
        "ZENITH-KAFKA-HEALTH (broker / topic / consumer-lag)",
        "ZENITH-SLO-BURN (multi-window SLO burn-rate alarms)",
    ]
    for d in dashes:
        doc.add_paragraph(d, style="List Bullet")

    doc.add_page_break()

    # ── 9. Runbook references ─────────────────────────────────────────
    doc.add_heading("9. On-Call & Runbooks", level=1)
    doc.add_paragraph(
        "Primary on-call rotation is managed in PagerDuty under schedule "
        "'zenith-prod-primary'. Escalation path is: primary → secondary "
        "(7 min) → engineering manager (15 min) → director (Sebastian Vogel, "
        "30 min)."
    )
    doc.add_paragraph(
        "All critical-path runbooks live in the 'zenith' Confluence space "
        "under /Runbooks. Each runbook MUST be reviewed quarterly; the next "
        "review window is July 6–17, 2026."
    )

    doc.add_page_break()

    # ── 10. Roadmap ───────────────────────────────────────────────────
    doc.add_heading("10. Forward Roadmap", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Quarter"
    hdr[1].text = "Initiative"
    hdr[2].text = "Owner"
    rm_rows = [
        ("Q3 2026", "Multi-region active/active (add eu-central-1)", "Sebastian Vogel"),
        ("Q3 2026", "Replace Auth0 with internal Identity v4",       "Devon Cole"),
        ("Q4 2026", "Order pipeline rewrite in Rust (services 12/23)","Akiko Tamura"),
        ("Q4 2026", "TimescaleDB 2.16 upgrade",                      "Platform"),
        ("Q1 2027", "ML-driven order routing (replace static rules)","Mira Chen"),
    ]
    for q, init, owner in rm_rows:
        r = table.add_row().cells
        r[0].text = q
        r[1].text = init
        r[2].text = owner

    doc.add_paragraph()
    doc.add_paragraph(
        "This document supersedes the v2 architecture reference dated "
        "October 2024. The next scheduled revision is November 2026."
    ).italic = True

    out = OUT_DIR / "05_zenith_v3_architecture_with_diagrams.docx"
    doc.save(out)
    print(f"Wrote {out} ({out.stat().st_size} bytes)")


if __name__ == "__main__":
    build_excel()
    build_zenith_docx()
    print("Done.")
