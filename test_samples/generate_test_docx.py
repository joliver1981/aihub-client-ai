"""
Generate sample .docx files to test the enhanced docx extraction.

Each sample stresses a different content-bearing area of the OOXML structure:
  - Sample 1: paragraphs + body table (pre-fix should already work)
  - Sample 2: paragraphs + headers + footers (pre-fix would miss header/footer)
  - Sample 3: paragraphs + body text box (pre-fix would miss text-box content)
  - Sample 4: everything combined

Run with the project's python:
  C:\\Users\\james\\miniconda3\\envs\\aihub2.1\\python.exe generate_test_docx.py
"""
from pathlib import Path
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

OUT_DIR = Path(__file__).parent / "docx"
OUT_DIR.mkdir(parents=True, exist_ok=True)


def _text_box_run_xml(text: str) -> str:
    """OOXML for a run containing an inline text box with `text` inside.

    The actual extracted content lives in `<w:txbxContent>`, which python-docx
    does NOT expose via doc.paragraphs/doc.tables — this is the case the
    enhanced extractor handles.
    """
    return f'''
<w:r {nsdecls('w')}
     xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"
     xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
     xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
     xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
  <mc:AlternateContent>
    <mc:Choice Requires="wps">
      <w:drawing>
        <wp:inline distT="0" distB="0" distL="0" distR="0">
          <wp:extent cx="3200400" cy="900000"/>
          <wp:effectExtent l="0" t="0" r="0" b="0"/>
          <wp:docPr id="100" name="Text Box"/>
          <wp:cNvGraphicFramePr/>
          <a:graphic>
            <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
              <wps:wsp>
                <wps:cNvSpPr txBox="1"/>
                <wps:spPr>
                  <a:xfrm><a:off x="0" y="0"/><a:ext cx="3200400" cy="900000"/></a:xfrm>
                  <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
                  <a:solidFill><a:srgbClr val="FFFFFF"/></a:solidFill>
                  <a:ln w="9525"><a:solidFill><a:srgbClr val="000000"/></a:solidFill></a:ln>
                </wps:spPr>
                <wps:txbx>
                  <w:txbxContent>
                    <w:p>
                      <w:r>
                        <w:t xml:space="preserve">{text}</w:t>
                      </w:r>
                    </w:p>
                  </w:txbxContent>
                </wps:txbx>
                <wps:bodyPr rot="0" vert="horz" wrap="square"
                            lIns="91440" tIns="45720" rIns="91440" bIns="45720"
                            anchor="t" anchorCtr="0"/>
              </wps:wsp>
            </a:graphicData>
          </a:graphic>
        </wp:inline>
      </w:drawing>
    </mc:Choice>
  </mc:AlternateContent>
</w:r>
'''.strip()


def add_text_box(paragraph, text: str) -> None:
    """Append an inline text box containing `text` to the given paragraph."""
    run_element = parse_xml(_text_box_run_xml(text))
    paragraph._p.append(run_element)


# ---------------------------------------------------------------------------
# Sample 1 — paragraphs + table (control: should work before AND after fix)
# ---------------------------------------------------------------------------
def make_sample_1():
    doc = Document()
    doc.add_heading("Acme Widget Co. — Q3 2025 Operations Report", level=1)
    doc.add_paragraph(
        "Acme Widget Co. shipped 47,302 Model-X widgets during Q3 2025, "
        "exceeding the quarterly target of 45,000 units. The Cleveland plant "
        "led production at 18,400 units, followed by Phoenix (14,902) and "
        "Atlanta (14,000)."
    )
    doc.add_paragraph(
        "Defect rate fell to 0.37 percent, the lowest in plant history, "
        "driven by the new automated optical inspection line installed in July."
    )

    doc.add_heading("Plant Performance", level=2)
    table = doc.add_table(rows=4, cols=3)
    table.style = "Light Grid Accent 1"
    headers = table.rows[0].cells
    headers[0].text = "Plant"
    headers[1].text = "Units Shipped"
    headers[2].text = "Defect Rate"
    rows = [
        ("Cleveland", "18,400", "0.29%"),
        ("Phoenix",   "14,902", "0.41%"),
        ("Atlanta",   "14,000", "0.44%"),
    ]
    for i, (plant, units, defect) in enumerate(rows, start=1):
        cells = table.rows[i].cells
        cells[0].text = plant
        cells[1].text = units
        cells[2].text = defect

    out = OUT_DIR / "sample_1_paragraphs_and_table.docx"
    doc.save(out)
    print(f"wrote {out}")


# ---------------------------------------------------------------------------
# Sample 2 — paragraphs + headers + footers (pre-fix would miss header/footer)
# ---------------------------------------------------------------------------
def make_sample_2():
    doc = Document()

    section = doc.sections[0]
    header = section.header
    header.paragraphs[0].text = "CONFIDENTIAL — Project Halcyon Memo #HAL-2025-118"
    footer = section.footer
    footer.paragraphs[0].text = "Distribution restricted to clearance level Blue or higher."

    doc.add_heading("Project Halcyon — Status Update", level=1)
    doc.add_paragraph(
        "The Halcyon team completed integration testing of the new payload "
        "telemetry subsystem on 2025-09-14. All twelve checkout milestones "
        "passed without exception."
    )
    doc.add_paragraph(
        "The next gate review is scheduled for 2025-11-03 in Hangar 7. "
        "Stakeholders should submit pre-read materials by 2025-10-27."
    )

    out = OUT_DIR / "sample_2_headers_and_footers.docx"
    doc.save(out)
    print(f"wrote {out}")


# ---------------------------------------------------------------------------
# Sample 3 — paragraphs + text box (THE main fix)
# Most of the load-bearing content lives in the text box. Pre-fix this would
# extract only the surrounding paragraphs and miss the critical numbers.
# ---------------------------------------------------------------------------
def make_sample_3():
    doc = Document()
    doc.add_heading("Northstar Capital — Investor Brief", level=1)
    doc.add_paragraph(
        "The following call-out summarizes the fund's headline metrics for "
        "the period ending 2025-09-30. Full methodology is provided in the "
        "appendix on request."
    )

    # Critical content placed inside a text box — what users would call out
    # in a brochure or marketing one-pager. Pre-fix: invisible to extractor.
    callout_para = doc.add_paragraph()
    add_text_box(
        callout_para,
        "KEY METRICS — Northstar Opportunity Fund III. "
        "Net IRR: 18.4 percent. "
        "Total commitments: 412 million USD. "
        "Fund inception: 2022-04-15. "
        "Lead portfolio manager: Dana Rivera."
    )

    doc.add_paragraph(
        "Please direct questions to investor.relations@northstar.example."
    )

    out = OUT_DIR / "sample_3_text_box.docx"
    doc.save(out)
    print(f"wrote {out}")


# ---------------------------------------------------------------------------
# Sample 4 — everything combined
# ---------------------------------------------------------------------------
def make_sample_4():
    doc = Document()

    section = doc.sections[0]
    section.header.paragraphs[0].text = (
        "Helios Pharmaceuticals — Internal Trial Brief — Document ID HEL-TR-2208"
    )
    section.footer.paragraphs[0].text = (
        "Approved for internal distribution only. Author: Dr. Priya Anand."
    )

    doc.add_heading("HELIOS-22 Phase II Interim Readout", level=1)
    doc.add_paragraph(
        "The HELIOS-22 phase II trial enrolled 312 patients across 14 sites "
        "between 2024-06 and 2025-03. The primary endpoint was a 30 percent "
        "reduction in serum biomarker BX-7 at 12 weeks."
    )

    doc.add_heading("Site Enrollment", level=2)
    table = doc.add_table(rows=4, cols=2)
    table.style = "Light Grid Accent 1"
    table.rows[0].cells[0].text = "Region"
    table.rows[0].cells[1].text = "Patients Enrolled"
    for i, (region, n) in enumerate([("North America", "148"),
                                     ("Europe",        "102"),
                                     ("Asia-Pacific",   "62")], start=1):
        table.rows[i].cells[0].text = region
        table.rows[i].cells[1].text = n

    doc.add_paragraph(
        "An interim safety summary is presented in the callout below."
    )

    callout_para = doc.add_paragraph()
    add_text_box(
        callout_para,
        "SAFETY SUMMARY — Serious adverse events: 4 of 312 patients (1.3 percent). "
        "No treatment-related deaths. "
        "Dose-limiting toxicity observed at the 800 mg arm — recommended "
        "phase III dose is 600 mg twice daily."
    )

    doc.add_paragraph(
        "The independent data monitoring committee meets on 2025-12-08 to "
        "review the interim analysis and recommend whether HELIOS-22 should "
        "proceed to phase III."
    )

    out = OUT_DIR / "sample_4_combined.docx"
    doc.save(out)
    print(f"wrote {out}")


if __name__ == "__main__":
    make_sample_1()
    make_sample_2()
    make_sample_3()
    make_sample_4()
    print(f"\nAll samples in: {OUT_DIR}")
