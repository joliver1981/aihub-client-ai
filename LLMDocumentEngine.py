import anthropic
import base64
# import httpx
import time
import config as cfg

import os
import json
import re
import uuid
import shutil
import logging
from logging.handlers import WatchedFileHandler
import yaml
import pyodbc
from typing import List, Dict, Any, Optional, Tuple, Union
import pandas as pd

import chromadb

#from anthropic import Anthropic
from datetime import datetime
import PyPDF2
import io
import math

from CommonUtils import get_db_connection, get_db_connection_string, normalize_boolean, AnthropicProxyClient, rotate_logs_on_startup, get_log_path
from api_keys_config import get_anthropic_config, create_anthropic_client
from anthropic_streaming_helper import anthropic_messages_create

# Backward compatability
from LLMDocumentVectorAdapter import LLMDocumentVectorAdapter

from LLMDocumentSummarizer import DocumentPageSummarizer
from vector_engine_client import VectorEngineClient

# Fast PDF extraction - tries PyMuPDF first, falls back to AI when needed
try:
    from fast_pdf_extractor import FastPDFExtractor, PDFType
    FAST_PDF_AVAILABLE = True
except ImportError:
    FAST_PDF_AVAILABLE = False


# Configure logging
def setup_logging():
    """Configure logging"""
    logger = logging.getLogger("LLMDocumentEngine")
    log_level_name = os.getenv('LOG_LEVEL', 'DEBUG')
    log_level = getattr(logging, log_level_name, logging.DEBUG)
    logger.setLevel(log_level)
    formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
    handler = WatchedFileHandler(filename=os.getenv('LLM_DOCUMENT_ENGINE', get_log_path('llm_document_engine_log.txt')), encoding='utf-8')
    handler.setFormatter(formatter)
    logger.addHandler(handler)
    
    return logger

rotate_logs_on_startup(os.getenv('LLM_DOCUMENT_ENGINE', get_log_path('llm_document_engine_log.txt')))

logger = setup_logging()


database_server = cfg.DATABASE_SERVER
database_name = cfg.DATABASE_NAME
username = cfg.DATABASE_UID
password = cfg.DATABASE_PWD

class MultiPagePDFHandler:
    """
    Handles extraction of text from multi-page PDF documents using Claude's vision capabilities.
    """
    
    def __init__(self, anthropic_client, logger=None, anthropic_config=None, anthropic_proxy_client=None):
        """
        Initialize the PDF handler.
        
        Args:
            anthropic_client: Initialized Anthropic API client
            logger: Optional logger instance
            anthropic_config: Configuration dict from get_anthropic_config()
            anthropic_proxy_client: AnthropicProxyClient instance for proxy mode
        """
        self.anthropic_client = anthropic_client
        self.logger = logger or logging.getLogger("MultiPagePDFHandler")
        self.anthropic_proxy_client = anthropic_proxy_client
        
        # Store config, default to direct API if client provided and no config
        if anthropic_config is None:
            self._anthropic_config = {
                'use_direct_api': anthropic_client is not None,
                'source': 'legacy'
            }
        else:
            self._anthropic_config = anthropic_config
        
    def extract_from_pdf(
        self, 
        file_path: str, 
        document_type: str = "document",
        use_batch_processing: bool = False,
        batch_size: int = 3,
        include_page_numbers: bool = True
    ) -> List[Dict[str, Any]]:
        """
        Extract text from a multi-page PDF file.
        
        Args:
            file_path: Path to the PDF file
            document_type: The type of document for customized extraction
            use_batch_processing: Whether to process multiple pages in a single request
            batch_size: Number of pages per batch if batch processing is enabled
            include_page_numbers: Whether to add page number prefixes to extracted text
            
        Returns:
            List of dictionaries with page number and extracted text
        """
        self.logger.info(f"Processing PDF: {file_path}")
        
        # Read the PDF file
        try:
            pdf_reader = PyPDF2.PdfReader(file_path)
            total_pages = len(pdf_reader.pages)
            self.logger.info(f"PDF has {total_pages} pages")
            
            extracted_pages = []
            
            if use_batch_processing and batch_size > 1:
                # Process pages in batches
                self.logger.info(f"Using batch processing with batch size {batch_size}")
                for start_page in range(0, total_pages, batch_size):
                    end_page = min(start_page + batch_size, total_pages)
                    page_range = range(start_page, end_page)
                    
                    self.logger.info(f"Processing pages {start_page+1}-{end_page} of {total_pages}")
                    
                    # Create a temporary PDF with just this batch of pages
                    temp_pdf_bytes = self._create_subset_pdf(file_path, page_range)
                    
                    # Extract text using Claude Vision
                    batch_pages = self._extract_batch_with_claude(
                        temp_pdf_bytes, 
                        document_type,
                        page_range,
                        include_page_numbers
                    )
                    
                    extracted_pages.extend(batch_pages)
            else:
                # Process pages individually (default)
                self.logger.info("Processing pages individually")
                for page_num in range(total_pages):
                    self.logger.info(f"Processing page {page_num+1} of {total_pages}")
                    
                    # Create a temporary PDF with just this single page
                    temp_pdf_bytes = self._create_subset_pdf(file_path, [page_num])
                    
                    # Extract text using Claude Vision
                    single_page = self._extract_single_page_with_claude(
                        temp_pdf_bytes,
                        document_type,
                        page_num + 1,  # Convert to 1-based page number
                        include_page_numbers
                    )
                    
                    extracted_pages.append(single_page)
                
            return extracted_pages
            
        except Exception as e:
            self.logger.error(f"Error extracting from PDF: {str(e)}")
            raise
            
    def _create_subset_pdf(self, file_path: str, pages: List[int]) -> bytes:
        """
        Create a subset PDF containing only the specified pages.
        
        Args:
            file_path: Path to the original PDF
            pages: List of page indices to include (0-indexed)
            
        Returns:
            Bytes of the subset PDF
        """
        pdf_reader = PyPDF2.PdfReader(file_path)
        pdf_writer = PyPDF2.PdfWriter()
        total_pages = len(pdf_reader.pages)

        # Filter pages to only include valid indices
        valid_pages = [page for page in pages if 0 <= page < total_pages]
        
        # Add the selected pages to the writer
        for i in valid_pages:
            pdf_writer.add_page(pdf_reader.pages[i])
            
        # Write to a bytes buffer
        output_buffer = io.BytesIO()
        pdf_writer.write(output_buffer)
        output_buffer.seek(0)
        
        return output_buffer.getvalue()
    
    def _extract_single_page_with_claude(
        self,
        pdf_bytes: bytes,
        document_type: str,
        page_number: int,
        include_page_numbers: bool
    ) -> Dict[str, Any]:
        """
        Extract text from a single PDF page using Claude.
        
        Args:
            pdf_bytes: PDF content as bytes
            document_type: Type of document
            page_number: Page number (1-indexed)
            include_page_numbers: Whether to add page numbers to extracted text
            
        Returns:
            Dictionary with page number and extracted text
        """
        # Create a system prompt based on document type
        system_prompt = (
            f"You are an expert at extracting text and data from {document_type} documents. "
            "Extract ALL text from the page maintaining its structure and formatting. "
            "Your task is to convert the PDF content to plain text that includes ALL information visible on the page. "
            "This includes headers, footers, tables, form fields, and any other text elements."
        )
        
        # User message for single page extraction
        user_text = (
            f"Extract the complete text from this {document_type} document page. "
            "Include ALL visible text and maintain the structure of tables and forms. "
            "Return ONLY the extracted text without any additional commentary."
        )
        
        # Call Claude Vision API
        try:
            if self._anthropic_config['use_direct_api']:
                # Encode the PDF for the API
                base64_pdf = base64.b64encode(pdf_bytes).decode('utf-8')

                # Use streaming wrapper (required for newer Anthropic models)
                response = anthropic_messages_create(
                    client=self.anthropic_client,
                    model=cfg.ANTHROPIC_MODEL,
                    max_tokens=int(cfg.ANTHROPIC_MAX_TOKENS),
                    system=system_prompt,
                    messages=[
                        {
                            "role": "user",
                            "content": [
                                {
                                    "type": "document",
                                    "source": {
                                        "type": "base64",
                                        "media_type": "application/pdf",
                                        "data": base64_pdf
                                    }
                                },
                                {
                                    "type": "text", 
                                    "text": user_text
                                }
                            ]
                        }
                    ]
                )

                # Extract the text from Claude's response
                extracted_text = response.content[0].text.strip()
            else:
                # Encode the PDF for the API
                base64_pdf = base64.b64encode(pdf_bytes).decode('utf-8')

                messages = [
                        {
                            "role": "user",
                            "content": [
                                {
                                    "type": "document",
                                    "source": {
                                        "type": "base64",
                                        "media_type": "application/pdf",
                                        "data": base64_pdf
                                    }
                                },
                                {
                                    "type": "text", 
                                    "text": user_text
                                }
                            ]
                        }
                    ]
                
                if not self.anthropic_proxy_client:
                    client = AnthropicProxyClient()
                else:
                    client = self.anthropic_proxy_client
                response = client.messages_create(
                    messages=messages,
                    model=cfg.ANTHROPIC_MODEL,
                    max_tokens=int(cfg.ANTHROPIC_MAX_TOKENS),
                    system=system_prompt
                )

                # Extract the text from Claude's response
                extracted_text = response['content'][0]['text']
            
            # Add page number prefix if requested
            if include_page_numbers:
                extracted_text = f"[Page {page_number}]\n{extracted_text}"

            if cfg.ANTHROPIC_API_THROTTLE_CALLS:
                print('API throttling active, waiting before submitting next request...')
                time.sleep(int(cfg.ANTHROPIC_API_THROTTLE_DELAY))
            
            return {
                "page_number": page_number,
                "text": extracted_text
            }
            
        except Exception as e:
            self.logger.error(f"Error calling Claude API for page {page_number}: {str(e)}")
            raise
    
    def _extract_batch_with_claude(
        self, 
        pdf_bytes: bytes, 
        document_type: str,
        page_range: range,
        include_page_numbers: bool
    ) -> List[Dict[str, Any]]:
        """
        Extract text from a batch of PDF pages using Claude.
        
        Args:
            pdf_bytes: PDF content as bytes
            document_type: Type of document
            page_range: Range of pages in this batch (0-indexed)
            include_page_numbers: Whether to add page numbers to extracted text
            
        Returns:
            List of dictionaries with page number and extracted text
        """
        # Encode the PDF for the API
        base64_pdf = base64.b64encode(pdf_bytes).decode('utf-8')
        
        # Create a system prompt based on document type
        system_prompt = (
            f"You are an expert at extracting text and data from {document_type} documents. "
            "Extract ALL text from each page maintaining its structure and formatting. "
            "Your task is to convert the PDF content to plain text that includes ALL information visible on the page. "
            "This includes headers, footers, tables, form fields, and any other text elements. "
            "For each page, clearly indicate the page number at the beginning of the page's content like this: 'PAGE X:'"
        )
        
        # Generate user message for batch processing
        user_text = (
            f"Extract the complete text from this {document_type} document. "
            "The PDF contains multiple pages. Process each page separately and clearly indicate where each page begins and ends. "
            "For each page, start with 'PAGE X:' where X is the page number. "
            "Include ALL visible text and maintain the structure of tables and forms."
        )

        messages=[
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "document",
                                "source": {
                                    "type": "base64",
                                    "media_type": "application/pdf",
                                    "data": base64_pdf
                                }
                            },
                            {
                                "type": "text", 
                                "text": user_text
                            }
                        ]
                    }
                ]
        
        # Call Claude Vision API
        try:
            if self._anthropic_config['use_direct_api']:
                # Use streaming wrapper (required for newer Anthropic models)
                response = anthropic_messages_create(
                    client=self.anthropic_client,
                    model=cfg.ANTHROPIC_MODEL,
                    max_tokens=int(cfg.ANTHROPIC_MAX_TOKENS),
                    system=system_prompt,
                    messages=messages
                )
                # print('Raw response:', response)
                # Extract and process the response
                extracted_text = response.content[0].text
            else:
                if not self.anthropic_proxy_client:
                    client = AnthropicProxyClient()
                else:
                    client = self.anthropic_proxy_client
                response = client.messages_create(
                    messages=messages,
                    model=cfg.ANTHROPIC_MODEL,
                    max_tokens=int(cfg.ANTHROPIC_MAX_TOKENS),
                    system=system_prompt
                )
                #print('Raw response content:', response['content'])
                # Extract and process the response
                extracted_text = response['content'][0]['text']

            # print('Raw response:', response)
            print(86 * '#')
            print('Extracted text:', extracted_text)
            
            # Split the response by page indicators
            return self._parse_claude_response(extracted_text, page_range, include_page_numbers)
            
        except Exception as e:
            self.logger.error(f"Error calling Claude API for batch: {str(e)}")
            raise


    def _remove_json_comments(self, json_str):
        # Remove all //... comments
        json_str = re.sub(r'//.*', '', json_str)
        # Optionally, remove /* ... */ comments too:
        json_str = re.sub(r'/\*.*?\*/', '', json_str, flags=re.DOTALL)
        return json_str

    
    def _parse_claude_response(
        self, 
        response_text: str, 
        page_range: range,
        include_page_numbers: bool
    ) -> List[Dict[str, Any]]:
        """
        Parse Claude's response to extract individual pages.
        
        Args:
            response_text: Text response from Claude
            page_range: Range of pages in this batch (0-indexed)
            include_page_numbers: Whether to keep page numbers in text
            
        Returns:
            List of dictionaries with page number and extracted text
        """
        result = []

        # Convert range to list for easier indexing
        page_numbers = list(page_range)

        print('page_range (input):', page_range)
        #print('page_numbers:', page_numbers)

        # Look for common page separator patterns
        page_patterns = [
            r'(?i)PAGE\s*(\d+):',  # "PAGE X:" or "Page X:" (our requested format)
            r'(?i)Page\s*(\d+)[:\-]?\s*\n',  # "Page X:" or "Page X" or "PAGE X:"
            r'(?i)---+\s*Page\s*(\d+)\s*---+',  # "---- Page X ----"
            r'(?i)##\s*Page\s*(\d+)\s*##',  # "## Page X ##"
            r'(?i)Page\s*(\d+)\s*of\s*\d+',  # "Page X of Y"
        ]
        
        # Try each pattern to find page breaks
        page_breaks = []
        for pattern in page_patterns:
            matches = list(re.finditer(pattern, response_text))
            if matches:
                page_breaks = [(m.start(), m.end(), int(m.group(1))) for m in matches]
                break
                
        # If we couldn't find page breaks with patterns, try alternative approach
        if not page_breaks:
            # Look for lines that might indicate page breaks
            lines = response_text.split('\n')
            for i, line in enumerate(lines):
                # Check if line has "page" and a number
                match = re.search(r'(?i)page\s*(\d+)', line)
                if match and (len(line) < 80):  # Likely a header, not content
                    page_num = int(match.group(1))
                    start_pos = sum(len(lines[j]) + 1 for j in range(i))
                    end_pos = start_pos + len(line) + 1
                    page_breaks.append((start_pos, end_pos, page_num))
        
        print('page_breaks:', page_breaks)

        # Extract text between page breaks
        if page_breaks:

            if len(page_range) == 1:
                print('Parsing single page document w/ breaks...')
                # Single page
                result.append({
                    "page_number": page_range[0] + 1,  # Convert to 1-based while respecting current place in batch
                    "text": response_text
                })
            else:
                print('Parsing multi-page document w/ breaks...')
                # Sort by position in text
                page_breaks.sort(key=lambda x: x[0])
                print('page_breaks (sorted):', page_breaks)

                # Extract text for each page
                print('Extracting text for each page...')
                for i, (start, end, extracted_page_num) in enumerate(page_breaks):
                    print('i, start, end, page_num', i, start, end, extracted_page_num)
                    if i < len(page_breaks) - 1:
                        next_start = page_breaks[i+1][0]
                        page_text = response_text[end:next_start]
                        print('next_start', next_start)
                    else:
                        page_text = response_text[end:]
                    
                    # Clean up the text
                    page_text = page_text.strip()

                    # Determine the actual page number from our page_numbers list
                    # If we have matched page numbers from the text, try to map them
                    if i < len(page_numbers):
                        actual_page_num = page_numbers[i]  # Use the page number from our range
                    elif i > len(page_numbers):
                        # Check if this page is a duplicate of the previous page
                        if i > 0 and page_text == result[-1]["text"].replace(f"[Page {result[-1]['page_number']}]\n", ""):
                            # It's an exact duplicate - skip this page
                            print('Duplicate page found, skipping page...')
                            self.logger.warning("Duplicate page found, skipping page...")
                            continue
                        else:
                            # Not a duplicate - assign it a new incremented page number
                            if result:
                                actual_page_num = result[-1]["page_number"]  # Already 1-based from previous iteration
                            else:
                                actual_page_num = page_numbers[-1] + 1 if page_numbers else 0
                    else:
                        # Use the last page number
                        actual_page_num = page_numbers[-1] if page_numbers else 0
                        
                    # Add page indicator if requested
                    if include_page_numbers:
                        # Convert to 1-based page number for display
                        display_page_num = actual_page_num + 1
                        page_text = f"[Page {display_page_num}]\n{page_text}"
                        print('display_page_num', display_page_num)
                    
                    result.append({
                        "page_number": actual_page_num + 1,  # Convert to 1-based
                        "text": page_text
                    })
        else:
            # If we couldn't find page breaks, treat the whole response as a single page
            # or split evenly if we know there are multiple pages
            if len(page_range) == 1:
                print('Parsing single page document...')
                # Single page
                result.append({
                    "page_number": page_range[0] + 1,  # Convert to 1-based
                    "text": response_text
                })
            else:
                # Multiple pages but couldn't identify breaks
                # Try a simple heuristic: split by character count
                self.logger.warning("Could not identify page breaks in Claude response. Using character-based splitting.")
                print("Could not identify page breaks in Claude response. Using character-based splitting.")
                chars_per_page = len(response_text) // len(page_range)
                for i, page_num in enumerate(page_range):
                    start = i * chars_per_page
                    end = start + chars_per_page if i < len(page_range) - 1 else len(response_text)
                    page_text = response_text[start:end].strip()
                    
                    if include_page_numbers:
                        absolute_page_num = page_num + 1  # Convert to 1-based
                        page_text = f"[Page {absolute_page_num}]\n{page_text}"
                    
                    result.append({
                        "page_number": page_num + 1,  # Convert to 1-based
                        "text": page_text
                    })
        
        return result

    def write_subset_pdf_to_temp_file(self, file_path: str, pages: List[int]) -> str:
        """
        Create a subset PDF containing only the specified pages and write it to a temporary file.
        
        Args:
            file_path: Path to the original PDF
            pages: List of page indices to include (0-indexed)
            
        Returns:
            Path to the temporary PDF file
        """
        # Create subset PDF bytes
        pdf_bytes = self._create_subset_pdf(file_path, pages)
        
        # Ensure tmp directory exists
        tmp_dir = "./tmp"
        os.makedirs(tmp_dir, exist_ok=True)
        
        # Generate unique temporary filename
        original_filename = os.path.splitext(os.path.basename(file_path))[0]
        temp_filename = f"{original_filename}_subset_{uuid.uuid4().hex[:8]}.pdf"
        temp_file_path = os.path.join(tmp_dir, temp_filename)
        
        # Write bytes to file
        with open(temp_file_path, 'wb') as f:
            f.write(pdf_bytes)
            
        self.logger.info(f"Created temporary subset PDF: {temp_file_path}")
        return temp_file_path



class LLMDocumentProcessor:
    """
    A flexible system to process various document types using Claude Vision,
    extract structured data, and store in vector and relational databases.
    """
    
    def __init__(
        self, 
        vector_db_path: str = "./chroma_db",
        schema_dir: str = "./schemas",
        sql_connection_string: Optional[str] = f"DRIVER={{SQL Server}};SERVER={database_server};DATABASE={database_name};UID={username};PWD={password}",
        log_level: str = "DEBUG"
    ):
        """
        Initialize the document processor with necessary configurations.
        
        Args:
            vector_db_path: Path to store ChromaDB
            schema_dir: Directory containing document schemas
            sql_connection_string: Connection string for SQL Server (optional)
            log_level: Logging level
        """
        # Set up logging
        self.logger = logger

        # Initialize API client based on BYOK/config settings
        self._anthropic_config = get_anthropic_config()

        if self._anthropic_config['use_direct_api']:
            self.anthropic_client = anthropic.Anthropic(
                api_key=self._anthropic_config['api_key']
            )
            self.logger.info(f"Using direct Anthropic API (source: {self._anthropic_config['source']})")
        else:
            self.anthropic_client = None  # Will use proxy instead
            self.logger.info("Using Anthropic proxy for document processing")

        self.anthropic_proxy_client = AnthropicProxyClient()

        # Document summarization feature
        default_summary_types = cfg.DOC_AUTO_SUMMARY_TYPES
        enable_summarization = cfg.DOC_ENABLE_AUTO_SUMMARIZATION
        self.enable_summarization = enable_summarization
        self.default_summary_types = default_summary_types or ['standard']

        # Initialize the summarizer
        if self.enable_summarization:
            self.summarizer = DocumentPageSummarizer(
                anthropic_client=self.anthropic_client,
                logger=self.logger,
                anthropic_config=self._anthropic_config,
                anthropic_proxy_client=self.anthropic_proxy_client
            )
            self.logger.info("Document summarization enabled")
        else:
            self.summarizer = None
            self.logger.info("Document summarization disabled")
        
        # Initialize vector database
        # self.chroma_client = chromadb.PersistentClient(path=vector_db_path)
        # self.collection = self.chroma_client.get_or_create_collection(
        #     name="documents",
        #     metadata={"description": "Processed documents for analysis and retrieval", "hnsw:space": "cosine"}
        # )

        # Initialize vector database using the adapter
        self.vector_adapter = LLMDocumentVectorAdapter(
            use_remote=True,  # Change to True to use remote API
            vector_db_path=vector_db_path,
            collection_name="documents",
            log_level=log_level
        )
        # For backward compatibility with existing code
        self.collection = self.vector_adapter

        
        # Initialize SQL connection (if provided)
        self.sql_connection_string = sql_connection_string
        self.sql_conn = None
        if sql_connection_string:
            try:
                self.sql_conn = pyodbc.connect(sql_connection_string)
                self.logger.info("Connected to SQL Server database")
                self._ensure_database_tables()
            except Exception as e:
                self.logger.error(f"Failed to connect to SQL database: {str(e)}")
        
        # Load document schemas
        self.schema_dir = schema_dir
        self.schemas = self._load_schemas()

    def _load_schemas(self) -> Dict[str, Any]:
        """
        Load document schemas from the schema directory.
        
        Returns:
            Dictionary of document types and their schemas
        """
        schemas = {}
        
        # Create schema directory if it doesn't exist
        os.makedirs(self.schema_dir, exist_ok=True)
        
        # Load each schema file
        for filename in os.listdir(self.schema_dir):
            if filename.endswith(('.yaml', '.yml')):
                try:
                    with open(os.path.join(self.schema_dir, filename), 'r') as f:
                        schema = yaml.safe_load(f)
                        doc_type = schema.get('document_type')
                        if doc_type:
                            schemas[doc_type] = schema
                            self.logger.info(f"Loaded schema for document type: {doc_type}")
                except Exception as e:
                    self.logger.error(f"Error loading schema {filename}: {str(e)}")
                    
        return schemas
    
    def _ensure_database_tables(self):
        """Create necessary database tables if they don't exist"""
        sql_conn = get_db_connection()
        if not sql_conn:
            return
            
        cursor = sql_conn.cursor()
        
        # Create documents table if it doesn't exist
        try:
            # Main documents table
            cursor.execute("""
            IF NOT EXISTS (SELECT * FROM sys.tables WHERE name = 'Documents')
            BEGIN
                CREATE TABLE Documents (
                    document_id VARCHAR(100) PRIMARY KEY,
                    filename VARCHAR(255) NOT NULL,
                    original_path VARCHAR(1000) NOT NULL,
                    document_type VARCHAR(100) NOT NULL,
                    page_count INT NOT NULL,
                    processed_at DATETIME NOT NULL,
                    archived_path VARCHAR(1000),
                    hash_value VARCHAR(100)
                )
            END
            """)
            
            # Document pages table
            cursor.execute("""
            IF NOT EXISTS (SELECT * FROM sys.tables WHERE name = 'DocumentPages')
            BEGIN
                CREATE TABLE DocumentPages (
                    page_id VARCHAR(100) PRIMARY KEY,
                    document_id VARCHAR(100) NOT NULL,
                    page_number INT NOT NULL,
                    full_text NVARCHAR(MAX),
                    vector_id VARCHAR(100),
                    FOREIGN KEY (document_id) REFERENCES Documents(document_id)
                )
            END
            """)
            
            # Document fields table (to store extracted key/value pairs)
            cursor.execute("""
            IF NOT EXISTS (SELECT * FROM sys.tables WHERE name = 'DocumentFields')
            BEGIN
                CREATE TABLE DocumentFields (
                    field_id INT IDENTITY(1,1) PRIMARY KEY,
                    page_id VARCHAR(100) NOT NULL,
                    field_name NVARCHAR(255) NOT NULL,
                    field_value NVARCHAR(MAX),
                    field_path NVARCHAR(500), -- JSON path to nested fields
                    confidence FLOAT,
                    FOREIGN KEY (page_id) REFERENCES DocumentPages(page_id)
                )
            END
            """)
            
            sql_conn.commit()
            self.logger.info("Database tables created or verified")
            
        except Exception as e:
            self.logger.error(f"Error creating database tables: {str(e)}")
            sql_conn.rollback()

    def process_document(
        self, 
        file_path: str, 
        document_type: Optional[str] = None,
        force_ai_extraction: bool = False,
        use_batch_processing: bool = False,  # New parameter
        batch_size: int = 3,
        execution_id: str = None   ,
        is_knowledge_document: bool = False, 
        do_not_store: bool = False, 
        extract_fields: bool = True, 
        detect_document_type: bool = True
    ) -> Dict[str, Any]:
        """
        Process a document using Claude Vision to extract structured data.
        
        Args:
            file_path: Path to the document file (PDF)
            document_type: Type of document to determine schema (optional)
            force_ai_extraction: If True, use AI to determine structure even if schema exists
            use_batch_processing: Whether to process multiple pages in a single API request
            batch_size: Number of pages per batch if batch processing is enabled
            
        Returns:
            Dictionary with processing results and document data
        """
        try:
            self.anthropic_proxy_client._set_tracking_params('document_processor')

            page_data_list = []
            extracted_pages = []
            # Get filename and basic file info
            filename = os.path.basename(file_path)
            file_extension = os.path.splitext(filename)[1].lower()
            file_id = str(uuid.uuid4())
            created_at = datetime.now().isoformat()
            
            print(86 * '-')
            print(f"Processing document: {filename} (ID: {file_id})")
            self.logger.info(f"Processing document: {filename} (ID: {file_id})")
            file_start_time = time.time()

            sql_conn = get_db_connection()
            cursor = sql_conn.cursor()

            print(f"Setting context...")
            self.logger.debug(f"Setting context...")
            cursor.execute("EXEC tenant.sp_setTenantContext ?", os.getenv('API_KEY'))

            # Insert file processing record
            if execution_id > 0:
                print(f"Inserting processing record...")
                self.logger.debug(f"Inserting processing record...")
                cursor.execute("""
                    INSERT INTO DocumentJobFileDetails (
                        ExecutionID, ProcessedAt, FileName, OriginalPath, Status
                    ) VALUES (?, getutcdate(), ?, ?, 'PROCESSING')
                """, (execution_id, filename, file_path))
                sql_conn.commit()

                #Get the file detail ID
                print(f"Getting identity...")
                self.logger.debug(f"Getting identity...")
                cursor.execute("SELECT @@IDENTITY")
                file_detail_id = cursor.fetchone()[0]
            else:
                file_detail_id = ''
        
            # Detect document type if not provided
            print(f"Detecting document type...")
            self.logger.debug(f"Detecting document type...")
            detected_document_type = document_type
            
            if detect_document_type:
                if not detected_document_type or force_ai_extraction:
                    detected_document_type = self._detect_document_type(file_path)
                    print(f"Detected document type: {detected_document_type}")
                    self.logger.info(f"Detected document type: {detected_document_type}")
            else:
                print("Skipping automatic detection of document type...")
                self.logger.info(f"Skipping automatic detection of document type...")

            if not detected_document_type:
                detected_document_type = 'unknown'

            # Process based on file type
            print(f"Calling document handler based on type {file_extension}...")
            self.logger.debug(f"Calling document handler based on type {file_extension}...")
            if file_extension == '.pdf':
                extracted_pages = self._process_pdf(file_path, detected_document_type, use_batch_processing, batch_size)
            elif file_extension in ['.jpg', '.jpeg', '.png', '.gif', '.webp']:
                extracted_pages = self._process_image(file_path, detected_document_type)
            elif file_extension in ['.docx', '.doc']:
                extracted_pages = self._process_word_document(file_path, detected_document_type)
            elif file_extension in ['.xlsx', '.xls']:
                extracted_pages = self._process_excel(file_path, detected_document_type)
            # elif file_extension in ['.pptx', '.ppt']:
            #     extracted_pages = self._process_powerpoint(file_path, detected_document_type)
            elif file_extension in ['.txt', '.md', '.csv', '.json', '.xml', '.html', '.htm']:
                extracted_pages = self._process_text_file(file_path, detected_document_type)
            else:
                self.logger.error(f"Invalid document type selected: {document_type}")
                print(f'ERROR: Invalid document type selected: {document_type}')
                extracted_pages = []
                
            # Process each page
            self.logger.debug(f"Processing pages...")
            page_data_list = []
            for page in extracted_pages:
                page_num = page["page_number"]
                page_content = page["text"]
                
                print(86 * '-')
                print(f'Processing page {page_num}')
                
                # Determine structure and extract fields
                extracted_data = {}
                if extract_fields:
                    if (detected_document_type in self.schemas) and not force_ai_extraction:
                        # Use predefined schema to extract fields
                        print('Extracting fields with schema...')
                        extracted_data = self._extract_with_schema(
                            page_content, 
                            self.schemas[detected_document_type]
                        )
                    else:
                        # Use AI to determine structure and extract fields
                        print('Extracting fields with ai...')
                        extracted_data = self._extract_with_ai(page_content, detected_document_type)
                
                # Create a structured record for this page
                page_data = {
                    "document_id": file_id,
                    "page_id": f"{file_id}_p{page_num}",
                    "filename": filename,
                    "page_number": page_num,
                    "document_type": detected_document_type,
                    "extraction_timestamp": created_at,
                    "full_text": page_content,
                    "extracted_data": extracted_data
                }
                
                page_data_list.append(page_data)
                
            if not do_not_store:
                # Store in vector database and SQL database
                VECTOR_LOADED = False
                if not is_knowledge_document and cfg.DOC_ENABLE_VECTOR_STORE:
                    print(f"Storing in vector database...")
                    print(f"Page data list: {page_data_list}")
                    self._store_in_vector_db(page_data_list)
                    VECTOR_LOADED = True

                if sql_conn:
                    print(f"Storing in SQL database...")
                    try:
                        self._store_in_sql_db(file_path, file_id, detected_document_type, page_data_list, is_knowledge_document)
                    except:
                        print(f"Error storing in SQL database")
                        if VECTOR_LOADED:
                            print(f"Deleting from vector database...")
                            ids = []
                            for page in page_data_list:
                                ids.append(page["page_id"])
                            self.vector_adapter.delete(ids=ids)
                        raise
            else:
                self.logger.info(f"Do not store flag set, document will not be saved...")
                print("Do not store flag set, document will not be saved...")

            # Update file processing record with success
            processing_time = time.time() - file_start_time
            # Update job file status in database
            sql_conn = get_db_connection()
            if sql_conn:
                print('Updating job file status...')
                cursor = sql_conn.cursor()
                cursor.execute("EXEC tenant.sp_setTenantContext ?", os.getenv('API_KEY'))
                cursor.execute("""
                    UPDATE DocumentJobFileDetails
                    SET Status = 'SUCCEEDED', 
                        DocumentType = ?,
                        DocumentID = ?,
                        PageCount = ?,
                        ProcessingDurationSeconds = ?
                    WHERE ExecutionID = ? AND FileName = ?
                """, (detected_document_type, file_id, len(page_data_list), processing_time, execution_id, filename))
                sql_conn.commit()
        except Exception as e:
            print(f"Error processing document: {str(e)}")
            self.logger.error(f"Error processing document: {str(e)}")
            # Update job file status in database
            processing_time = time.time() - file_start_time
            sql_conn = get_db_connection()
            if sql_conn:
                print('Updating job file status to FAILED...')
                cursor = sql_conn.cursor()
                cursor.execute("EXEC tenant.sp_setTenantContext ?", os.getenv('API_KEY'))
                cursor.execute("""
                    UPDATE DocumentJobFileDetails
                    SET Status = 'FAILED', 
                        DocumentType = ?,
                        DocumentID = ?,
                        PageCount = ?,
                        ProcessingDurationSeconds = ?,
                        ErrorMessage = ?
                    WHERE ExecutionID = ? AND FileName = ?
                """, (detected_document_type, file_id, len(page_data_list), processing_time, str(e), execution_id, filename))
                sql_conn.commit()
            
        # Return processing results
        return {
            "document_id": file_id,
            "filename": filename,
            "document_type": detected_document_type,
            "page_count": len(page_data_list),
            "processed_at": created_at,
            "pages": page_data_list,
            "file_detail_id": file_detail_id
        }
    
     
    def _detect_document_type(self, file_path: str) -> str:
        """
        Detect the document type using AI.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Detected document type
        """
        try:
            pdf_handler = MultiPagePDFHandler(
                self.anthropic_client, 
                self.logger,
                anthropic_config=self._anthropic_config,
                anthropic_proxy_client=self.anthropic_proxy_client
            )
            temp_pdf_file = None  # Initialize variable to track temp file for cleanup

            # Get filename and basic file info
            filename = os.path.basename(file_path)

            media_type, doc_type = self._get_file_and_document_type(file_name=filename)

            system_prompt = "You are a document classification expert. Your task is to identify the exact type of document shown in the image. Respond with ONLY the document type as a single term, like 'invoice', 'bill_of_lading', 'purchase_order', etc. Do not include any explanation or extra text."
            user_text = "What type of document is this? Respond with only the document type."

            if self._anthropic_config['use_direct_api']:
                # Use vision API
                if doc_type in ['document','image']:
                    if doc_type == 'document':
                        # Create a temporary PDF with just a subset of pages for type detection
                        print(f'Creating temporary PDF with pages {cfg.DOC_TYPE_DETECTION_PAGES_TO_USE}...')
                        temp_pdf_bytes = pdf_handler._create_subset_pdf(file_path, cfg.DOC_TYPE_DETECTION_PAGES_TO_USE)
                        file_content = base64.standard_b64encode(temp_pdf_bytes).decode("utf-8")
                    else:
                        with open(file_path, "rb") as f:
                            file_content = base64.standard_b64encode(f.read()).decode("utf-8")
                    
                    # Ask Claude to identify the document type
                    response = anthropic_messages_create(
                        client=self.anthropic_client,
                        model=cfg.ANTHROPIC_MODEL,
                        max_tokens=int(cfg.ANTHROPIC_MAX_TOKENS),
                        system=system_prompt,
                        messages=[
                            {
                                "role": "user",
                                "content": [
                                    {
                                        "type": doc_type,
                                        "source": {
                                            "type": "base64",
                                            "media_type": media_type,
                                            "data": file_content
                                        }
                                    },
                                    {
                                        "type": "text", 
                                        "text": user_text
                                    }
                                ]
                            }
                        ]
                    )
                    # Parse and clean the response
                    document_type = response.content[0].text.strip().lower()
                else:
                    # Get text from document
                    file_extension = os.path.splitext(filename)[1].lower()

                    document_text = self._get_text_from_other_documents(file_extension=file_extension, file_path=file_path)

                    user_text = (
                        f"Here is the document text:\n\n{document_text}\n\n"
                        "What type of document is this? Respond with only the document type."
                    )
                    messages = [
                                {
                                    "role": "user",
                                    "content": [
                                        {
                                            "type": "text", 
                                            "text": user_text
                                        }
                                    ]
                                }
                            ]
                    
                    # Ask Claude to extract structured data
                    response = anthropic_messages_create(
                        client=self.anthropic_client,
                        model=cfg.ANTHROPIC_MODEL,
                        max_tokens=int(cfg.ANTHROPIC_MAX_TOKENS),
                        system=system_prompt,
                        messages=messages
                    )
                    
                    # Parse and clean the response
                    document_type = response.content[0].text.strip().lower()
            else:
                #client = AnthropicProxyClient()
                client = self.anthropic_proxy_client
                if doc_type in ['document','image']:
                    if doc_type == 'document':
                        # Create a temporary PDF with just a subset of pages for type detection
                        print(f'Creating temporary PDF with pages {cfg.DOC_TYPE_DETECTION_PAGES_TO_USE}...')
                        temp_pdf_file = pdf_handler.write_subset_pdf_to_temp_file(file_path, cfg.DOC_TYPE_DETECTION_PAGES_TO_USE)
                        response = client.messages_with_document(
                            file_path=temp_pdf_file,
                            user_text=user_text,
                            model=cfg.ANTHROPIC_MODEL,
                            max_tokens=int(cfg.ANTHROPIC_MAX_TOKENS),
                            system=system_prompt
                        )
                    else:
                        response = client.messages_with_document(
                            file_path=file_path,
                            user_text=user_text,
                            model=cfg.ANTHROPIC_MODEL,
                            max_tokens=int(cfg.ANTHROPIC_MAX_TOKENS),
                            system=system_prompt
                        )
                else:
                    file_extension = os.path.splitext(filename)[1].lower()

                    document_text = self._get_text_from_other_documents(file_extension=file_extension, file_path=file_path)

                    user_text = (
                        f"Here is the document text:\n\n{document_text}\n\n"
                        "What type of document is this? Respond with only the document type."
                    )

                    messages = [
                                {
                                    "role": "user",
                                    "content": [
                                        {
                                            "type": "text", 
                                            "text": user_text
                                        }
                                    ]
                                }
                            ]
                    
                    response = client.messages_create(
                        messages=messages,
                        model=cfg.ANTHROPIC_MODEL,
                        max_tokens=int(cfg.ANTHROPIC_MAX_TOKENS),
                        system=system_prompt
                    )

                print(86 * '=')
                print('Response:', response)
                print(86 * '=')

                # Check for errors
                if "error" in response:
                    print('Error processing document with proxy', str(response["error"]))
                    document_type = 'unknown'
                else:
                    # Parse and clean the response
                    document_type = response['content'][0]['text'].strip().lower()

            # Clean up temporary PDF file if it was created
            if temp_pdf_file and os.path.exists(temp_pdf_file):
                try:
                    os.remove(temp_pdf_file)
                    self.logger.info(f"Deleted temporary PDF file: {temp_pdf_file}")
                except Exception as cleanup_error:
                    self.logger.warning(f"Failed to delete temporary PDF file {temp_pdf_file}: {str(cleanup_error)}")

            # Remove any non-alphanumeric characters and replace spaces with underscores
            document_type = ''.join(c if c.isalnum() or c == '_' else '_' for c in document_type)
            if document_type.startswith('_'):
                document_type = document_type[1:]

            if cfg.ANTHROPIC_API_THROTTLE_CALLS:
                print('API throttling active, waiting before submitting next request...')
                time.sleep(int(cfg.ANTHROPIC_API_THROTTLE_DELAY))
                
            return document_type
            
        except Exception as e:
            self.logger.error(f"Error detecting document type: {str(e)}")
            
            # Clean up temporary PDF file if it was created (even in case of error)
            if 'temp_pdf_file' in locals() and temp_pdf_file and os.path.exists(temp_pdf_file):
                try:
                    os.remove(temp_pdf_file)
                    self.logger.info(f"Deleted temporary PDF file: {temp_pdf_file}")
                except Exception as cleanup_error:
                    self.logger.warning(f"Failed to delete temporary PDF file {temp_pdf_file}: {str(cleanup_error)}")
            
            return "unknown_document"
        
    def _get_text_from_other_documents(self, file_extension, file_path):
        if file_extension in ['.docx', '.doc']:
            extracted_pages = self._process_word_document(file_path, 'unknown')
        elif file_extension in ['.xlsx', '.xls']:
            extracted_pages = self._process_excel(file_path, 'unknown')
        # elif file_extension in ['.pptx', '.ppt']:
        #     extracted_pages = self._process_powerpoint(file_path, detected_document_type)
        elif file_extension in ['.txt', '.md', '.csv', '.json', '.xml', '.html', '.htm']:
            extracted_pages = self._process_text_file(file_path, 'unknown')
        else:
            self.logger.error(f"Invalid document type selected: {file_extension}")
            print(f'ERROR: Invalid document type selected: {file_extension}')
            extracted_pages = []

        if not extracted_pages:
            return {"text": ""}
        
        # For type detection, only return the first page (truncated to 10K chars)
        # — sending millions of chars to Claude for classification is wasteful
        page = extracted_pages[0]
        if isinstance(page, dict) and len(page.get("text", "")) > 10_000:
            page = {**page, "text": page["text"][:10_000] + "\n\n[... truncated for type detection ...]"}
        return page
    
    def _call_claude_vision(self, file_content: bytes, document_type: str) -> List[str]:
        """
        Call Claude Vision API to extract text from document pages.
        
        Args:
            file_content: Binary content of the document file
            document_type: Type of document for specialized extraction
            
        Returns:
            List of extracted text content, one item per page
        """
        try:
            # Create a specialized system prompt based on document type
            system_prompt = f"You are an expert at extracting detailed information from {document_type} documents. "
            system_prompt += "Extract all text and data fields precisely, maintaining their relationships. "
            system_prompt += "Include all identifiers, dates, addresses, quantities, and other relevant details."
            
            user_text = f"Extract all text and data from this {document_type} document. Format the output as plain text that preserves all information, including all numbers, dates, and identifiers."

            if self._anthropic_config['use_direct_api']:
                response = anthropic_messages_create(
                    client=self.anthropic_client,
                    model=cfg.ANTHROPIC_MODEL,
                    max_tokens=int(cfg.ANTHROPIC_MAX_TOKENS),
                    system=system_prompt,
                    messages=[
                        {
                            "role": "user",
                            "content": [
                                {
                                    "type": "document",
                                    "source": {
                                        "type": "base64",
                                        "media_type": "application/pdf",
                                        "data": file_content
                                    }
                                },
                                {
                                    "type": "text", 
                                    "text": user_text
                                }
                            ]
                        }
                    ]
                )
            else:
                # Decode base64 string to bytes (proxy client will encode)
                file_bytes = base64.b64decode(file_content)

                client = AnthropicProxyClient()
                response = client.messages_with_document(
                    file_bytes=file_bytes,
                    user_text=user_text,
                    model=cfg.ANTHROPIC_MODEL,
                    max_tokens=int(cfg.ANTHROPIC_MAX_TOKENS),
                    system=system_prompt
                )

            print(86 * '=')
            print(response)
            print(86 * '=')
            
            # Parse response - in a real implementation, would need to handle multi-page PDFs
            # For simplicity, treating as a single page for now
            extracted_text = response.content[0].text

            # For now, we'll simulate page extraction as a single page
            # In production, you'd need to process multi-page PDFs properly
            pages = [extracted_text]
            
            return pages
            
        except Exception as e:
            self.logger.error(f"Error calling Claude Vision API: {str(e)}")
            return ["Error extracting text: " + str(e)]
    
    def _extract_with_schema(self, text: str, schema: Dict[str, Any]) -> Dict[str, Any]:
        """
        Extract structured data using a predefined schema.
        
        Args:
            text: Extracted text from a document page
            schema: Schema definition for the document type
            
        Returns:
            Dictionary of structured fields
        """
        extracted_data = {}
        
        # Get field definitions from schema
        fields = schema.get('fields', {})
        
        # Extract each field using patterns defined in schema
        for field_name, field_info in fields.items():
            pattern = field_info.get('pattern')
            
            if pattern:
                value = self._extract_field(text, pattern)
                
                # Apply any transformations defined in the schema
                transform = field_info.get('transform')
                if transform == 'to_number' and value:
                    try:
                        value = float(value.replace(',', ''))
                    except ValueError:
                        pass
                elif transform == 'to_date' and value:
                    # Keep as string but ensure consistent format if possible
                    pass
                
                # Handle nested fields using dot notation in field name
                if '.' in field_name:
                    parts = field_name.split('.')
                    current = extracted_data
                    for part in parts[:-1]:
                        if part not in current:
                            current[part] = {}
                        current = current[part]
                    current[parts[-1]] = value
                else:
                    extracted_data[field_name] = value
        
        return extracted_data
    
    def _remove_json_comments(self, json_str):
        # Remove all //... comments
        json_str = re.sub(r'//.*', '', json_str)
        # Optionally, remove /* ... */ comments too:
        json_str = re.sub(r'/\*.*?\*/', '', json_str, flags=re.DOTALL)
        return json_str
    
    def _extract_with_ai(self, text: str, document_type: str) -> Dict[str, Any]:
        """
        Use AI to determine structure and extract fields.
        
        Args:
            text: Extracted text from a document page
            document_type: Type of document
            
        Returns:
            Dictionary of structured fields
        """
        try:
            system_prompt = f"You are an expert in extracting structured data from {document_type} documents. Extract all key fields and their values into a clean JSON structure. Use nested objects for logical grouping. Be comprehensive but focus on important business information."
            user_text = f"Extract all the key information from this {document_type} document text into a clean, well-structured JSON format. Include all identifiers, dates (formatted as YYYY-MM-DD), amounts, parties involved, and any other relevant information.\n\nDocument Text:\n{text}"
            messages = [
                        {
                            "role": "user",
                            "content": [
                                {
                                    "type": "text", 
                                    "text": user_text
                                }
                            ]
                        }
                    ]
            # Ask Claude to extract structured data
            if self._anthropic_config['use_direct_api']:
                response = anthropic_messages_create(
                    client=self.anthropic_client,
                    model=cfg.ANTHROPIC_MODEL,
                    max_tokens=int(cfg.ANTHROPIC_MAX_TOKENS),
                    system=system_prompt,
                    messages=messages
                )
                # Extract and parse the JSON from the response
                ai_response = response.content[0].text
            else:
                if not self.anthropic_proxy_client:
                    client = AnthropicProxyClient()
                else:
                    client = self.anthropic_proxy_client
                response = client.messages_create(
                    messages=messages,
                    model=cfg.ANTHROPIC_MODEL,
                    max_tokens=int(cfg.ANTHROPIC_MAX_TOKENS),
                    system=system_prompt
                )
                # Extract and parse the JSON from the response
                ai_response = response['content'][0]['text']

            # print('RAW RESPONSE:', response)
            print(86 * '$')
            print('===== _extract_with_ai (DEBUG - before parse) =====')
            print('Claude Response:', ai_response)
            print(86 * '$')
            
            # Find JSON in the response (it might be wrapped in backticks or explanation)
            import re
            json_match = re.search(r'```(?:json)?\s*([\s\S]+?)\s*```', ai_response)
            if json_match:
                json_str = json_match.group(1)
            else:
                # If not in code block, try to find a JSON-like structure
                json_str = ai_response
                
            # Clean up the string for parsing
            json_str = json_str.strip()
            if not (json_str.startswith('{') and json_str.endswith('}')):
                # If still no valid JSON found, try to extract just the JSON part
                json_match = re.search(r'(\{[\s\S]+\})', json_str)
                if json_match:
                    json_str = json_match.group(1)
                
            # Parse the JSON
            try:
                extracted_data = json.loads(json_str)
            except:
                self.logger.error(f"Failed to load JSON response, attempting to correct output...")
                print('Failed to load JSON response, attempting to correct output...')
                json_str = self._remove_json_comments(json_str=json_str)
                extracted_data = json.loads(json_str)
                self.logger.info(f"Corrected JSON response.")
                print("Corrected JSON response.")
            
            # Save the AI-determined schema for future use
            self._save_ai_schema(document_type, extracted_data)

            if cfg.ANTHROPIC_API_THROTTLE_CALLS:
                print('API throttling active, waiting before submitting next request...')
                time.sleep(int(cfg.ANTHROPIC_API_THROTTLE_DELAY))
            
            return extracted_data
            
        except Exception as e:
            self.logger.error(f"Error extracting with AI: {str(e)}")
            # Return a minimal structure with error information
            return {
                "extraction_error": str(e),
                "raw_text_sample": text[:500] + ("..." if len(text) > 500 else "")
            }
    
    def _save_ai_schema(self, document_type: str, extracted_data: Dict[str, Any]):
        """
        Save AI-determined schema for future use.
        
        Args:
            document_type: Type of document
            extracted_data: Extracted structured data
        """
        # Only save if we don't already have a schema for this document type
        if document_type in self.schemas:
            return
            
        # Create a schema based on the AI extraction
        schema = {
            'document_type': document_type,
            'fields': {}
        }
        
        # Recursively build field definitions
        def add_fields(data, prefix=''):
            if isinstance(data, dict):
                for key, value in data.items():
                    path = f"{prefix}.{key}" if prefix else key
                    if isinstance(value, (dict, list)):
                        add_fields(value, path)
                    else:
                        # Add as a field with empty pattern (will need to be filled in later)
                        schema['fields'][path] = {'description': f"Auto-detected field: {path}"}
            elif isinstance(data, list) and data and isinstance(data[0], dict):
                # For lists of objects, just process the first one as an example
                add_fields(data[0], prefix + '[0]')
        
        add_fields(extracted_data)
        
        # Save the schema
        os.makedirs(self.schema_dir, exist_ok=True)
        schema_path = os.path.join(self.schema_dir, f"{document_type}_auto.yml")
        
        with open(schema_path, 'w') as f:
            yaml.dump(schema, f, default_flow_style=False)
            
        self.logger.info(f"Saved auto-generated schema for {document_type} to {schema_path}")
        
        # Add to loaded schemas
        self.schemas[document_type] = schema
    
    def _extract_field(self, text: str, pattern: str) -> str:
        """Extract a single field using regex pattern"""
        import re
        match = re.search(pattern, text)
        return match.group(1).strip() if match else ""
    
    def _extract_all_fields(self, text: str, pattern: str) -> List[str]:
        """Extract all occurrences of a field using regex pattern"""
        import re
        matches = re.findall(pattern, text)
        return [match.strip() for match in matches if match.strip()]
    
    def _store_in_vector_db(self, page_data_list: List[Dict[str, Any]]):
        """
        Store document pages in the vector database.
        
        Args:
            page_data_list: List of page data dictionaries
        """
        documents = []
        metadatas = []
        ids = []

        for page_data in page_data_list:
            # Create metadata for the vector DBs
            metadata = {
                "document_id": page_data["document_id"],
                "page_id": page_data["page_id"],
                "filename": page_data["filename"],
                "page_number": page_data["page_number"],
                "document_type": page_data["document_type"],
                "extraction_timestamp": page_data["extraction_timestamp"]
            }
            
            # Add important extracted fields to metadata for filtering
            for key, value in page_data["extracted_data"].items():
                if isinstance(value, (str, int, float, bool)) and key not in metadata:
                    # Truncate long values for metadata
                    if isinstance(value, str) and len(value) > 100:
                        value = value[:100] + "..."
                    metadata[key] = value

            # Add to our lists
            documents.append(page_data["full_text"])
            metadatas.append(metadata)
            ids.append(page_data["page_id"])
            
            # Add to vector DB
            # self.collection.add(
            #     documents=[page_data["full_text"]],
            #     metadatas=[metadata],
            #     ids=[page_data["page_id"]]
            # )

            #self.logger.debug(f"Added page {page_data['page_number']} to vector DB with ID {page_data['page_id']}")

        # Add all pages in a single batch operation
        self.vector_adapter.add(
            documents=documents,
            metadatas=metadatas,
            ids=ids
        )
        print(f"Added {len(page_data_list)} pages to vector DB")
        self.logger.debug(f"Added {len(page_data_list)} pages to vector DB")
    
    def _store_in_sql_db(
        self, 
        file_path: str,
        document_id: str, 
        document_type: str, 
        page_data_list: List[Dict[str, Any]],
        is_knowledge_document: bool = False
    ):
        """
        Store document data in SQL database.
        
        Args:
            file_path: Original file path
            document_id: Document ID
            document_type: Document type
            page_data_list: List of page data dictionaries
        """
        try:
            sql_conn = get_db_connection()
            if not sql_conn:
                print('SQL connection not available, skipping database storage')
                self.logger.warning("SQL connection not available, skipping database storage")
                return
            cursor = sql_conn.cursor()

            cursor.execute("EXEC tenant.sp_setTenantContext ?", os.getenv('API_KEY'))
            
            # Insert document record
            cursor.execute("""
                INSERT INTO Documents (
                    document_id, filename, original_path, document_type, 
                    page_count, processed_at, is_knowledge_document
                ) VALUES (?, ?, ?, ?, ?, ?, ?)
            """, (
                document_id,
                os.path.basename(file_path),
                file_path,
                document_type,
                len(page_data_list),
                datetime.now(),
                1 if normalize_boolean(is_knowledge_document) else 0
            ))
            
            # Insert page records and fields
            for page_data in page_data_list:
                # Insert page record
                cursor.execute("""
                    INSERT INTO DocumentPages (
                        page_id, document_id, page_number, full_text, vector_id
                    ) VALUES (?, ?, ?, ?, ?)
                """, (
                    page_data["page_id"],
                    document_id,
                    page_data["page_number"],
                    page_data["full_text"],
                    page_data["page_id"]  # Using same ID for vector and page
                ))
                
                # Insert extracted fields (flattening nested structures)
                self._insert_fields(cursor, page_data["page_id"], page_data["extracted_data"])

            # Commit the transaction
            sql_conn.commit()
            self.logger.info(f"Stored document {document_id} in SQL database")

            # NEW: Generate and store summaries if enabled
            if self.enable_summarization:
                try:
                    print('Generating document summaries...')
                    for page_data in page_data_list:
                        page_id = page_data["page_id"]
                        page_number = page_data["page_number"]
                        full_text = page_data["full_text"]
                        if self.summarizer and full_text.strip():
                            self._generate_and_store_page_summaries(
                                page_id=page_id,
                                document_id=document_id,
                                page_number=page_number,
                                page_content=full_text,
                                document_type=document_type
                            )
                except:
                    print('Something went wrong generating document summaries, skipping...')
                    self.logger.error('Something went wrong generating document summaries, skipping...')
        except Exception as e:
            self.logger.error(f"Error storing in SQL database: {str(e)}")
            sql_conn.rollback()
            raise

    def _generate_and_store_page_summaries(
        self, 
        page_id: str, 
        document_id: str, 
        page_number: int, 
        page_content: str, 
        document_type: str
    ):
        """
        Generate and store summaries for a page during document processing.
        This method is called automatically during document processing if summarization is enabled.
        """
        try:
            self.logger.info(f"Generating summaries for page {page_number} of document {document_id}")
            
            # Generate summaries for each configured type
            for summary_type in self.default_summary_types:
                try:
                    self.logger.info(f"Generating summaries of type {summary_type}...")
                    print('Summarizing page...')
                    # Generate summary
                    summary_data = self.summarizer.summarize_page(
                        page_content=page_content,
                        document_type=document_type,
                        summary_type=summary_type,
                        custom_instructions=None  # Could be made configurable
                    )
                    print('Summary Data:', summary_data)
                    
                    print('Saving page summary...')
                    # Save to database
                    success = self.summarizer.save_page_summary(
                        page_id=page_id,
                        document_id=document_id,
                        page_number=page_number,
                        summary_data=summary_data
                    )
                    
                    if success:
                        self.logger.info(f"Generated {summary_type} summary for page {page_id}")
                    else:
                        self.logger.error(f"Failed to save {summary_type} summary for page {page_id}")
                        
                except Exception as e:
                    self.logger.error(f"Error generating {summary_type} summary for page {page_id}: {str(e)}")
                    
        except Exception as e:
            self.logger.error(f"Error in summary generation process for page {page_id}: {str(e)}")
    
    def _insert_fields(self, cursor, page_id: str, data: Dict[str, Any], path: str = ""):
        """
        Recursively insert extracted fields into the database.
        
        Args:
            cursor: Database cursor
            page_id: Page ID
            data: Extracted data dictionary
            path: Current JSON path (for nested fields)
        """
        for key, value in data.items():
            current_path = f"{path}.{key}" if path else key
            
            if isinstance(value, dict):
                # Recursively process nested dictionaries
                self._insert_fields(cursor, page_id, value, current_path)
            elif isinstance(value, list) and value and isinstance(value[0], dict):
                # For lists of objects, store the list index in the path
                for i, item in enumerate(value):
                    if isinstance(item, dict):
                        self._insert_fields(cursor, page_id, item, f"{current_path}[{i}]")
            else:
                # Store scalar values (converting to string if needed)
                str_value = json.dumps(value) if not isinstance(value, str) else value
                cursor.execute("""
                    INSERT INTO DocumentFields (
                        page_id, field_name, field_value, field_path
                    ) VALUES (?, ?, ?, ?)
                """, (
                    page_id,
                    key,
                    str_value,
                    current_path
                ))

        # After processing all fields, ensure document_type exists (only at root level)
        if not path:  # Only do this at the root level to avoid repetition
            # Check if document_type field was already extracted and inserted
            cursor.execute("""
                SELECT COUNT(*) 
                FROM DocumentFields 
                WHERE page_id = ? AND field_name = 'document_type'
            """, (page_id,))
            
            document_type_exists = cursor.fetchone()[0] > 0
            
            # If document_type field doesn't exist, add it from the Documents table
            if not document_type_exists:
                cursor.execute("""
                    SELECT d.document_type 
                    FROM Documents d
                    JOIN DocumentPages dp ON d.document_id = dp.document_id
                    WHERE dp.page_id = ?
                """, (page_id,))
                
                result = cursor.fetchone()
                if result:
                    document_type = result[0]
                    cursor.execute("""
                        INSERT INTO DocumentFields (
                            page_id, field_name, field_value, field_path
                        ) VALUES (?, ?, ?, ?)
                    """, (
                        page_id,
                        'document_type',
                        document_type,
                        'document_type'
                    ))
    
    def process_directory(
        self, 
        directory_path: str, 
        archive_dir: str = "processed_archive",
        file_pattern: str = "*.pdf",
        recursive: bool = False,
        document_type: str = None,
        execution_id: str = None
    ) -> pd.DataFrame:
        """
        Process all matching files in a directory and archive them.
        
        Args:
            directory_path: Path to directory containing documents
            archive_dir: Subdirectory to move processed files to
            file_pattern: Pattern to match files (e.g., "*.pdf")
            recursive: Whether to process subdirectories recursively
            
        Returns:
            DataFrame summarizing the processed documents
        """
        import glob
        
        # Create the archive directory if it doesn't exist
        archive_path = os.path.join(directory_path, archive_dir)
        os.makedirs(archive_path, exist_ok=True)
        
        # Find all matching files
        pattern = os.path.join(directory_path, "**" if recursive else "", file_pattern)
        file_paths = glob.glob(pattern, recursive=recursive)

        self.logger.info(f"Processing file paths (before filter): {file_paths}")

        # Normalize the archive path for proper comparison
        archive_path_normalized = os.path.normpath(archive_path)

        # Filter out files that are actually IN the archive directory
        # Not just files that have the archive folder name in their path
        file_paths = [
            p for p in file_paths 
            if not os.path.normpath(os.path.dirname(p)).startswith(archive_path_normalized)
        ]

        self.logger.info(f"Processing file paths (after filter): {file_paths}")
        
        # Filter out files in the archive directory
        #file_paths = [p for p in file_paths if archive_dir not in p]
        
        results = []

        self.logger.info(f"directory_path (input): {directory_path}")
        self.logger.info(f"archive_dir (input): {archive_dir}")
        self.logger.info(f"archive_path: {archive_path}")
        self.logger.info(f"file_pattern (input): {file_pattern}")
        self.logger.info(f"pattern: {pattern}")
        self.logger.info(f"document_type: {document_type}")
        self.logger.info(f"Processing file paths: {file_paths}")
        
        for file_path in file_paths:
            try:
                # Process the file
                self.logger.info(f"Processing file: {file_path}")
                if document_type is not None:
                    result = self.process_document(file_path, document_type=document_type, force_ai_extraction=True, execution_id=execution_id)
                else:
                    result = self.process_document(file_path, force_ai_extraction=True, execution_id=execution_id)
                
                # Create summary record
                summary = {
                    "filename": result["filename"],
                    "document_id": result["document_id"],
                    "document_type": result["document_type"],
                    "page_count": result["page_count"],
                    "processed_at": result["processed_at"],
                    "status": "success",
                    "original_path": file_path,
                    "archived_path": archive_dir
                }

                print('Processed file:', result)
                print(86 * '-')
                print('Summary:', summary)
                
                print('Archiving file...', result["document_id"])
                # Archive the file
                archive_file_path = os.path.join(archive_path, os.path.basename(file_path))
                shutil.move(file_path, archive_file_path)
                summary["archived_path"] = archive_file_path
                
                # Update archive path in database if SQL connection exists
                sql_conn = get_db_connection()
                if sql_conn:
                    print('Updating archive path...', result["document_id"], archive_file_path)
                    cursor = sql_conn.cursor()
                    cursor.execute("EXEC tenant.sp_setTenantContext ?", os.getenv('API_KEY'))
                    cursor.execute(
                        "UPDATE Documents SET archived_path = ? WHERE document_id = ?", 
                        (archive_file_path, result["document_id"])
                    )
                    sql_conn.commit()

                    cursor.execute(
                        "UPDATE DocumentJobFileDetails SET ArchivedPath = ? WHERE FileDetailID = ?", 
                        (archive_file_path, result["file_detail_id"])
                    )
                    sql_conn.commit()
                
                results.append(summary)
                self.logger.info(f"Successfully processed and archived: {file_path}")
                print(f"Successfully processed and archived: {file_path}")
                
            except Exception as e:
                self.logger.error(f"Error processing {file_path}: {str(e)}")
                results.append({
                    "filename": os.path.basename(file_path),
                    "document_id": None,
                    "document_type": None,
                    "page_count": 0,
                    "processed_at": datetime.now().isoformat(),
                    "status": f"error: {str(e)}",
                    "original_path": file_path,
                    "archived_path": None
                })
        
        # Convert to DataFrame
        df = pd.DataFrame(results)
        return df
    
    def search_documents(
        self, 
        query: str, 
        document_type: Optional[str] = None,
        filters: Optional[Dict[str, Any]] = None,
        n_results: int = 5,
        min_score: float = 0.0
    ) -> List[Dict[str, Any]]:
        """
        Search for documents using vector similarity and metadata filters.
        
        Args:
            query: Text query to search for
            document_type: Optional document type to filter by
            filters: Optional additional metadata filters
            n_results: Maximum number of results to return
            min_score: Minimum relevance score (0-1) to include in results
            
        Returns:
            List of search results with references to original documents
        """
        # Set up where clause for filtering
        where_clause = {}
        
        # ChromaDB expects the where clause to be in a specific format for multiple conditions
        # We need to use $and operator for multiple conditions
        conditions = []
        
        # Add document_type filter if provided
        if document_type:
            conditions.append({"document_type": document_type})
        
        # Add other filters if provided
        if filters:
            for key, value in filters.items():
                conditions.append({key: value})
        
        # Build the final where clause
        if len(conditions) > 1:
            # Multiple conditions use $and operator
            where_clause = {"$and": conditions}
        elif len(conditions) == 1:
            # Single condition can be used directly
            where_clause = conditions[0]
        
        # print(86 * '-')
        # print(query)
        # print(where_clause)
        # print(86 * '-')

        # Perform the search
        # results = self.collection.query(
        #     query_texts=[query],
        #     where=where_clause if where_clause else None,
        #     #where_document={"$contains":query},
        #     n_results=n_results
        # )
        # Perform the search using the adapter
        try:
            results = self.vector_adapter.query(
                query_texts=[query],
                where=where_clause if where_clause else None,
                n_results=n_results,
                min_score=min_score
            )
        except Exception as e:
            self.logger.error(f"Error searching documents: {str(e)}")
            # Return empty results on error
            return []
        
        # Process and enhance the results
        processed_results = []
        
        for i, (doc_id, document, metadata, distance) in enumerate(zip(
                results['ids'][0],
                results['documents'][0], 
                results['metadatas'][0],
                results['distances'][0]
        )):
            # Calculate relevance score
            #relevance_score = math.exp(-float(distance))   #  For Euclidean
            #relevance_score = 1 - distance
            relevance_score = distance

            print('=============================================')
            print('distance', distance)
            print('relevance_score', relevance_score)
            print('min_score', min_score)
            print('=============================================')
            
            # Skip low-relevance results
            if relevance_score > min_score:
                continue
                
            # Get additional data from SQL if available
            additional_data = self._get_additional_data(doc_id) if get_db_connection() else {}
            
            # Create a result entry
            result = {
                "result_position": i + 1,
                "relevance_score": 1 - relevance_score,
                "page_id": doc_id,
                "document_id": metadata.get("document_id", ""),
                "filename": metadata.get("filename", ""),
                "page_number": metadata.get("page_number", 0),
                "document_type": metadata.get("document_type", ""),
                "snippet": self._create_snippet(document, query, max_length=200),
                "metadata": metadata,
                "extracted_fields": additional_data
            }
            
            processed_results.append(result)
            
        return processed_results
    
    def _get_additional_data(self, page_id: str) -> Dict[str, Any]:
        """
        Get additional data from SQL database for a document page.
        
        Args:
            page_id: Page ID
            
        Returns:
            Dictionary of field data
        """
        result = {}
        
        try:
            sql_conn = get_db_connection()
            cursor = sql_conn.cursor()
            
            # Query fields for this page
            cursor.execute(
                "SELECT field_name, field_value, field_path FROM DocumentFields WHERE page_id = ?", 
                (page_id,)
            )
            
            rows = cursor.fetchall()
            
            # Organize fields by path
            for field_name, field_value, field_path in rows:
                # Handle nested paths
                if '.' in field_path:
                    # Build nested structure
                    parts = field_path.split('.')
                    current = result
                    for part in parts[:-1]:
                        # Handle array indices in path
                        if '[' in part and ']' in part:
                            base_name = part.split('[')[0]
                            index = int(part.split('[')[1].split(']')[0])
                            if base_name not in current:
                                current[base_name] = []
                            # Ensure list is long enough
                            while len(current[base_name]) <= index:
                                current[base_name].append({})
                            current = current[base_name][index]
                        else:
                            if part not in current:
                                current[part] = {}
                            current = current[part]
                    current[parts[-1]] = field_value
                else:
                    # Add top-level field
                    result[field_name] = field_value
                    
        except Exception as e:
            self.logger.error(f"Error retrieving additional data: {str(e)}")
            
        return result
    
    def _create_snippet(self, text: str, query: str, max_length: int = 200) -> str:
        """Create a relevant text snippet containing the query terms"""
        # Find position of query terms
        query_terms = query.lower().split()
        text_lower = text.lower()
        
        # Find the best position to start the snippet
        best_pos = 0
        max_term_count = 0
        
        for i in range(len(text) - max_length):
            window = text_lower[i:i+max_length]
            term_count = sum(1 for term in query_terms if term in window)
            
            if term_count > max_term_count:
                max_term_count = term_count
                best_pos = i
        
        # Extract and clean the snippet
        end_pos = min(best_pos + max_length, len(text))
        snippet = text[best_pos:end_pos].strip()
        
        # Add ellipsis if we're not at the beginning/end
        if best_pos > 0:
            snippet = "..." + snippet
        if end_pos < len(text):
            snippet = snippet + "..."
            
        return snippet
    
    def export_document_data(self, document_id: str, format: str = 'json') -> Union[str, pd.DataFrame]:
        """
        Export all data for a document in the specified format.
        
        Args:
            document_id: Document ID
            format: Output format ('json', 'csv', or 'dataframe')
            
        Returns:
            String (JSON/CSV) or DataFrame containing document data
        """
        sql_conn = get_db_connection()
        if not sql_conn:
            raise ValueError("SQL connection is required for data export")
            
        try:
            cursor = sql_conn.cursor()
            
            # Get document info
            cursor.execute(
                "SELECT * FROM Documents WHERE document_id = ?", 
                (document_id,)
            )
            doc_row = cursor.fetchone()
            
            if not doc_row:
                raise ValueError(f"Document ID {document_id} not found")
                
            # Convert row to dict (PYODBC doesn't have column names in row)
            columns = [column[0] for column in cursor.description]
            doc_data = dict(zip(columns, doc_row))
            
            # Get pages
            cursor.execute(
                "SELECT * FROM DocumentPages WHERE document_id = ? ORDER BY page_number", 
                (document_id,)
            )
            pages = []
            for row in cursor.fetchall():
                columns = [column[0] for column in cursor.description]
                page_data = dict(zip(columns, row))
                
                # Get fields for this page
                cursor.execute(
                    "SELECT field_name, field_value, field_path FROM DocumentFields WHERE page_id = ?", 
                    (page_data['page_id'],)
                )
                
                fields = {}
                for field_name, field_value, field_path in cursor.fetchall():
                    fields[field_path] = field_value
                    
                page_data['extracted_fields'] = fields
                pages.append(page_data)
                
            # Combine all data
            full_data = {
                **doc_data,
                'pages': pages
            }
            
            # Return in requested format
            if format == 'json':
                return json.dumps(full_data, indent=2)
            elif format in ('csv', 'dataframe'):
                # Flatten the data for CSV/DataFrame
                rows = []
                for page in pages:
                    base_row = {
                        'document_id': doc_data['document_id'],
                        'filename': doc_data['filename'],
                        'document_type': doc_data['document_type'],
                        'page_number': page['page_number'],
                        'page_id': page['page_id']
                    }
                    
                    # Add fields as columns
                    for field_path, value in page['extracted_fields'].items():
                        # Sanitize column name
                        col_name = field_path.replace('[', '_').replace(']', '').replace('.', '_')
                        base_row[col_name] = value
                        
                    rows.append(base_row)
                    
                df = pd.DataFrame(rows)
                
                if format == 'csv':
                    return df.to_csv(index=False)
                return df
                
            else:
                raise ValueError(f"Unsupported format: {format}")
                
        except Exception as e:
            self.logger.error(f"Error exporting document data: {str(e)}")
            raise
    
    def close(self):
        """Close all connections"""
        if self.sql_conn:
            self.sql_conn.close()
            self.sql_conn = None
            self.logger.info("SQL connection closed")

        # Close vector adapter connections
        if hasattr(self, 'vector_adapter'):
            self.vector_adapter.close()
            self.logger.info("Vector adapter connections closed")

    def debug_database_contents(self):
        """Print information about what's in the database"""
        # Get all documents in the collection
        #results = self.collection.get()
         # Get vector database stats through the adapter
        stats = self.vector_adapter.get_collection_stats() if hasattr(self.vector_adapter, 'get_collection_stats') else {}

        # print(f"Total documents in collection: {len(results['ids'])}")
        # print(f"Available metadata fields: {set().union(*[set(m.keys()) for m in results['metadatas']])}")

        print(f"Total documents in collection: {stats.get('count', 'unknown')}")
        print(f"Available metadata fields: {stats.get('metadata_fields', [])}")
        
        # Print sample of documents
        # for i in range(min(5, len(results['ids']))):
        #     print(f"\nDocument {i+1}:")
        #     print(f"ID: {results['ids'][i]}")
        #     print(f"Metadata: {results['metadatas'][i]}")

    ##########################################
    # IN TESTING
    ##########################################
    def _process_pdf_DEPRECATED(self, file_path, document_type, use_batch_processing=False, batch_size=3):
        """Process PDF documents using the existing PDF handler"""
        pdf_handler = MultiPagePDFHandler(
            self.anthropic_client, 
            self.logger,
            anthropic_config=self._anthropic_config,
            anthropic_proxy_client=self.anthropic_proxy_client
        )
        return pdf_handler.extract_from_pdf(
            file_path,
            document_type=document_type,
            use_batch_processing=use_batch_processing,
            batch_size=batch_size,
            include_page_numbers=True
        )
    
    def _process_pdf(self, file_path, document_type, use_batch_processing=False, batch_size=3):
        """
        Process PDF documents using optimized extraction strategy.
        
        Strategy:
        1. Analyze PDF to determine if it's text-based or scanned/image-based
        2. For text-based PDFs: Use fast PyMuPDF extraction (no API calls)
        3. For scanned/mixed PDFs: Use AI extraction (Claude Vision)
        4. If fast extraction produces poor results: Fall back to AI extraction
        
        This can significantly reduce API costs and processing time for text-based PDFs
        while still handling scanned documents correctly.
        
        Args:
            file_path: Path to the PDF file
            document_type: Type of document for extraction hints
            use_batch_processing: Whether to batch AI requests (when using AI)
            batch_size: Number of pages per AI batch
            
        Returns:
            List of page dictionaries with extracted text
        """
        # Check if fast extraction is available and enabled
        use_fast = FAST_PDF_AVAILABLE and getattr(cfg, 'DOC_USE_FAST_PDF_EXTRACTION', True)
        
        if use_fast:
            try:
                min_chars = getattr(cfg, 'DOC_FAST_PDF_MIN_CHARS_PER_PAGE', 100)
                always_try_fast = getattr(cfg, 'DOC_FAST_PDF_ALWAYS_TRY_FAST', True)
                
                self.logger.info(f"Using FastPDFExtractor for: {file_path}")
                
                fast_extractor = FastPDFExtractor(
                    anthropic_client=self.anthropic_client,
                    logger=self.logger,
                    anthropic_config=self._anthropic_config,
                    anthropic_proxy_client=self.anthropic_proxy_client,
                    min_chars_per_page=min_chars,
                    always_try_fast=always_try_fast
                )
                
                # Use extract_with_details for logging purposes
                result = fast_extractor.extract_with_details(
                    file_path=file_path,
                    document_type=document_type,
                    use_batch_processing=use_batch_processing,
                    batch_size=batch_size,
                    include_page_numbers=True
                )
                
                # Log extraction method used
                self.logger.info(
                    f"PDF extraction complete - Method: {result.method_used.value}, "
                    f"PDF Type: {result.pdf_analysis.pdf_type.value}, "
                    f"Pages: {len(result.pages)}, "
                    f"Fast attempted: {result.fast_extraction_attempted}, "
                    f"Fast success: {result.fast_extraction_success}"
                )
                
                if result.fallback_reason:
                    self.logger.info(f"Fallback reason: {result.fallback_reason}")
                
                return result.pages
                
            except Exception as e:
                self.logger.warning(f"FastPDFExtractor failed: {e}, falling back to AI-only extraction")
        
        # Fall back to original AI-only method
        self.logger.info(f"Using AI extraction for: {file_path}")
        pdf_handler = MultiPagePDFHandler(
            self.anthropic_client, 
            self.logger,
            anthropic_config=self._anthropic_config,
            anthropic_proxy_client=self.anthropic_proxy_client
        )
        return pdf_handler.extract_from_pdf(
            file_path,
            document_type=document_type,
            use_batch_processing=use_batch_processing,
            batch_size=batch_size,
            include_page_numbers=True
        )

    def _process_image(self, file_path, document_type):
        """Process an image file using Claude Vision"""
        self.logger.info(f"Processing image: {file_path}")
        
        try:
            # Read the image file
            with open(file_path, "rb") as f:
                file_bytes = f.read()
                
            # Create a base64 representation
            base64_image = base64.b64encode(file_bytes).decode("utf-8")
            
            # System prompt for image processing
            system_prompt = (
                f"You are an expert at extracting text and data from {document_type} images. "
                "Extract ALL text from the image maintaining its structure and formatting. "
                "Your task is to convert the image content to plain text that includes ALL information visible in the image. "
                "This includes all text elements, tables, form fields, and other textual information."
            )
            
            # User message for image extraction
            user_text = (
                f"Extract the complete text from this {document_type} image. "
                "Include ALL visible text and maintain the structure of tables and forms. "
                "Return ONLY the extracted text without any additional commentary."
            )
            
            # Call Claude Vision API
            if self._anthropic_config['use_direct_api']:
                print('Bypassing document proxy...')
                response = anthropic_messages_create(
                    client=self.anthropic_client,
                    model=cfg.ANTHROPIC_MODEL,
                    max_tokens=int(cfg.ANTHROPIC_MAX_TOKENS),
                    system=system_prompt,
                    messages=[
                        {
                            "role": "user",
                            "content": [
                                {
                                    "type": "image",
                                    "source": {
                                        "type": "base64",
                                        "media_type": self._get_image_media_type(file_path),
                                        "data": base64_image
                                    }
                                },
                                {
                                    "type": "text", 
                                    "text": user_text
                                }
                            ]
                        }
                    ]
                )
                
                # Extract the text from Claude's response
                extracted_text = response.content[0].text.strip()
            else:
                print('Calling document proxy...')
                client = AnthropicProxyClient()
                response = client.messages_with_document(
                    file_path=file_path,
                    user_text=user_text,
                    model=cfg.ANTHROPIC_MODEL,
                    max_tokens=int(cfg.ANTHROPIC_MAX_TOKENS),
                    system=system_prompt
                )
                
                # Extract the text from response
                extracted_text = response['content'][0]['text'].strip()
                
            # Apply API throttling if configured
            if cfg.ANTHROPIC_API_THROTTLE_CALLS:
                self.logger.info('API throttling active, waiting before submitting next request...')
                time.sleep(int(cfg.ANTHROPIC_API_THROTTLE_DELAY))
                
            # Return as a single page
            return [{
                "page_number": 1,
                "text": extracted_text
            }]
            
        except Exception as e:
            self.logger.error(f"Error processing image: {str(e)}")
            return [{
                "page_number": 1,
                "text": f"Error processing image: {str(e)}"
            }]

    def _get_image_media_type(self, file_path):
        """Determine the media type for an image file"""
        ext = os.path.splitext(file_path)[1].lower()
        media_types = {
            '.jpg': 'image/jpeg',
            '.jpeg': 'image/jpeg',
            '.png': 'image/png',
            '.gif': 'image/gif',
            '.webp': 'image/webp',
            '.bmp': 'image/bmp',
            '.tiff': 'image/tiff',
            '.tif': 'image/tiff'
        }
        return media_types.get(ext, 'image/jpeg')  # Default to JPEG if unknown    
    def _get_file_and_document_type(self, file_name):
        filename = str(file_name).lower()
        # If content type wasn't provided or is generic, try to determine from filename
        if filename:
            if filename.endswith('.pdf'):
                file_type = 'application/pdf'
                doc_type = 'document'
            elif filename.endswith('.docx'):
                file_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                doc_type = 'text'
            elif filename.endswith('.doc'):
                file_type = 'application/msword'
                doc_type = 'text'
            elif filename.endswith('.pptx'):
                file_type = 'application/vnd.openxmlformats-officedocument.presentationml.presentation'
                doc_type = 'text'
            elif filename.endswith('.xlsx'):
                file_type = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
                doc_type = 'text'
            elif filename.endswith('.txt'):
                file_type = 'text/plain'
                doc_type = 'text'
            elif filename.endswith('.csv'):
                file_type = 'text/csv'
                doc_type = 'text'
            elif filename.endswith('.jpg') or filename.endswith('.jpeg'):
                file_type = 'image/jpeg'
                doc_type = 'image'
            elif filename.endswith('.png'):
                file_type = 'image/png'
                doc_type = 'image'
            else:
                # Default to binary if we can't determine
                file_type = 'application/octet-stream'
                doc_type = 'document'

        return file_type, doc_type
                

    ##########################################
    # REALLY IN TESTING
    ##########################################
    def _process_word_document(self, file_path, document_type):
        """Process Word documents"""
        self.logger.info(f"Processing Word document: {file_path}")
        
        try:
            # Try to use python-docx for .docx files
            if file_path.lower().endswith('.docx'):
                import docx
                doc = docx.Document(file_path)
                
                # Extract text from paragraphs
                pages = []
                full_text = ""
                
                for para in doc.paragraphs:
                    if para.text.strip():
                        full_text += para.text.strip() + "\n\n"
                
                # Extract text from tables
                for table in doc.tables:
                    table_text = ""
                    for row in table.rows:
                        row_text = " | ".join([cell.text.strip() for cell in row.cells])
                        table_text += row_text + "\n"
                    full_text += "\n" + table_text + "\n\n"
                    
                # Handle as a single page document
                pages.append({
                    "page_number": 1,
                    "text": full_text.strip()
                })
                
                return pages
                
            # For .doc files or if python-docx fails, use Claude Vision
            return self._process_generic_file(file_path, document_type)
            
        except Exception as e:
            self.logger.error(f"Error processing Word document: {str(e)}")
            # Fallback to generic handling
            return self._process_generic_file(file_path, document_type)


    # =========================================================================
    # Excel Processing Helpers
    # =========================================================================

    def _excel_read_sample_rows(self, file_path, sheet_name, max_rows=20):
        """
        Read a sample of rows from an Excel sheet using openpyxl for structure analysis.

        Returns dict with keys:
          - 'rows': list of lists (raw cell values, None for empty)
          - 'merged_ranges': list of str (e.g., ["A1:D1", "B3:B5"])
          - 'is_visible': bool
          - 'total_rows': int (approximate, from sheet.max_row)
          - 'total_cols': int (from sheet.max_column)
        """
        import openpyxl

        wb = openpyxl.load_workbook(file_path, read_only=False, data_only=True)
        try:
            sheet = wb[sheet_name]

            is_visible = sheet.sheet_state == 'visible'
            merged_ranges = [str(r) for r in sheet.merged_cells.ranges]
            total_rows = sheet.max_row or 0
            total_cols = sheet.max_column or 0

            rows = []
            for row_idx, row in enumerate(sheet.iter_rows(
                min_row=1, max_row=min(max_rows, total_rows), values_only=True
            )):
                rows.append(list(row))

            wb.close()
            return {
                'rows': rows,
                'merged_ranges': merged_ranges,
                'is_visible': is_visible,
                'total_rows': total_rows,
                'total_cols': total_cols
            }
        except Exception:
            wb.close()
            raise

    def _excel_format_sample_as_markdown(self, rows, max_cols=0):
        """
        Format raw sample rows as a row-numbered markdown-style text table for AI analysis.
        Each row is prefixed with its 0-based index (Row 0, Row 1, etc.) so the AI
        can reference row numbers unambiguously.
        """
        if not rows:
            return "(empty sheet)"

        if not max_cols:
            max_cols = max(len(r) for r in rows) if rows else 0

        lines = []
        for row_idx, row in enumerate(rows):
            padded = list(row) + [None] * (max_cols - len(row))
            cells = []
            for cell in padded:
                if cell is None:
                    cells.append("")
                elif isinstance(cell, (int, float)):
                    cells.append(str(cell))
                else:
                    cells.append(str(cell).replace("|", "\\|").replace("\n", " ").strip())
            line = f"Row {row_idx}: | " + " | ".join(cells) + " |"
            lines.append(line)

        return "\n".join(lines)

    def _excel_detect_structure(self, sample_markdown, sheet_name, merged_ranges, total_rows):
        """
        Use AI to analyze the structure of an Excel sheet sample.

        Sends the first N rows (as text, not images) to Claude and asks it to identify
        the real header row, title/metadata rows to skip, and whether the sheet is empty.

        Returns dict with keys:
          - 'header_row': int (0-indexed row number of the header, or -1 if none)
          - 'skip_rows': list of int (0-indexed rows to skip)
          - 'is_empty_or_irrelevant': bool
          - 'notes': str (brief description of what the AI observed)

        Falls back to default assumptions on any error.
        """
        default_result = {
            'header_row': 0,
            'skip_rows': [],
            'is_empty_or_irrelevant': False,
            'notes': 'Default: assumed row 0 is header (AI detection skipped or failed)'
        }

        if not sample_markdown or sample_markdown == "(empty sheet)":
            return {
                'header_row': -1,
                'skip_rows': [],
                'is_empty_or_irrelevant': True,
                'notes': 'Sheet appears empty'
            }

        system_prompt = (
            "You are a spreadsheet structure analyzer. You will receive the first rows of an Excel sheet "
            "formatted as a text table. Each row is labeled with its 0-based index (Row 0, Row 1, etc.).\n\n"
            "Analyze the structure and respond with ONLY a JSON object (no markdown fencing, no explanation) "
            "with these exact keys:\n"
            '- "header_row": integer, the 0-based row index that contains the column headers. '
            "Use -1 if there is no clear header row.\n"
            '- "skip_rows": array of integers, 0-based row indices of title/metadata/blank rows '
            "that appear BEFORE the header row and should be skipped during data extraction.\n"
            '- "is_empty_or_irrelevant": boolean, true if the sheet contains no meaningful tabular data.\n'
            '- "notes": string, brief 1-sentence description of what you observed about the sheet structure.'
        )

        merged_info = ""
        if merged_ranges:
            merged_info = f"\n\nMerged cell regions in this sheet: {', '.join(merged_ranges)}"

        user_text = (
            f"Sheet name: \"{sheet_name}\"\n"
            f"Total rows in sheet: {total_rows}\n"
            f"{merged_info}\n\n"
            f"First rows:\n{sample_markdown}"
        )

        try:
            if self._anthropic_config['use_direct_api']:
                response = anthropic_messages_create(
                    client=self.anthropic_client,
                    model=cfg.ANTHROPIC_MODEL,
                    max_tokens=int(getattr(cfg, 'EXCEL_STRUCTURE_MAX_TOKENS', 1024)),
                    system=system_prompt,
                    messages=[{
                        "role": "user",
                        "content": [{"type": "text", "text": user_text}]
                    }]
                )
                response_text = response.content[0].text.strip()
            else:
                response = self.anthropic_proxy_client.messages_create(
                    model=cfg.ANTHROPIC_MODEL,
                    max_tokens=int(getattr(cfg, 'EXCEL_STRUCTURE_MAX_TOKENS', 1024)),
                    system=system_prompt,
                    messages=[{
                        "role": "user",
                        "content": [{"type": "text", "text": user_text}]
                    }]
                )
                if "error" in response:
                    self.logger.warning(f"AI structure detection proxy error: {response['error']}")
                    return default_result
                response_text = response['content'][0]['text'].strip()

            # Apply API throttling if configured
            if cfg.ANTHROPIC_API_THROTTLE_CALLS:
                self.logger.info('API throttling active, waiting before submitting next request...')
                time.sleep(int(cfg.ANTHROPIC_API_THROTTLE_DELAY))

            # Parse JSON response - strip markdown fencing if present
            if response_text.startswith('```'):
                response_text = response_text.split('\n', 1)[1] if '\n' in response_text else response_text[3:]
                if response_text.endswith('```'):
                    response_text = response_text[:-3]
                response_text = response_text.strip()

            result = json.loads(response_text)

            # Validate required keys
            required_keys = ['header_row', 'skip_rows', 'is_empty_or_irrelevant']
            if not all(k in result for k in required_keys):
                self.logger.warning(f"AI structure response missing keys: {result}")
                return default_result

            # Validate and coerce types
            if not isinstance(result['header_row'], int):
                result['header_row'] = int(result['header_row'])
            if not isinstance(result['skip_rows'], list):
                result['skip_rows'] = []
            if not isinstance(result['is_empty_or_irrelevant'], bool):
                result['is_empty_or_irrelevant'] = False
            if 'notes' not in result:
                result['notes'] = ''

            self.logger.info(f"AI structure detection for sheet '{sheet_name}': header_row={result['header_row']}, "
                           f"skip_rows={result['skip_rows']}, is_empty={result['is_empty_or_irrelevant']}, "
                           f"notes={result.get('notes', '')}")
            return result

        except json.JSONDecodeError as e:
            self.logger.warning(f"Failed to parse AI structure response as JSON: {e}")
            return default_result
        except Exception as e:
            self.logger.warning(f"AI structure detection failed for sheet '{sheet_name}': {e}")
            return default_result

    def _excel_build_markdown_table(self, df, max_rows=5000):
        """
        Convert a pandas DataFrame to a clean markdown table string.

        Handles NaN replacement, pipe escaping, newline flattening, and row truncation.
        Adapted from extract_xlsx_text() pattern in attachment_text_extractor.py.
        """
        df = df.fillna("")

        # Build column headers
        headers = [str(col).replace("|", "\\|").replace("\n", " ") for col in df.columns]

        lines = []
        # Header row
        lines.append("| " + " | ".join(headers) + " |")
        # Separator
        lines.append("| " + " | ".join(["---"] * len(headers)) + " |")

        # Data rows
        row_count = 0
        for _, row in df.iterrows():
            if row_count >= max_rows:
                break
            cells = []
            for val in row:
                if pd.isna(val) or val == "":
                    cells.append("")
                elif isinstance(val, (int, float)):
                    cells.append(str(val))
                else:
                    cells.append(str(val).replace("|", "\\|").replace("\n", " "))
            lines.append("| " + " | ".join(cells) + " |")
            row_count += 1

        total_rows = len(df)
        if total_rows > max_rows:
            lines.append(f"\n... (truncated, showing {max_rows} of {total_rows} rows)")

        return "\n".join(lines)

    # =========================================================================
    # Excel Main Processor
    # =========================================================================

    def _process_excel(self, file_path, document_type):
        """Process Excel spreadsheets with AI-guided structure detection"""
        self.logger.info(f"Processing Excel document: {file_path}")

        try:
            import openpyxl

            max_rows = int(getattr(cfg, 'EXCEL_MAX_ROWS_PER_SHEET', 5000))
            sample_rows_count = int(getattr(cfg, 'EXCEL_STRUCTURE_SAMPLE_ROWS', 20))
            use_ai_detection = getattr(cfg, 'EXCEL_AI_STRUCTURE_DETECTION', True)

            # Get sheet names and visibility info using openpyxl
            wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
            all_sheets = wb.sheetnames
            sheet_states = {}
            for name in all_sheets:
                sheet_states[name] = wb[name].sheet_state
            wb.close()

            pages = []
            page_number = 1

            for sheet_name in all_sheets:
                # Skip hidden/very hidden sheets
                if sheet_states.get(sheet_name) != 'visible':
                    self.logger.info(f"Skipping hidden sheet: '{sheet_name}' (state: {sheet_states.get(sheet_name)})")
                    continue

                # --- Structure detection phase ---
                structure = None
                if use_ai_detection:
                    try:
                        sample_data = self._excel_read_sample_rows(
                            file_path, sheet_name, max_rows=sample_rows_count
                        )

                        if sample_data['rows']:
                            sample_md = self._excel_format_sample_as_markdown(
                                sample_data['rows'],
                                max_cols=sample_data['total_cols']
                            )
                            structure = self._excel_detect_structure(
                                sample_md, sheet_name,
                                sample_data['merged_ranges'],
                                sample_data['total_rows']
                            )
                    except Exception as e:
                        self.logger.warning(f"Structure detection failed for sheet '{sheet_name}': {e}")
                        structure = None

                # Apply defaults if no structure detected
                if structure is None:
                    structure = {
                        'header_row': 0,
                        'skip_rows': [],
                        'is_empty_or_irrelevant': False,
                        'notes': 'Using default settings (AI detection disabled or failed)'
                    }

                # Skip empty/irrelevant sheets
                if structure.get('is_empty_or_irrelevant', False):
                    self.logger.info(f"Skipping irrelevant sheet: '{sheet_name}' ({structure.get('notes', '')})")
                    continue

                # --- Guided extraction phase ---
                header_row = structure['header_row']
                skip_rows = structure.get('skip_rows', [])

                try:
                    if header_row < 0:
                        # No header row detected - read without headers
                        df = pd.read_excel(
                            file_path, sheet_name=sheet_name,
                            header=None, skiprows=skip_rows if skip_rows else None
                        )
                    else:
                        # Compute rows to skip so that the header ends up at position 0
                        rows_before_header = [r for r in range(header_row) if r not in skip_rows]
                        all_skip = sorted(set(skip_rows + rows_before_header))

                        df = pd.read_excel(
                            file_path, sheet_name=sheet_name,
                            header=0, skiprows=all_skip if all_skip else None
                        )

                    # Build markdown table output
                    sheet_text = f"## Sheet: {sheet_name}\n\n"

                    if structure.get('notes') and use_ai_detection:
                        sheet_text += f"_Structure: {structure['notes']}_\n\n"

                    sheet_text += self._excel_build_markdown_table(df, max_rows=max_rows)

                    pages.append({
                        "page_number": page_number,
                        "text": sheet_text
                    })
                    page_number += 1

                except Exception as e:
                    self.logger.warning(f"Guided extraction failed for sheet '{sheet_name}', "
                                      f"falling back to basic read: {e}")
                    # Fallback: basic read with no skip/header guidance
                    try:
                        df = pd.read_excel(file_path, sheet_name=sheet_name)
                        df = df.fillna("")
                        sheet_text = f"## Sheet: {sheet_name}\n\n"
                        sheet_text += self._excel_build_markdown_table(df, max_rows=max_rows)
                        pages.append({
                            "page_number": page_number,
                            "text": sheet_text
                        })
                        page_number += 1
                    except Exception as e2:
                        self.logger.error(f"Basic read also failed for sheet '{sheet_name}': {e2}")

            if not pages:
                # All sheets were hidden, empty, or failed
                pages.append({
                    "page_number": 1,
                    "text": "No visible or relevant sheets found in this Excel file."
                })

            return pages

        except Exception as e:
            self.logger.error(f"Error processing Excel document: {str(e)}")
            # Fallback to generic handling
            return self._process_generic_file(file_path, document_type)

    # def _process_powerpoint(self, file_path, document_type):
    #     """Process PowerPoint presentations"""
    #     self.logger.info(f"Processing PowerPoint document: {file_path}")
        
    #     try:
    #         from pptx import Presentation
            
    #         # Open the presentation
    #         prs = Presentation(file_path)
            
    #         pages = []
            
    #         # Process each slide as a separate "page"
    #         for i, slide in enumerate(prs.slides):
    #             slide_text = f"Slide {i+1}:\n\n"
                
    #             # Extract text from shapes
    #             for shape in slide.shapes:
    #                 if hasattr(shape, "text") and shape.text.strip():
    #                     slide_text += shape.text.strip() + "\n\n"
                
    #             pages.append({
    #                 "page_number": i + 1,
    #                 "text": slide_text.strip()
    #             })
                
    #         return pages
            
    #     except Exception as e:
    #         self.logger.error(f"Error processing PowerPoint document: {str(e)}")
    #         # Fallback to generic handling
    #         return self._process_generic_file(file_path, document_type)

    def _process_text_file(self, file_path, document_type):
        """Process text files (TXT, CSV, JSON, etc.) with chunking for large files."""
        self.logger.info(f"Processing text file: {file_path}")
        
        # Maximum characters per chunk/page (~4 chars per token, ~50K tokens safe for Claude)
        MAX_CHARS_PER_PAGE = 200_000  # ~50K tokens
        
        try:
            # Determine encoding
            import chardet
            with open(file_path, 'rb') as f:
                raw_data = f.read()
                result = chardet.detect(raw_data)
                encoding = result['encoding']
            
            # Read the file with detected encoding
            with open(file_path, 'r', encoding=encoding) as f:
                content = f.read()
            
            # Handle CSV files specially
            if file_path.lower().endswith('.csv'):
                import pandas as pd
                try:
                    df = pd.read_csv(file_path)
                    content = df.to_string(index=False)
                except:
                    # If pandas fails, keep the raw content
                    pass
                    
            # Handle JSON files specially
            if file_path.lower().endswith('.json'):
                import json
                try:
                    data = json.loads(content)
                    content = json.dumps(data, indent=2)
                except:
                    # If JSON parsing fails, keep the raw content
                    pass
            
            # Chunk large files into multiple pages
            if len(content) <= MAX_CHARS_PER_PAGE:
                return [{
                    "page_number": 1,
                    "text": content
                }]
            
            # Split into chunks at line boundaries for clean breaks
            self.logger.info(f"Large text file ({len(content):,} chars) — chunking into ~{MAX_CHARS_PER_PAGE:,} char pages")
            pages = []
            lines = content.split('\n')
            current_chunk = []
            current_size = 0
            page_num = 1
            
            for line in lines:
                line_len = len(line) + 1  # +1 for newline
                if current_size + line_len > MAX_CHARS_PER_PAGE and current_chunk:
                    pages.append({
                        "page_number": page_num,
                        "text": '\n'.join(current_chunk)
                    })
                    page_num += 1
                    current_chunk = []
                    current_size = 0
                current_chunk.append(line)
                current_size += line_len
            
            # Don't forget the last chunk
            if current_chunk:
                pages.append({
                    "page_number": page_num,
                    "text": '\n'.join(current_chunk)
                })
            
            self.logger.info(f"Split into {len(pages)} pages")
            return pages
            
        except Exception as e:
            self.logger.error(f"Error processing text file: {str(e)}")
            return [{
                "page_number": 1,
                "text": f"Error processing text file: {str(e)}"
            }]

    def _process_generic_file(self, file_path, document_type):
        """Process any file using Claude Vision as a fallback"""
        self.logger.info(f"Processing generic file with Claude Vision: {file_path}")
        
        try:
            # Let Claude process the file directly
            client = AnthropicProxyClient()
            response = client.messages_with_document(
                file_path=file_path,
                user_text=f"Extract all text content from this {document_type} file. Preserve the structure and formatting as much as possible. Return ONLY the extracted text without additional commentary.",
                model=cfg.ANTHROPIC_MODEL,
                max_tokens=int(cfg.ANTHROPIC_MAX_TOKENS),
                system=f"You are an expert at extracting text and data from documents. Your task is to extract ALL visible content from this {document_type} file and convert it to plain text."
            )
            
            # Extract the text from response
            if isinstance(response, dict) and 'content' in response:
                extracted_text = response['content'][0]['text'].strip()
            else:
                extracted_text = response.content[0].text.strip()
                
            # Return as a single page
            return [{
                "page_number": 1,
                "text": extracted_text
            }]
            
        except Exception as e:
            self.logger.error(f"Error processing generic file: {str(e)}")
            return [{
                "page_number": 1,
                "text": f"Error processing file: {str(e)}"
            }]
        
    def reprocess_document_vectors(
        self,
        document_id: str = None,
        document_type: str = None,
        batch_size: int = 100,
        force_recreate: bool = False
    ) -> Dict[str, Any]:
        """
        Reprocess document vectors from existing database text data.
        
        This function rebuilds the vector database from the text stored in DocumentPages
        without needing to reprocess the original documents. Useful for:
        - Vector database corruption recovery
        - Changing vector models/settings
        - Bulk vector database maintenance
        
        Args:
            document_id: Specific document ID to reprocess (optional)
            document_type: Filter by document type (optional)
            batch_size: Number of pages to process in each batch
            force_recreate: If True, delete existing vectors before recreating
            
        Returns:
            Dictionary with processing results and statistics
        """
        try:
            self.logger.info(f"Starting vector reprocessing - document_id: {document_id}, document_type: {document_type}")
            
            # Get database connection
            sql_conn = get_db_connection()
            if not sql_conn:
                raise Exception("SQL connection not available")
                
            cursor = sql_conn.cursor()
            cursor.execute("EXEC tenant.sp_setTenantContext ?", os.getenv('API_KEY'))
            
            # Build query to get pages that need reprocessing
            base_query = """
                SELECT 
                    dp.page_id,
                    dp.document_id,
                    dp.page_number,
                    dp.full_text,
                    d.filename,
                    d.document_type,
                    d.processed_at,
                    d.is_knowledge_document
                FROM DocumentPages dp
                INNER JOIN Documents d ON dp.document_id = d.document_id
                WHERE dp.full_text IS NOT NULL 
                AND d.is_knowledge_document = 0
            """
            
            params = []
            
            # Add filters if specified
            if document_id:
                base_query += " AND dp.document_id = ?"
                params.append(document_id)
                
            if document_type:
                base_query += " AND d.document_type = ?"
                params.append(document_type)
                
            base_query += " ORDER BY dp.document_id, dp.page_number"
            
            # Execute query to get total count first
            count_query = base_query.replace(" ORDER BY dp.document_id, dp.page_number", "")
            count_query = f"SELECT COUNT(*) FROM ({count_query}) as counted"
            cursor.execute(count_query, params)
            total_pages = cursor.fetchone()[0]
            
            if total_pages == 0:
                self.logger.info("No pages found to reprocess")
                return {
                    "status": "success",
                    "message": "No pages found to reprocess",
                    "pages_processed": 0,
                    "documents_processed": 0,
                    "errors": []
                }
            
            self.logger.info(f"Found {total_pages} pages to reprocess")
            
            # Get all page data
            cursor.execute(base_query, params)
            pages_data = cursor.fetchall()
            
            # Track statistics
            processed_pages = 0
            processed_documents = set()
            errors = []
            
            # If force_recreate is True OR we're doing selective reprocessing, delete existing vectors first
            if force_recreate or document_id or document_type:
                self.logger.info("Deleting existing vectors for selected documents")
                print("Deleting existing vectors for selected documents...")
                page_ids_to_delete = [row[0] for row in pages_data]  # page_id is first column
                
                if page_ids_to_delete:
                    try:
                        for page_id in page_ids_to_delete:
                            print(f"Deleting vector(s) for page: {page_id}")
                            # Delete from vector database using VectorEngineClient for each page_id
                            vector_client = VectorEngineClient()
                            
                            # Delete each page from vector database
                            response = vector_client.delete_document(page_id)

                            if response.get('status') != 'success':
                                self.logger.warning(f"Page {page_id}: {response.get('message')}")
                                print(f"Page {page_id}: {response.get('message')}")
                            else:
                                self.logger.info(f"Vector(s) deleted successfully for page: {page_id}")
                                print(f"Vector(s) deleted successfully for page: {page_id}")
                    except Exception as e:
                        self.logger.warning(f"Error deleting existing vectors: {str(e)}")
                        # Continue anyway - might be because vectors don't exist
            
                # Process pages in batches to match main pipeline
                for i in range(0, len(pages_data), batch_size):
                    batch_data = pages_data[i:i + batch_size]
                    
                    # Build page_data_list exactly like main pipeline
                    page_data_list = []
                    
                    self.logger.info(f"Processing batch {i//batch_size + 1}/{(len(pages_data) + batch_size - 1)//batch_size}")
                    print(f"Processing batch {i//batch_size + 1}/{(len(pages_data) + batch_size - 1)//batch_size}")
                    
                    for row in batch_data:
                        (page_id, doc_id, page_number, full_text, filename, 
                        doc_type, processed_at, is_knowledge_doc) = row
                        
                        try:
                            # Skip if no text content
                            if not full_text or full_text.strip() == "":
                                self.logger.warning(f"Skipping page {page_id} - no text content")
                                continue
                            
                            # Get extracted fields for this page from SQL
                            cursor.execute("""
                                SELECT field_name, field_value 
                                FROM DocumentFields 
                                WHERE page_id = ?
                            """, (page_id,))
                            
                            # Build extracted_data dict like main pipeline
                            extracted_data = {}
                            field_rows = cursor.fetchall()
                            for field_name, field_value in field_rows:
                                extracted_data[field_name] = field_value
                            
                            # Create page_data exactly like main pipeline
                            page_data = {
                                "document_id": doc_id,
                                "page_id": page_id,
                                "filename": filename,
                                "page_number": page_number,
                                "document_type": doc_type,
                                "extraction_timestamp": processed_at.isoformat() if processed_at else None,
                                "full_text": full_text,
                                "extracted_data": extracted_data
                            }
                            
                            page_data_list.append(page_data)
                            processed_documents.add(doc_id)
                            processed_pages += 1
                            
                        except Exception as e:
                            error_msg = f"Error processing page {page_id}: {str(e)}"
                            self.logger.error(error_msg)
                            errors.append(error_msg)
                    
                    # Use the EXACT same method as main pipeline
                    if page_data_list:
                        try:
                            self._store_in_vector_db(page_data_list)
                            self.logger.info(f"Added batch of {len(page_data_list)} pages to vector database")
                            print(f"Added batch of {len(page_data_list)} pages to vector database")
                            
                        except Exception as e:
                            error_msg = f"Error adding batch to vector database: {str(e)}"
                            self.logger.error(error_msg)
                            errors.append(error_msg)
            
            # Prepare results
            result = {
                "status": "success" if not errors or len(errors) < total_pages * 0.1 else "partial_success",
                "message": f"Reprocessed {processed_pages}/{total_pages} pages",
                "pages_processed": processed_pages,
                "documents_processed": len(processed_documents),
                "total_pages_found": total_pages,
                "errors": errors,
                "error_count": len(errors)
            }
            
            if errors:
                result["message"] += f" with {len(errors)} errors"
                
            self.logger.info(f"Vector reprocessing completed: {result['message']}")
            return result
            
        except Exception as e:
            error_msg = f"Error in vector reprocessing: {str(e)}"
            self.logger.error(error_msg)
            return {
                "status": "error",
                "message": error_msg,
                "pages_processed": 0,
                "documents_processed": 0,
                "errors": [error_msg]
            }
        
        finally:
            if sql_conn:
                sql_conn.close()


    def reprocess_all_vectors(self, batch_size: int = 50, force_recreate: bool = False) -> Dict[str, Any]:
        """
        Convenience method to reprocess all document vectors.
        
        Args:
            batch_size: Number of pages to process in each batch
            force_recreate: If True, delete all existing vectors before recreating
            
        Returns:
            Dictionary with processing results
        """
        return self.reprocess_document_vectors(
            document_id=None,
            document_type=None,
            batch_size=batch_size,
            force_recreate=force_recreate
        )


    def reprocess_vectors_by_type(self, document_type: str, batch_size: int = 50, force_recreate: bool = False) -> Dict[str, Any]:
        """
        Convenience method to reprocess vectors for a specific document type.
        
        Args:
            document_type: Type of documents to reprocess
            batch_size: Number of pages to process in each batch
            force_recreate: If True, delete existing vectors for this type before recreating
            
        Returns:
            Dictionary with processing results
        """
        return self.reprocess_document_vectors(
            document_id=None,
            document_type=document_type,
            batch_size=batch_size,
            force_recreate=force_recreate
        )
            
    ####################################################
    # Document Summarization
    ####################################################
    def process_page_with_summary(
        self, 
        page_content: str, 
        page_data: Dict[str, Any], 
        document_type: str,
        custom_summary_instructions: str = None
    ) -> Dict[str, Any]:
        """
        Process a page and generate summaries.
        
        This method would be called during your existing document processing pipeline
        after text extraction but before database storage.
        
        Args:
            page_content: Extracted text content of the page
            page_data: Existing page data dictionary
            document_type: Type of document being processed
            custom_summary_instructions: Optional custom instructions
            
        Returns:
            Enhanced page data with summary information
        """
        enhanced_page_data = page_data.copy()
        
        if not self.enable_summarization:
            return enhanced_page_data
        
        try:
            # Generate summaries for each configured type
            summaries = {}
            
            for summary_type in self.default_summary_types:
                self.logger.info(f"Generating {summary_type} summary for page {page_data.get('page_number', 'unknown')}")
                
                summary_data = self.summarizer.summarize_page(
                    page_content=page_content,
                    document_type=document_type,
                    summary_type=summary_type,
                    custom_instructions=custom_summary_instructions
                )
                
                summaries[summary_type] = summary_data
            
            # Add summaries to page data
            enhanced_page_data['summaries'] = summaries
            
            return enhanced_page_data
            
        except Exception as e:
            self.logger.error(f"Error processing page summaries: {str(e)}")
            return enhanced_page_data
    
    def save_page_summaries(self, page_data: Dict[str, Any]) -> bool:
        """
        Save all page summaries to the database.
        
        Args:
            page_data: Page data containing summaries
            
        Returns:
            True if all summaries saved successfully
        """
        if not self.enable_summarization or 'summaries' not in page_data:
            return True
        
        success = True
        page_id = page_data.get('page_id')
        document_id = page_data.get('document_id')
        page_number = page_data.get('page_number')
        
        for summary_type, summary_data in page_data['summaries'].items():
            if not self.summarizer.save_page_summary(
                page_id=page_id,
                document_id=document_id,
                page_number=page_number,
                summary_data=summary_data
            ):
                success = False
                self.logger.error(f"Failed to save {summary_type} summary for page {page_id}")
        
        return success
